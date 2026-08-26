"""In-browser preview pages for generated CAISO packet files.

Every document type gets an inline preview: PDFs are embedded, spreadsheets are
rendered as tables, KMZ boundaries are drawn as a map figure, PSLF model files
get a syntax-styled code view, and markdown is converted to HTML.
"""

from __future__ import annotations

import html
import re
import zipfile
from pathlib import Path
from typing import Any

PAGE = """<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8" />
<meta name="viewport" content="width=device-width, initial-scale=1" />
<title>{title} — GridPilot</title>
<link rel="preconnect" href="https://fonts.googleapis.com" />
<link href="https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600&family=JetBrains+Mono:wght@400;500&display=swap" rel="stylesheet" />
<style>
  :root {{
    --ink: #1c211e; --muted: #6b7370; --line: #dfe3df; --soft: #f4f5f3;
    --accent: #2159c0; --ok: #1d7a44; --warn-bg: #fdf6e3; --canvas: #fafbf9;
  }}
  * {{ box-sizing: border-box; }}
  body {{ margin: 0; background: var(--canvas); color: var(--ink);
    font: 14px/1.55 Inter, system-ui, sans-serif; }}
  header {{ position: sticky; top: 0; z-index: 5; display: flex; flex-wrap: wrap;
    align-items: center; gap: 10px 16px; padding: 12px 22px; background: #fffffff2;
    border-bottom: 1px solid var(--line); backdrop-filter: blur(6px); }}
  header .t {{ min-width: 0; flex: 1; }}
  header h1 {{ margin: 0; font-size: 15px; font-weight: 600; letter-spacing: -0.01em; }}
  header .f {{ font-family: "JetBrains Mono", monospace; font-size: 11px; color: var(--muted); }}
  .chip {{ display: inline-block; padding: 2px 9px; border: 1px solid var(--line);
    border-radius: 99px; font-family: "JetBrains Mono", monospace; font-size: 10px;
    text-transform: uppercase; letter-spacing: 0.06em; color: var(--muted); background: var(--soft); }}
  .chip.ok {{ color: var(--ok); border-color: #bcd9c6; background: #eef7f0; }}
  .btn {{ display: inline-block; padding: 7px 16px; border-radius: 99px; text-decoration: none;
    font-family: "JetBrains Mono", monospace; font-size: 11px; text-transform: uppercase;
    letter-spacing: 0.07em; }}
  .btn.primary {{ background: var(--ink); color: #fff; }}
  .btn.ghost {{ border: 1px solid var(--line); color: var(--ink); background: #fff; }}
  main {{ max-width: 1060px; margin: 0 auto; padding: 22px; }}
  .note {{ margin: 0 0 16px; padding: 10px 14px; border: 1px solid var(--line);
    border-radius: 10px; background: var(--soft); color: var(--muted); font-size: 12.5px; }}
  iframe.pdf {{ width: 100%; height: calc(100vh - 130px); border: 1px solid var(--line);
    border-radius: 10px; background: #fff; }}
  h2.sheet {{ margin: 26px 0 10px; font-size: 13px; text-transform: uppercase;
    letter-spacing: 0.08em; color: var(--accent); font-family: "JetBrains Mono", monospace; }}
  table {{ width: 100%; border-collapse: collapse; background: #fff; border: 1px solid var(--line);
    border-radius: 10px; overflow: hidden; }}
  td, th {{ padding: 7px 12px; border-bottom: 1px solid var(--line); text-align: left;
    vertical-align: top; font-size: 13px; }}
  tr:last-child td {{ border-bottom: none; }}
  tr.section td {{ background: var(--soft); font-weight: 600; font-size: 11.5px;
    text-transform: uppercase; letter-spacing: 0.05em; color: var(--accent); }}
  td.k {{ color: var(--muted); width: 34%; }}
  td.c {{ color: var(--muted); font-size: 12px; }}
  th {{ background: var(--soft); font-size: 11.5px; text-transform: uppercase;
    letter-spacing: 0.05em; color: var(--muted); }}
  pre.code {{ margin: 0; padding: 18px; background: #14211a; color: #d8e6dc; border-radius: 10px;
    overflow-x: auto; font: 12px/1.7 "JetBrains Mono", monospace; }}
  pre.code .cm {{ color: #6f8f7c; }}
  pre.code .kw {{ color: #8ec5ff; }}
  pre.code .num {{ color: #ffd08a; }}
  .md h1 {{ font-size: 22px; letter-spacing: -0.01em; margin: 0 0 14px; }}
  .md h2 {{ font-size: 16px; margin: 24px 0 10px; }}
  .md p {{ margin: 0 0 10px; }}
  .md li {{ margin: 3px 0; }}
  figure.map {{ margin: 0 0 18px; padding: 16px; background: #fff; border: 1px solid var(--line);
    border-radius: 10px; }}
  figure.map > svg {{ width: 100%; height: auto; display: block; }}
  figcaption {{ margin-top: 10px; font-size: 12px; color: var(--muted); }}
  .grid2 {{ display: grid; gap: 14px; grid-template-columns: 1fr; }}
  @media (min-width: 860px) {{ .grid2 {{ grid-template-columns: 3fr 2fr; }} }}
  .sum {{ display: flex; flex-wrap: wrap; gap: 10px; margin: 0 0 16px; }}
  .sum .card {{ flex: 1 1 150px; padding: 10px 14px; background: #fff;
    border: 1px solid var(--line); border-radius: 10px; }}
  .sum .card b {{ display: block; font-size: 16px; letter-spacing: -0.01em; }}
  .docx {{ background: #fff; border: 1px solid var(--line); border-radius: 10px;
    padding: 46px 54px; max-width: 860px; margin: 0 auto;
    font: 13.5px/1.6 Arial, Helvetica, sans-serif; }}
  .docx p {{ margin: 0 0 8px; white-space: pre-wrap; }}
  .docx p.gap {{ margin: 0 0 4px; }}
  .docx .tab {{ display: inline-block; width: 28px; }}
  .docx .cbx {{ color: var(--muted); }}
  .docx .cbx.on {{ color: var(--ok); font-weight: 700; }}
  .docx .fillin {{ color: var(--accent); font-weight: 600; text-decoration: underline;
    text-underline-offset: 3px; }}
  .docx .cell {{ display: inline-block; min-width: 46px; padding: 0 4px;
    border: 1px solid #3a3f3c; line-height: 1.35; text-align: center; }}
  .docx table {{ margin: 10px 0 14px; border-radius: 0; }}
  .docx td, .docx th {{ border: 1px solid var(--line); font-size: 12px; }}
  .docx .hd {{ display: flex; align-items: center; justify-content: space-between;
    gap: 16px; margin: 0 0 22px; padding-bottom: 14px; border-bottom: 1px solid var(--line); }}
  .docx .hd img {{ height: 34px; width: auto; }}
  .docx .hd .t {{ text-align: right; font-size: 12.5px; font-weight: 700; line-height: 1.45; }}
  .docx .ft {{ margin-top: 26px; padding-top: 10px; border-top: 1px solid var(--line);
    font-size: 11px; color: var(--muted); }}
  .sum .card span {{ font-family: "JetBrains Mono", monospace; font-size: 10px;
    text-transform: uppercase; letter-spacing: 0.07em; color: var(--muted); }}
  figure.sld {{ margin: 0 0 18px; padding: 16px; background: #fff;
    border: 1px solid var(--line); border-radius: 10px; }}
  figure.sld > svg {{ width: 100%; height: auto; display: block; }}
  details.src {{ margin-top: 18px; }}
  details.src summary {{ cursor: pointer; font-family: "JetBrains Mono", monospace;
    font-size: 11px; text-transform: uppercase; letter-spacing: 0.07em; color: var(--muted);
    padding: 6px 0; }}
  details.src pre.code {{ margin-top: 8px; }}
  .rec {{ margin: 0 0 12px; padding: 14px 16px; background: #fff;
    border: 1px solid var(--line); border-radius: 10px; }}
  .rec .head {{ display: flex; flex-wrap: wrap; align-items: baseline; gap: 8px 12px;
    margin-bottom: 9px; }}
  .rec .model {{ font-family: "JetBrains Mono", monospace; font-size: 12px; font-weight: 600;
    color: var(--accent); text-transform: uppercase; }}
  .rec .role {{ font-size: 12px; color: var(--muted); }}
  .rec .bus {{ margin-left: auto; font-family: "JetBrains Mono", monospace; font-size: 11px;
    color: var(--muted); }}
  .chips {{ display: flex; flex-wrap: wrap; gap: 6px; }}
  .chips code {{ padding: 2px 9px; background: var(--soft); border: 1px solid var(--line);
    border-radius: 6px; font-family: "JetBrains Mono", monospace; font-size: 11px; }}
  .chips code b {{ color: var(--accent); font-weight: 500; }}
  td.num {{ font-family: "JetBrains Mono", monospace; white-space: nowrap; font-size: 12px; }}
  .lgd {{ display: flex; flex-wrap: wrap; gap: 8px 18px; margin: 10px 0 0;
    font-size: 11.5px; color: var(--muted); }}
  .lgd i {{ display: inline-block; width: 22px; border-top: 2px solid var(--ink);
    vertical-align: middle; margin-right: 6px; }}
  .lgd i.thick {{ border-top-width: 4px; }}
  .lgd i.dash {{ border-top-style: dashed; }}
</style>
</head>
<body>
<header>
  <div class="t">
    <h1>{title}</h1>
    <div class="f">{filename} · {project}</div>
  </div>
  <span class="chip {chip_class}">{status_label}</span>
  {actions}
</header>
<main>{body}</main>
</body>
</html>"""


def _e(v: Any) -> str:
    return html.escape(str(v if v is not None else ""))


# ---------------------------------------------------------------------------
# Renderers by type
# ---------------------------------------------------------------------------

def _pdf_body(inline_url: str) -> str:
    return f'<iframe class="pdf" src="{_e(inline_url)}" title="PDF preview"></iframe>'


def _xlsx_body(path: Path) -> str:
    import openpyxl

    wb = openpyxl.load_workbook(path, data_only=True)
    parts: list[str] = []
    for ws in wb.worksheets:
        rows = []
        for row in ws.iter_rows(values_only=True):
            cells = ["" if v is None else str(v) for v in row]
            while cells and cells[-1] == "":
                cells.pop()
            if cells:
                rows.append(cells)
        if not rows:
            continue
        width = max(len(r) for r in rows)
        body_rows = []
        for r in rows:
            padded = r + [""] * (width - len(r))
            # Single-cell rows read as section banners.
            if sum(1 for c in padded if c.strip()) == 1 and padded[0].strip():
                body_rows.append(
                    f'<tr class="section"><td colspan="{width}">{_e(padded[0])}</td></tr>'
                )
                continue
            tds = []
            for i, c in enumerate(padded):
                cls = "k" if i == 0 else ("c" if i == width - 1 and width > 2 else "")
                tds.append(f'<td class="{cls}">{_e(c)}</td>')
            body_rows.append(f"<tr>{''.join(tds)}</tr>")
        parts.append(f'<h2 class="sheet">{_e(ws.title)}</h2><table>{"".join(body_rows)}</table>')
    parts.append(
        '<p class="note" style="margin-top:16px">Transfer this data into CAISO\'s official '
        "Attachment A (.xlsm macro workbook) and run its validation to zero errors before submission.</p>"
    )
    return "".join(parts)


def _docx_body(path: Path) -> str:
    """Word-form preview: renders the filled document faithfully — paragraphs
    and tables in original order, with checked boxes and filled entries
    highlighted."""
    import docx
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    doc = docx.Document(str(path))

    from docx.oxml.ns import qn as _qn

    from docx.text.run import Run

    def _run_html(r) -> str:
        rpr = r._element.find(_qn("w:rPr"))
        style = rpr.find(_qn("w:rStyle")) if rpr is not None else None
        if style is not None and style.get(_qn("w:val")) == "PlaceholderText":
            # Unfilled content control — show an empty form cell.
            return '<span class="cell">&nbsp;&nbsp;&nbsp;&nbsp;</span>'
        t = _e(r.text).replace("\t", '<span class="tab"></span>')
        t = t.replace("☒", '<span class="cbx on">☒</span>')
        t = t.replace("☐", '<span class="cbx">☐</span>')
        if r.underline:
            t = f'<span class="fillin">{t}</span>'
        if r.bold:
            t = f"<b>{t}</b>"
        if rpr is not None and rpr.find(_qn("w:bdr")) is not None:
            t = f'<span class="cell">{t}</span>'
        return t

    def para_html(par) -> str:
        out = []
        for child in par._element:
            if child.tag == _qn("w:r"):
                out.append(_run_html(Run(child, par)))
            elif child.tag == _qn("w:sdt"):
                # Content controls: render the sdtContent runs inline.
                content = child.find(_qn("w:sdtContent"))
                if content is not None:
                    out.extend(_run_html(Run(rl, par))
                               for rl in content.findall(_qn("w:r")))
        return "".join(out)

    def _header_html() -> str:
        """Section header: text lines plus any embedded image (CAISO logo),
        rendered as a page-top strip like Word shows it."""
        import base64 as b64mod

        from docx.oxml.ns import qn

        hdr = doc.sections[0].header
        img_html = ""
        blips = hdr._element.findall(".//" + qn("a:blip"))
        if blips:
            rid = blips[0].get(qn("r:embed"))
            part = hdr.part.related_parts.get(rid)
            if part is not None:
                data = b64mod.b64encode(part.blob).decode()
                img_html = f'<img src="data:{part.content_type};base64,{data}" alt="">'
        lines = [para_html(p) for p in hdr.paragraphs if p.text.strip()]
        if not (lines or img_html):
            return ""
        return (f'<div class="hd">{img_html}<div class="t">'
                + "<br>".join(lines) + "</div></div>")

    def _footer_html() -> str:
        ftr = doc.sections[0].footer
        lines = []
        for p in ftr.paragraphs:
            t = " ".join(p.text.split())
            if not t:
                continue
            # Page-number fields are computed by Word; show a static stand-in.
            if t.startswith("Page"):
                t = "Page 1"
            lines.append(_e(t))
        return f'<div class="ft">{" · ".join(lines)}</div>' if lines else ""

    parts: list[str] = ['<div class="docx">', _header_html()]
    try:
        blocks = list(doc.iter_inner_content())
    except AttributeError:
        blocks = list(doc.paragraphs) + list(doc.tables)
    for block in blocks:
        if isinstance(block, Paragraph):
            html = para_html(block)
            if html.strip():
                cls = ' class="head"' if (block.style and "Heading" in (block.style.name or "")) else ""
                parts.append(f"<p{cls}>{html}</p>")
            else:
                parts.append('<p class="gap"></p>')
        elif isinstance(block, Table):
            rows_html = []
            for ri, row in enumerate(block.rows):
                tag = "th" if ri == 0 else "td"
                cells = "".join(
                    f"<{tag}>{'<br>'.join(para_html(cp) for cp in c.paragraphs if cp.text.strip() or cp.runs)}</{tag}>"
                    for c in row.cells)
                rows_html.append(f"<tr>{cells}</tr>")
            parts.append(f"<table>{''.join(rows_html)}</table>")
    parts.append(_footer_html())
    parts.append("</div>")
    parts.append(
        '<p class="note" style="margin-top:16px">This is the official CAISO Word form '
        "filled in place — download the .docx for the submission copy and execute it "
        "electronically in RIMS5.</p>"
    )
    return "".join(parts)


def _code_pre(text: str) -> str:
    lines_out = []
    for line in text.splitlines():
        esc = _e(line)
        if line.lstrip().startswith(("!", "#", "/")):
            lines_out.append(f'<span class="cm">{esc}</span>')
        else:
            esc = re.sub(r"^(\w[\w_]*)", r'<span class="kw">\1</span>', esc)
            lines_out.append(esc)
    return f'<pre class="code">{chr(10).join(lines_out)}</pre>'


def _source_details(text: str, label: str = "View file source") -> str:
    return f'<details class="src"><summary>{_e(label)}</summary>{_code_pre(text)}</details>'


def _code_body(text: str, kind: str) -> str:
    label = {
        ".epc": "GE PSLF load-flow model (.epc) — steady-state buses, branches, transformers, generators.",
        ".dyd": "GE PSLF dynamic model (.dyd) — WECC REGC_A / REEC / REPC_A controller records.",
        ".raw": "Siemens PSS/E load-flow model (.raw) — steady-state buses, branches, transformers, generators.",
        ".dyr": "Siemens PSS/E dynamic model (.dyr) — standard-library controller records.",
    }.get(kind, "Plain-text model file.")
    return (
        f'<p class="note">{_e(label)} Validate against the current ISO base case '
        f'before submission.</p>{_code_pre(text)}'
    )


# ---------------------------------------------------------------------------
# DXF — parse the R12 entity stream and render the actual drawing as SVG
# ---------------------------------------------------------------------------

def _parse_dxf_entities(text: str) -> list[dict[str, Any]]:
    lines = text.splitlines()
    pairs = [(lines[i].strip(), lines[i + 1]) for i in range(0, len(lines) - 1, 2)]
    ents: list[dict[str, Any]] = []
    section = None
    cur: dict[str, Any] | None = None
    expect_section_name = False
    for code, val in pairs:
        if expect_section_name:
            if code == "2":
                section = val.strip()
            expect_section_name = False
            continue
        if code == "0":
            if cur is not None:
                ents.append(cur)
                cur = None
            v = val.strip()
            if v == "SECTION":
                expect_section_name = True
            elif v == "ENDSEC":
                section = None
            elif section == "ENTITIES" and v in ("LINE", "CIRCLE", "TEXT"):
                cur = {"type": v}
            continue
        if cur is not None:
            cur[code] = val
    if cur is not None:
        ents.append(cur)
    return ents


def _dxf_body(text: str) -> str:
    ents = _parse_dxf_entities(text)
    if not ents:
        return _code_body(text, ".dxf")

    def f(ent: dict, code: str, default: float = 0.0) -> float:
        try:
            return float(ent.get(code, default))
        except (TypeError, ValueError):
            return default

    # Bounding box (DXF y-up); flip to screen coordinates for display.
    xs: list[float] = []
    ys: list[float] = []
    for en in ents:
        if en["type"] == "LINE":
            xs += [f(en, "10"), f(en, "11")]
            ys += [f(en, "20"), f(en, "21")]
        elif en["type"] == "CIRCLE":
            r = f(en, "40")
            xs += [f(en, "10") - r, f(en, "10") + r]
            ys += [f(en, "20") - r, f(en, "20") + r]
        elif en["type"] == "TEXT":
            xs.append(f(en, "10"))
            ys.append(f(en, "20"))
    if not xs:
        return _code_body(text, ".dxf")
    pad = 12.0
    x0, x1 = min(xs) - pad, max(xs) + pad
    y0, y1 = min(ys) - pad, max(ys) + pad

    def fy(y: float) -> float:
        return (y1 + y0) - y  # restore screen orientation inside the bbox

    svg = [f'<svg xmlns="http://www.w3.org/2000/svg" viewBox="{x0:.1f} {y0:.1f} '
           f'{x1 - x0:.1f} {y1 - y0:.1f}" font-family="Helvetica, Arial, sans-serif">']
    counts = {"LINE": 0, "CIRCLE": 0, "TEXT": 0}
    layers: dict[str, int] = {}
    for en in ents:
        layer = str(en.get("8", "SLD")).strip()
        layers[layer] = layers.get(layer, 0) + 1
        counts[en["type"]] = counts.get(en["type"], 0) + 1
        if en["type"] == "LINE":
            sw = 3.2 if layer == "SLD-BUS" else 1.1
            dash = ' stroke-dasharray="7 5"' if layer == "SLD-BOUNDARY" else ""
            svg.append(f'<line x1="{f(en, "10"):.1f}" y1="{fy(f(en, "20")):.1f}" '
                       f'x2="{f(en, "11"):.1f}" y2="{fy(f(en, "21")):.1f}" '
                       f'stroke="#111" stroke-width="{sw}"{dash}/>')
        elif en["type"] == "CIRCLE":
            svg.append(f'<circle cx="{f(en, "10"):.1f}" cy="{fy(f(en, "20")):.1f}" '
                       f'r="{f(en, "40"):.1f}" stroke="#111" stroke-width="1.1" fill="none"/>')
        elif en["type"] == "TEXT":
            anchor = {"1": "middle", "2": "end"}.get(str(en.get("72", "")).strip(), "start")
            svg.append(f'<text x="{f(en, "10"):.1f}" y="{fy(f(en, "20")):.1f}" '
                       f'font-size="{f(en, "40", 8.0):.1f}" text-anchor="{anchor}" '
                       f'fill="#111">{_e(str(en.get("1", "")))}</text>')
    svg.append("</svg>")

    layer_rows = "".join(
        f'<tr><td class="k">{_e(name)}</td><td class="num">{n} entities</td></tr>'
        for name, n in sorted(layers.items())
    )
    return f"""
    <p class="note">CAD-editable DXF (R12) rendered directly from the file's entity stream —
    the same drawing your CAD tool will open. Layers carry the bus work, symbols, text, and
    ownership boundary separately.</p>
    <div class="sum">
      <div class="card"><b>{counts.get("LINE", 0)}</b><span>Line entities</span></div>
      <div class="card"><b>{counts.get("CIRCLE", 0)}</b><span>Circle entities</span></div>
      <div class="card"><b>{counts.get("TEXT", 0)}</b><span>Text entities</span></div>
      <div class="card"><b>{len(layers)}</b><span>CAD layers</span></div>
    </div>
    <div class="grid2">
      <figure class="sld">{"".join(svg)}
        <figcaption>Rendered from the DXF LINE / CIRCLE / TEXT entities. Open in AutoCAD-class
        tools for editing; the PDF sheet in this packet is the stamped-format equivalent.</figcaption>
      </figure>
      <div>
        <table><tr><th colspan="2">CAD layers</th></tr>{layer_rows}</table>
        <div class="lgd" style="margin-top:12px">
          <span><i class="thick"></i>SLD-BUS — buses</span>
          <span><i></i>SLD — equipment</span>
          <span><i class="dash"></i>SLD-BOUNDARY — ownership</span>
        </div>
      </div>
    </div>
    {_source_details(text, "View DXF source (R12 tagged data)")}"""


# ---------------------------------------------------------------------------
# PSLF .epc — solved-case dashboard with parsed data tables
# ---------------------------------------------------------------------------

_EPC_BUS_RE = re.compile(
    r'^\s*(\d+)\s+"([^"]+)"\s+([\d.]+)\s*:\s*#9\s+(\d)\s+([\d.]+)\s+([-\d.]+)\s*:')
_EPC_BRANCH_RE = re.compile(
    r'^\s*(\d+)\s+"([^"]+)"\s+[\d.]+\s+(\d+)\s+"([^"]+)"\s+[\d.]+\s+"\s*(\d+)"\s*:\s*'
    r'([-\d.]+)\s+([-\d.]+)\s+([-\d.]+)\s*:\s*([\d.]+)')
_EPC_XFMR_RE = re.compile(
    r'^\s*(\d+)\s+"([^"]+)"\s+[\d.]+\s+(\d+)\s+"([^"]+)"\s+[\d.]+\s+"\s*(\d+)"\s*:\s*'
    r'([-\d.]+)\s+([-\d.]+)\s*:\s*([\d.]+)')
_EPC_GEN_RE = re.compile(
    r'^\s*(\d+)\s+"([^"]+)"\s+([\d.]+)\s+"\s*(\d+)"\s*:\s*([-\d.]+)\s+([-\d.]+)\s+'
    r'([-\d.]+)\s+(-[\d.]+)\s*:\s*([\d.]+)')
_EPC_LOAD_RE = re.compile(
    r'^\s*(\d+)\s+"([^"]+)"\s+([\d.]+)\s+"\s*(\d+)"\s*:\s*([-\d.]+)\s+([-\d.]+)')

_EPC_BUS_TYPE = {"0": "Swing (utility)", "1": "PQ", "2": "Generator"}


def _epc_sections(text: str) -> tuple[list[str], dict[str, list[str]]]:
    comments: list[str] = []
    sections: dict[str, list[str]] = {}
    cur: str | None = None
    for line in text.splitlines():
        s = line.strip()
        if not s:
            continue
        if s.startswith("!"):
            comments.append(s.lstrip("! "))
            continue
        low = s.lower()
        if low in ("title", "comments") or low.endswith(" data"):
            cur = low
            sections[cur] = []
            continue
        if low == "end":
            cur = None
            continue
        if cur is not None:
            sections[cur].append(line)
    return comments, sections


def _rows_table(headers: list[str], rows: list[list[str]], num_from: int = 2) -> str:
    if not rows:
        return ""
    head = "".join(f"<th>{_e(h)}</th>" for h in headers)
    body = "".join(
        "<tr>" + "".join(
            f'<td class="{"num" if i >= num_from else ""}">{_e(c)}</td>'
            for i, c in enumerate(r)
        ) + "</tr>"
        for r in rows
    )
    return f"<table><tr>{head}</tr>{body}</table>"


def _epc_body(text: str) -> str:
    comments, sec = _epc_sections(text)

    # Header cards from the writer's solved-case comments.
    iters = losses = dispatch = ""
    m = re.search(r"(\d+)\s+iterations", " ".join(comments))
    if m:
        iters = m.group(1)
    m = re.search(r"losses\s+([\d.]+)\s*MW", " ".join(comments))
    if m:
        losses = m.group(1)
    m = re.search(r"Dispatch:\s*(.+?)\s*(?:$|—)", next((c for c in comments if "Dispatch" in c), ""))
    if m:
        dispatch = m.group(1)

    parts: list[str] = ['<p class="note">GE PSLF steady-state model rendered from the file '
                        "contents — every number below traces to the solved load-flow case, the "
                        "project graph, or an OEM datasheet. Validate against the current ISO "
                        "base case before submission.</p>"]
    cards = []
    if iters:
        cards.append(f'<div class="card"><b>{_e(iters)}</b><span>Solver iterations</span></div>')
    if losses:
        cards.append(f'<div class="card"><b>{_e(losses)} MW</b><span>Series losses</span></div>')
    buses = sec.get("bus data", [])
    cards.append(f'<div class="card"><b>{len(buses)}</b><span>Buses</span></div>')
    cards.append(f'<div class="card"><b>{len(sec.get("generator data", []))}</b><span>Generators</span></div>')
    parts.append(f'<div class="sum">{"".join(cards)}</div>')
    if dispatch:
        parts.append(f'<p class="note">Dispatch: {_e(dispatch)}</p>')
    if sec.get("title"):
        parts.append(f'<h2 class="sheet">Case title</h2><p>{_e(" ".join(sec["title"]))}</p>')

    rows = []
    for line in buses:
        m = _EPC_BUS_RE.match(line)
        if m:
            num, name, kv, btype, v, ang = m.groups()
            rows.append([num, name.strip(), f"{float(kv):g}", _EPC_BUS_TYPE.get(btype, btype),
                         v, ang])
    if rows:
        parts.append('<h2 class="sheet">Bus data — solved voltages</h2>')
        parts.append(_rows_table(["Bus", "Name", "kV", "Type", "V (pu)", "Angle (°)"], rows))

    rows = []
    for line in sec.get("branch data", []):
        m = _EPC_BRANCH_RE.match(line)
        if m:
            f_, fn, t_, tn, cid, r, x, b, rate = m.groups()
            rows.append([f"{fn.strip()} → {tn.strip()}", cid, r, x, b, f"{float(rate):g}"])
    if rows:
        parts.append('<h2 class="sheet">Branch data</h2>')
        parts.append(_rows_table(["Branch", "ID", "R (pu)", "X (pu)", "B (pu)", "Rating (MVA)"],
                                 rows, num_from=1))

    rows = []
    for line in sec.get("transformer data", []):
        m = _EPC_XFMR_RE.match(line)
        if m:
            f_, fn, t_, tn, cid, r, x, mva = m.groups()
            rows.append([f"{fn.strip()} → {tn.strip()}", cid, r, x, f"{float(mva):g}"])
    if rows:
        parts.append('<h2 class="sheet">Transformer data</h2>')
        parts.append(_rows_table(["Transformer", "ID", "R (pu)", "X (pu)", "Rating (MVA)"],
                                 rows, num_from=1))

    rows = []
    for line in sec.get("generator data", []):
        m = _EPC_GEN_RE.match(line)
        if m:
            bus, name, kv, cid, p, q, qmax, qmin, mva = m.groups()
            rows.append([name.strip(), f"{float(kv):g}", p, q, qmax, qmin, f"{float(mva):g}"])
    if rows:
        parts.append('<h2 class="sheet">Generator dispatch — from the solved case</h2>')
        parts.append(_rows_table(
            ["Unit", "kV", "P (MW)", "Q (Mvar)", "Qmax", "Qmin", "MVA base"], rows, num_from=1))

    rows = []
    for line in sec.get("load data", []):
        m = _EPC_LOAD_RE.match(line)
        if m:
            bus, name, kv, cid, p, q = m.groups()
            rows.append([name.strip(), f"{float(kv):g}", p, q])
    if rows:
        parts.append('<h2 class="sheet">Load data — auxiliary / station service</h2>')
        parts.append(_rows_table(["Bus", "kV", "P (MW)", "Q (Mvar)"], rows, num_from=1))

    parts.append(_source_details(text, "View .epc source (as submitted)"))
    return "".join(parts)


# ---------------------------------------------------------------------------
# PSLF .dyd / PSS/E .dyr — WECC controller records as cards
# ---------------------------------------------------------------------------

_WECC_ROLE = {
    "regc": "Converter interface — current injection",
    "reec": "Electrical controls — V/Q and P priority",
    "repc": "Plant controller — POI regulation & droop",
}


def _model_role(model: str) -> str:
    key = model.lower()[:4]
    return _WECC_ROLE.get(key, "Controller record")


# Record grammar shared with the engine's vendor-file parser.
from backend.app.engine.vendor_models import DYD_PARAM_RE as _DYD_PARAM_RE
from backend.app.engine.vendor_models import DYD_REC_RE as _DYD_REC_RE


def _dyd_body(text: str) -> str:
    parts: list[str] = ['<p class="note">GE PSLF dynamic data rendered from the file contents — '
                        "WECC standard-library models with the parameter blocks the packet "
                        "submits. Generic blocks are flagged for replacement with OEM "
                        "MOD-026/027 data.</p>"]
    for line in text.splitlines():
        s = line.strip()
        if not s:
            continue
        if s.startswith("#"):
            note = s.lstrip("# ")
            warn_style = ' style="border-color:#f59e0b66;background:#fffbeb;color:#92400e"'
            style = warn_style if "GENERIC" in note.upper() else ""
            parts.append(f'<p class="note"{style}>{_e(note)}</p>')
            continue
        m = _DYD_REC_RE.match(s)
        if not m:
            continue
        model, bus, name, kv, cid, params = m.groups()
        chips = []
        for pm in _DYD_PARAM_RE.finditer(params):
            k = pm.group(1) or pm.group(3)
            v = pm.group(2) or pm.group(4)
            chips.append(f"<code><b>{_e(k)}</b> = {_e(v)}</code>")
        parts.append(f"""
        <div class="rec">
          <div class="head">
            <span class="model">{_e(model.upper())}</span>
            <span class="role">{_e(_model_role(model))}</span>
            <span class="bus">bus {_e(bus)} · {_e(name.strip())} · {_e(f"{float(kv):g}")} kV · id {_e(cid)}</span>
          </div>
          <div class="chips">{"".join(chips)}</div>
        </div>""")
    parts.append(_source_details(text, "View .dyd source (as submitted)"))
    return "".join(parts)


_DYR_REC_RE = re.compile(r"^(\d+)\s+'([\w]+)'\s+(\d+)\s+(.*?)\s*/\s*$")


def _dyr_body(text: str) -> str:
    parts: list[str] = ['<p class="note">Siemens PSS/E dynamic data rendered from the file '
                        "contents — WECC standard-library models with positional parameter "
                        "lists in library order.</p>"]
    for line in text.splitlines():
        s = line.strip()
        if not s:
            continue
        if s.startswith("/"):
            parts.append(f'<p class="note">{_e(s.lstrip("/ "))}</p>')
            continue
        m = _DYR_REC_RE.match(s)
        if not m:
            continue
        bus, model, cid, params = m.groups()
        chips = "".join(f"<code>{_e(v)}</code>" for v in params.split())
        parts.append(f"""
        <div class="rec">
          <div class="head">
            <span class="model">{_e(model.upper())}</span>
            <span class="role">{_e(_model_role(model))}</span>
            <span class="bus">bus {_e(bus)} · id {_e(cid)}</span>
          </div>
          <div class="chips">{chips}</div>
        </div>""")
    parts.append(_source_details(text, "View .dyr source (as submitted)"))
    return "".join(parts)


# ---------------------------------------------------------------------------
# PSS/E .raw — section tables
# ---------------------------------------------------------------------------

def _raw_body(text: str) -> str:
    lines = text.splitlines()
    parts: list[str] = ['<p class="note">Siemens PSS/E v33 case rendered from the file '
                        "contents — solved bus voltages, dispatch, and network data. Validate "
                        "against the current ISO base case before submission.</p>"]
    sections: dict[str, list[str]] = {}
    cur = None
    for line in lines[1:]:
        s = line.strip()
        if not s or s == "Q":
            continue
        m = re.search(r"BEGIN (\w[\w ]*?) DATA", s)
        if m or "END OF" in s:
            cur = m.group(1).title() if m else None
            if cur:
                sections[cur] = []
            continue
        if cur:
            sections[cur].append(s)
        # Header comment lines before the first section read as notes.
        elif not sections and not s[0].isdigit():
            parts.append(f'<p class="note">{_e(s)}</p>')

    headers = {
        "Bus": ["Bus", "Name", "kV", "Type", "V (pu)", "Angle (°)"],
        "Load": ["Bus", "ID", "P (MW)", "Q (Mvar)"],
        "Generator": ["Bus", "ID", "P (MW)", "Q (Mvar)", "Qmax", "Qmin", "MVA base"],
        "Branch": ["From", "To", "ID", "R (pu)", "X (pu)", "B (pu)", "Rating (MVA)"],
    }
    pick = {
        "Bus": [0, 1, 2, 3, 7, 8],
        "Load": [0, 1, 5, 6],
        "Generator": [0, 1, 2, 3, 4, 5, 8],
        "Branch": [0, 1, 2, 3, 4, 5, 6],
    }
    type_map = {"3": "Swing", "2": "Generator", "1": "PQ"}
    for name, rows_raw in sections.items():
        if name in headers:
            rows = []
            for r in rows_raw:
                cells = [c.strip().strip("'").strip() for c in r.split(",")]
                sel = [cells[i] if i < len(cells) else "" for i in pick[name]]
                if name == "Bus" and len(sel) > 3:
                    sel[3] = type_map.get(sel[3], sel[3])
                rows.append(sel)
            parts.append(f'<h2 class="sheet">{_e(name)} data</h2>')
            parts.append(_rows_table(headers[name], rows, num_from=1))
        elif rows_raw:
            body = "".join(f'<tr><td class="num">{_e(r)}</td></tr>' for r in rows_raw)
            parts.append(f'<h2 class="sheet">{_e(name)} data</h2>'
                         f"<table><tr><th>Record</th></tr>{body}</table>")

    parts.append(_source_details(text, "View .raw source (as submitted)"))
    return "".join(parts)


# ---------------------------------------------------------------------------
# JSON — portal field export as tables
# ---------------------------------------------------------------------------

def _json_body(text: str) -> str:
    import json as _json

    try:
        data = _json.loads(text)
    except Exception:
        return _code_body(text, ".json")

    def kv_table(obj: dict) -> str:
        rows = []
        for k, v in obj.items():
            if isinstance(v, dict):
                v = ", ".join(f"{ik}: {iv}" for ik, iv in v.items())
            elif isinstance(v, list):
                v = "; ".join(str(i) for i in v)
            label = str(k).replace("_", " ").capitalize()
            rows.append(f'<tr><td class="k">{_e(label)}</td><td class="num">{_e(v)}</td></tr>')
        return f"<table>{''.join(rows)}</table>"

    parts: list[str] = ['<p class="note">Portal field export rendered from the JSON contents — '
                        "values keyed for direct entry into the ISO application portal, all "
                        "traced to the project graph.</p>"]
    if isinstance(data, dict):
        flat = {k: v for k, v in data.items() if not isinstance(v, (dict, list))}
        if flat:
            parts.append(kv_table(flat))
        for k, v in data.items():
            if isinstance(v, dict):
                parts.append(f'<h2 class="sheet">{_e(str(k).replace("_", " "))}</h2>')
                parts.append(kv_table(v))
            elif isinstance(v, list) and v and isinstance(v[0], dict):
                cols = list(v[0].keys())
                head = "".join(f"<th>{_e(c)}</th>" for c in cols)
                body = "".join(
                    "<tr>" + "".join(f'<td class="num">{_e(row.get(c, ""))}</td>' for c in cols) + "</tr>"
                    for row in v
                )
                parts.append(f'<h2 class="sheet">{_e(str(k).replace("_", " "))}</h2>'
                             f"<table><tr>{head}</tr>{body}</table>")
    else:
        parts.append(_code_body(text, ".json"))
    parts.append(_source_details(text, "View JSON source"))
    return "".join(parts)


def _markdown_body(text: str) -> str:
    def inline(s: str) -> str:
        s = _e(s)
        s = re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", s)
        return s

    out: list[str] = []
    lines = text.splitlines()
    i = 0
    in_list = False
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if stripped.startswith("|"):
            if in_list:
                out.append("</ul>")
                in_list = False
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append([c.strip() for c in lines[i].strip().strip("|").split("|")])
                i += 1
            header, body = rows[0], [r for r in rows[1:] if not set("".join(r)) <= set("-: ")]
            out.append("<table><tr>" + "".join(f"<th>{inline(c)}</th>" for c in header) + "</tr>")
            for r in body:
                out.append("<tr>" + "".join(f"<td>{inline(c)}</td>" for c in r) + "</tr>")
            out.append("</table>")
            continue
        if stripped.startswith("- "):
            if not in_list:
                out.append("<ul>")
                in_list = True
            out.append(f"<li>{inline(stripped[2:])}</li>")
        else:
            if in_list:
                out.append("</ul>")
                in_list = False
            if stripped.startswith("## "):
                out.append(f"<h2>{inline(stripped[3:])}</h2>")
            elif stripped.startswith("# "):
                out.append(f"<h1>{inline(stripped[2:])}</h1>")
            elif stripped:
                out.append(f"<p>{inline(stripped)}</p>")
        i += 1
    if in_list:
        out.append("</ul>")
    return f'<div class="md">{"".join(out)}</div>'


def _kmz_body(path: Path, manifest: dict) -> str:
    try:
        with zipfile.ZipFile(path) as z:
            kml = z.read("doc.kml").decode("utf-8")
    except Exception:
        return '<p class="note">Could not read KMZ contents — use Download and open in Google Earth.</p>'

    coord_blocks = re.findall(r"<coordinates>(.*?)</coordinates>", kml, re.S)
    polygon: list[tuple[float, float]] = []
    center: tuple[float, float] | None = None
    for block in coord_blocks:
        pts = []
        for token in block.split():
            parts = token.split(",")
            if len(parts) >= 2:
                pts.append((float(parts[0]), float(parts[1])))
        if len(pts) >= 4:
            polygon = pts
        elif len(pts) == 1:
            center = pts[0]

    name_m = re.search(r"<name>(.*?)</name>", kml)
    desc_m = re.search(r"<description>(.*?)</description>", kml)

    map_html = '<p class="note">No polygon found in KMZ.</p>'
    if polygon:
        # Leaflet map over satellite imagery — the polygon sits on the real terrain.
        latlngs = ",".join(f"[{la:.6f},{lo:.6f}]" for lo, la in polygon)
        center_js = f"[{center[1]:.6f},{center[0]:.6f}]" if center else "null"
        map_html = f"""
        <link rel="stylesheet" href="https://unpkg.com/leaflet@1.9.4/dist/leaflet.css" />
        <script src="https://unpkg.com/leaflet@1.9.4/dist/leaflet.js"></script>
        <div id="bdy-map" style="height:420px;border-radius:8px;border:1px solid var(--line);background:#e8ebe6"></div>
        <script>
        (function () {{
          var poly = [{latlngs}];
          var map = L.map("bdy-map", {{ scrollWheelZoom: false }});
          L.tileLayer("https://server.arcgisonline.com/ArcGIS/rest/services/World_Imagery/MapServer/tile/{{z}}/{{y}}/{{x}}", {{
            maxZoom: 19,
            attribution: "Imagery &copy; Esri, Maxar, Earthstar Geographics",
          }}).addTo(map);
          var boundary = L.polygon(poly, {{
            color: "#37d067", weight: 3, dashArray: "8 5", fillColor: "#37d067", fillOpacity: 0.12,
          }}).addTo(map);
          var c = {center_js};
          if (c) {{
            L.circleMarker(c, {{ radius: 6, color: "#fff", weight: 2, fillColor: "#2159c0", fillOpacity: 1 }})
              .addTo(map).bindTooltip("Site centroid", {{ direction: "top", offset: [0, -8] }});
          }}
          map.fitBounds(boundary.getBounds().pad(0.6));
        }})();
        </script>"""

    rows = [
        ("Placemark", name_m.group(1) if name_m else path.stem),
        ("Description", desc_m.group(1) if desc_m else "—"),
    ]
    if center:
        rows.append(("Site centroid (lat, lon)", f"{center[1]:.6f}, {center[0]:.6f}"))
    if polygon:
        rows.append(("Boundary vertices", ", ".join(f"({la:.4f}, {lo:.4f})" for lo, la in polygon[:-1])))
    table = "<table>" + "".join(
        f'<tr><td class="k">{_e(k)}</td><td>{_e(v)}</td></tr>' for k, v in rows
    ) + "</table>"

    return f"""
    <div class="grid2">
      <figure class="map">{map_html}
        <figcaption>KMZ boundary polygon over satellite terrain. Download and open in Google Earth Pro
        for the interactive view CAISO reviewers use.</figcaption>
      </figure>
      <div>{table}</div>
    </div>"""


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

def render_preview(manifest: dict, doc: dict, path: Path, download_url: str) -> str:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        body = _pdf_body(download_url)
    elif suffix == ".xlsx":
        body = _xlsx_body(path)
    elif suffix == ".docx":
        body = _docx_body(path)
    elif suffix == ".dxf":
        body = _dxf_body(path.read_text(encoding="utf-8"))
    elif suffix == ".epc":
        body = _epc_body(path.read_text(encoding="utf-8"))
    elif suffix == ".dyd":
        body = _dyd_body(path.read_text(encoding="utf-8"))
    elif suffix == ".raw":
        body = _raw_body(path.read_text(encoding="utf-8"))
    elif suffix == ".dyr":
        body = _dyr_body(path.read_text(encoding="utf-8"))
    elif suffix == ".json":
        body = _json_body(path.read_text(encoding="utf-8"))
    elif suffix == ".md":
        body = _markdown_body(path.read_text(encoding="utf-8"))
    elif suffix == ".kmz":
        body = _kmz_body(path, manifest)
    else:
        body = '<p class="note">No inline preview for this file type — use Download.</p>'

    return PAGE.format(
        title=_e(doc.get("title") or path.name),
        filename=_e(doc.get("file") or path.name),
        project=_e(manifest.get("project_name") or ""),
        status_label=_e(doc.get("status_label") or ""),
        chip_class="ok" if doc.get("status") == "generated" else "",
        actions=f'<a class="btn ghost" href="{_e(download_url)}" download>Download</a>',
        body=body,
    )


# ---------------------------------------------------------------------------
# Kickoff-document previews (the developer-provided example files in Step 2)
# ---------------------------------------------------------------------------

HL_CSS = """<style>
  mark.vhl { background: #fef3c7; box-shadow: 0 0 0 2px #fef3c7; border-bottom: 2px solid #f59e0b;
    border-radius: 2px; color: inherit; }
  .vhl-note { max-width: 720px; margin: 0 auto 14px; padding: 9px 14px; border: 1px solid #f59e0b66;
    background: #fffbeb; border-radius: 8px; font-size: 12.5px; color: #92400e; }
</style>
<script>
  addEventListener("DOMContentLoaded", () => {
    const m = document.querySelector("mark.vhl, tr.vhl-row");
    if (m) setTimeout(() => m.scrollIntoView({ block: "center", behavior: "smooth" }), 150);
  });
</script>"""


def _mark(active: bool, html: str) -> str:
    """Wrap a fragment in the validation highlight when this section is under review."""
    return f'<mark class="vhl">{html}</mark>' if active else html


def _hl_note(hl: list[str]) -> str:
    if not hl:
        return ""
    return ('<div class="vhl-note">Highlighted below: the part of this document examined by the '
            "validation check that opened this preview.</div>")


DOC_CSS = """<style>
  .paper { max-width: 720px; margin: 0 auto; padding: 42px 48px; background: #fff;
    border: 1px solid var(--line); border-radius: 10px; box-shadow: 0 1px 8px rgba(0,0,0,0.04); }
  .paper h1 { font-size: 17px; text-align: center; letter-spacing: 0.02em; margin: 0 0 4px;
    text-transform: uppercase; }
  .paper .sub { text-align: center; font-size: 12px; color: var(--muted); margin: 0 0 26px; }
  .paper h2 { font-size: 12px; text-transform: uppercase; letter-spacing: 0.07em;
    color: var(--accent); margin: 24px 0 8px; font-family: "JetBrains Mono", monospace; }
  .paper p { margin: 0 0 10px; font-size: 13.5px; }
  .paper .sig { display: grid; gap: 28px; grid-template-columns: 1fr 1fr; margin-top: 34px; }
  .paper .sig div { border-top: 1px solid var(--ink); padding-top: 6px; font-size: 12px;
    color: var(--muted); }
  .paper .sig strong { display: block; color: var(--ink); font-size: 13px; }
  .stamp { float: right; margin: -6px 0 10px 14px; padding: 5px 12px; border: 2px solid #b23b3b;
    border-radius: 6px; color: #b23b3b; font-family: "JetBrains Mono", monospace; font-size: 10px;
    text-transform: uppercase; letter-spacing: 0.1em; transform: rotate(2deg); }
</style>"""


def _lease_body(intake: dict, hl: list[str] | None = None) -> str:
    hl = hl or []
    return DOC_CSS + HL_CSS + f"""
    <p class="note">Example document preloaded by the demo — in production this is the developer's own
    executed agreement, uploaded at kickoff.</p>
    {_hl_note(hl)}
    <div class="paper">
      <span class="stamp">Executed copy</span>
      <h1>Solar Ground Lease Agreement</h1>
      <p class="sub">Dated as of March 14, 2026</p>
      <h2>1 · Parties</h2>
      <p>This Ground Lease ("Lease") is entered into between <strong>{_e(intake.get('site_owner'))}</strong>,
      a California limited partnership ("Lessor"), and {_mark('legal_name' in hl,
      f"<strong>{_e(intake.get('legal_name'))}</strong>, a {_e(intake.get('state_of_origin'))} limited liability company")}
      ("Lessee").</p>
      <h2>2 · Premises</h2>
      <p>Approximately <strong>{_e(intake.get('site_acreage'))} acres</strong> of real property in
      {_e(intake.get('county'))} County, {_e(intake.get('state'))}, centered near
      {_mark('gps' in hl, f"{_e(intake.get('gps_lat'))}, {_e(intake.get('gps_lon'))}")}, as further described in Exhibit A
      (legal description) and depicted in Exhibit B (parcel map).</p>
      <h2>3 · Purpose and exclusivity</h2>
      <p>{_mark('exclusivity' in hl,
      "Lessee shall have the <strong>exclusive right</strong> to use the Premises for the development, "
      "construction, and operation of a solar photovoltaic and battery energy storage facility, including "
      "interconnection facilities. Lessor shall not grant any interest in the Premises inconsistent with "
      "Lessee's exclusive rights during the Term.")}</p>
      <h2>4 · Term</h2>
      <p>A development term of five (5) years from the Effective Date, followed upon commercial operation
      by an operating term of thirty (30) years with two (2) five-year extension options.</p>
      <h2>5 · Assignment for interconnection</h2>
      <p>{_mark('exclusivity' in hl,
      "Lessor acknowledges that Lessee may present this Lease to the California Independent System "
      "Operator as evidence of site exclusivity in connection with its interconnection request.")}</p>
      <div class="sig">
        <div><strong>{_e(intake.get('site_owner'))}</strong>By: T. Willow, General Partner<br/>Date: 03/14/2026</div>
        <div><strong>{_e(intake.get('legal_name'))}</strong>By: {_e(intake.get('signatory_name'))}, {_e(intake.get('signatory_title'))}<br/>Date: 03/14/2026</div>
      </div>
    </div>"""


def _signatory_body(intake: dict, hl: list[str] | None = None) -> str:
    hl = hl or []
    return DOC_CSS + HL_CSS + f"""
    <p class="note">Example document preloaded by the demo — in production this is the developer's own
    officer certificate or board resolution.</p>
    {_hl_note(hl)}
    <div class="paper">
      <h1>Certificate of Authorized Signatory</h1>
      <p class="sub">{_mark('legal_name' in hl,
      f"{_e(intake.get('legal_name'))} — a {_e(intake.get('state_of_origin'))} limited liability company")}</p>
      <p>The undersigned, being the Secretary of {_mark('legal_name' in hl,
      f"<strong>{_e(intake.get('legal_name'))}</strong>")} (the
      "Company"), hereby certifies that:</p>
      <h2>1 · Authorization</h2>
      <p>{_mark('signatory' in hl,
      f"<strong>{_e(intake.get('signatory_name'))}</strong>, {_e(intake.get('signatory_title'))} of the "
      "Company, is duly authorized to execute and deliver, on behalf of the Company, any and all "
      "applications, agreements, and instruments in connection with the Company's generator "
      "interconnection request to the California Independent System Operator Corporation, including "
      "Appendix 1 (Interconnection Request) and submissions through the RIMS5 system.")}</p>
      <h2>2 · Incumbency</h2>
      <p>The signature set forth beside the officer's name below is such officer's true signature.</p>
      <div class="sig">
        <div><strong>{_e(intake.get('signatory_name'))}</strong>{_e(intake.get('signatory_title'))} — specimen signature</div>
        <div><strong>M. Chen</strong>Secretary — certifying officer<br/>Date: 03/02/2026</div>
      </div>
    </div>"""


# Rows of the developer technical data workbook: (intake key, sheet label, unit).
_WORKBOOK_ROWS: list[tuple[str, list[tuple[str, str, str]]]] = [
    ("Project technical parameters", [
        ("project_type", "Project type", ""),
        ("gross_mva", "Gross capacity", "MVA"),
        ("gross_mw", "Gross output", "MW"),
        ("aux_mw", "Auxiliary (station) load", "MW"),
        ("losses_mw", "Electrical losses to POI", "MW"),
        ("net_mw_poi", "Requested net output at POI", "MW"),
    ]),
    ("Equipment", [
        ("inverter", "Inverter manufacturer / model / qty", ""),
        ("module", "PV module", ""),
        ("transformer", "Main transformer (GSU)", ""),
        ("collector_kv", "Collector system voltage", "kV"),
    ]),
]

# Rows of the storage vendor's BESS specification sheet.
_BESS_ROWS: list[tuple[str, list[tuple[str, str, str]]]] = [
    ("Storage system", [
        ("bess_vendor", "Manufacturer / model", ""),
        ("bess_mw", "Power rating", "MW"),
        ("bess_mwh", "Energy capacity", "MWh"),
        ("bess_charging", "Charging configuration", ""),
    ]),
]

WORKBOOK_CSS = """<style>
  .wb { max-width: 720px; margin: 0 auto; background: #fff; border: 1px solid var(--line);
    border-radius: 10px; overflow: hidden; box-shadow: 0 1px 8px rgba(0,0,0,0.04); }
  .wb table { width: 100%; border-collapse: collapse; font-size: 13px; }
  .wb th { text-align: left; font-family: "JetBrains Mono", monospace; font-size: 10px;
    text-transform: uppercase; letter-spacing: 0.08em; color: var(--muted);
    background: var(--soft); padding: 8px 14px; border-bottom: 1px solid var(--line); }
  .wb td { padding: 8px 14px; border-bottom: 1px solid var(--line); vertical-align: top; }
  .wb td.val { font-family: "JetBrains Mono", monospace; white-space: nowrap; }
  .wb tr.sec td { background: var(--soft); font-family: "JetBrains Mono", monospace;
    font-size: 10.5px; text-transform: uppercase; letter-spacing: 0.08em; color: var(--accent); }
  .wb tr.vhl-row td { background: #fffbeb; box-shadow: inset 3px 0 0 #f59e0b; }
  .wb td .blank { color: #b23b3b; font-style: italic; }
</style>"""


def _sheet_body(intake: dict, sections: list, note: str, hl: list[str] | None = None) -> str:
    """Spreadsheet-style preview shared by the technical workbook and BESS spec sheet."""
    hl = hl or []
    rows_html = ""
    for section, rows in sections:
        rows_html += f'<tr class="sec"><td colspan="3">{_e(section)}</td></tr>'
        for key, rlabel, unit in rows:
            v = intake.get(key)
            if isinstance(v, float) and v.is_integer():
                v = int(v)
            blank = v is None or str(v).strip() == ""
            val = '<span class="blank">— blank —</span>' if blank else _e(v)
            cls = ' class="vhl-row"' if key in hl else ""
            rows_html += (f"<tr{cls}><td>{_e(rlabel)}</td>"
                          f'<td class="val">{val}</td><td class="val">{_e(unit)}</td></tr>')
    return WORKBOOK_CSS + HL_CSS + f"""
    <p class="note">{note}</p>
    {_hl_note(hl)}
    <div class="wb">
      <table>
        <thead><tr><th>Parameter</th><th>Value</th><th>Unit</th></tr></thead>
        <tbody>{rows_html}</tbody>
      </table>
    </div>"""


def _workbook_body(intake: dict, hl: list[str] | None = None) -> str:
    return _sheet_body(
        intake, _WORKBOOK_ROWS,
        "Developer technical data workbook — the source of the technical parameters in the "
        "intake form. Values shown reflect the file currently attached to the intake.", hl)


def _bess_body(intake: dict, hl: list[str] | None = None) -> str:
    return _sheet_body(
        intake, _BESS_ROWS,
        "Storage vendor specification sheet — the source of the storage parameters in the "
        "intake form. Values shown reflect the file currently attached to the intake.", hl)


def render_requirement_preview(rule_id: str, req: dict) -> str:
    """Preview page for the ISO requirement (ground truth) behind a validation check."""
    paras = "".join(f"<p>{_e(p)}</p>" for p in req.get("paragraphs", []))
    return PAGE.format(
        title=_e(req.get("title") or "Requirement"),
        filename=_e(req.get("source") or rule_id),
        project="CAISO requirement",
        status_label="Ground truth",
        chip_class="ok",
        actions="",
        body=DOC_CSS + HL_CSS + f"""
        <p class="note">Requirement summary drafted for this demo from the cited CAISO provision —
        the operative clause enforced by the validation check is highlighted.</p>
        <div class="paper">
          <h1>{_e(req.get('title'))}</h1>
          <p class="sub">{_e(req.get('source'))}</p>
          {paras}
          <h2>Operative clause</h2>
          <p><mark class="vhl">{_e(req.get('clause'))}</mark></p>
        </div>""",
    )


def render_kickoff_preview(key: str, intake: dict, meta: dict, hl: list[str] | None = None) -> str:
    """Preview page for a kickoff document (Step 2 uploads), with optional highlights."""
    hl = hl or []
    if key == "file_site_control":
        title, body = "Executed Site Exclusivity Agreement", _lease_body(intake, hl)
    elif key == "file_signatory":
        title, body = "Proof of Authorized Signatory", _signatory_body(intake, hl)
    elif key == "file_technical":
        title, body = "Technical Data Workbook", _workbook_body(intake, hl)
    elif key == "file_bess":
        title, body = "BESS Specification Sheet", _bess_body(intake, hl)
    elif key == "file_boundary":
        import tempfile
        from backend.app.services.caiso_packet import _derived, _gen_kmz

        with tempfile.NamedTemporaryFile(suffix=".kmz") as tmp:
            path = Path(tmp.name)
            _gen_kmz(intake, _derived(intake), path)
            body = (
                '<p class="note">Example boundary file preloaded by the demo — rendered from the parcel '
                "polygon. In production this is the developer's own KMZ or parcel map.</p>"
                + (HL_CSS + _hl_note(hl) if hl else "")
                + _kmz_body(path, {})
            )
        title = "Project Boundary (KMZ)"
    elif key == "file_dyd":
        from backend.app.engine.vendor_models import demo_vendor_dyd_text

        text = demo_vendor_dyd_text(intake)
        body = (
            '<p class="note">Vendor-supplied PSLF parameter package for the selected inverter — '
            "the OEM blocks GridPilot integrates into the plant-level dynamic model. Records are "
            "per-unit; the engine scales them to the fleet MVA base when it writes the packet "
            ".dyd.</p>"
            + (HL_CSS + _hl_note(hl) if hl else "")
            + _dyd_body(text)
        )
        title = "Vendor PSLF Dynamic Model (.dyd)"
    elif key == "file_base_case":
        poi = _e(str(intake.get("poi_name") or "POI"))
        kv = _e(str(intake.get("poi_voltage_kv") or "230"))
        body = f"""
        <p class="note">Demo stand-in for a <strong>customer-supplied ISO base-case extract</strong>
        covering the {_e(poi)} area. Real ISO base cases are confidential and stay in the customer's
        authorized environment — GridPilot never fabricates one. Without a real extract the plant
        model is validated against a flat system equivalent at the POI.</p>
        <div class="sum">
          <div class="card"><b>{poi}</b><span>Anchor POI</span></div>
          <div class="card"><b>{kv} kV</b><span>System voltage</span></div>
          <div class="card"><b>DEMO</b><span>Not a live ISO case</span></div>
          <div class="card"><b>PSLF</b><span>Extract format</span></div>
        </div>
        <table>
          <tr><th colspan="4">Area extract — representative buses (demo)</th></tr>
          <tr><th>Bus</th><th>Name</th><th>kV</th><th>Role</th></tr>
          <tr><td>11001</td><td>WHIRLWND</td><td>{kv}</td><td>POI / swing</td></tr>
          <tr><td>11002</td><td>WHIRL-HS</td><td>{kv}</td><td>Utility high-side</td></tr>
          <tr><td>11014</td><td>ANTELOPE</td><td>{kv}</td><td>Neighboring station</td></tr>
          <tr><td>11022</td><td>WINDHUB</td><td>{kv}</td><td>Neighboring station</td></tr>
        </table>
        <p class="note" style="margin-top:14px">In production the customer drops the licensed extract
        here; GridPilot records it on the validation report and reminds the engineer to re-run the
        final load-flow against that case in the authorized environment.</p>
        """ + (HL_CSS + _hl_note(hl) if hl else "")
        title = "ISO Base Case Extract (customer-supplied)"
    else:
        title = "Kickoff document"
        body = '<p class="note">No inline preview available for this document.</p>'

    return PAGE.format(
        title=_e(title),
        filename=_e(meta.get("name") or key),
        project=_e(intake.get("project_name") or ""),
        status_label="Example file — demo" if meta.get("example") else "Uploaded document",
        chip_class="",
        actions="",
        body=body,
    )
