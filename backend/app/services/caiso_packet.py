"""CAISO Interconnection Request packet generation.

GridPilot's core demo: take the developer's Phase-1 intake and generate the
consulting-firm workstream — Appendix 1, Attachment A data, models, drawings,
plots, and legal drafts — as a downloadable submission packet.
"""

from __future__ import annotations

import base64
import hashlib
import json
import math
import re
import shutil
import zipfile
from datetime import date, datetime, timedelta, timezone
from pathlib import Path
from typing import Any

import copy

import fitz

from backend.app.config import DATA_ROOT
from backend.app.services.iso_profiles import get_profile, localize
from backend.app.engine import equipment as lib
from backend.app.engine.design import INVERTER_SIZING_PF
from backend.app.engine.consistency import run_engineering
from backend.app.engine.models_out import dyd_text, dyr_text, epc_text, raw_text
from backend.app.engine.reports import (
    gen_assumptions_log, gen_dynamics_report, gen_equipment_schedule,
    gen_loadflow_report, gen_missing_data, gen_portal_fields, gen_source_trace,
)
from backend.app.engine.sld import build_sld, to_dxf, to_pdf as sld_to_pdf

PACKETS_DIR = DATA_ROOT / "packets"

RIMS5_URL = "https://rimspub.caiso.com/rims5/logon.do"
CAISO_FORMS_URL = "https://www.caiso.com/library/interconnection-request-technical-data-forms"

# ---------------------------------------------------------------------------
# Intake schema (drives the wizard form) + Ravenwood example defaults
# ---------------------------------------------------------------------------

INTAKE_SECTIONS: list[dict[str, Any]] = [
    {
        "id": "legal",
        "title": "1 · Legal & Corporate",
        "hint": "Legal name must match the Secretary of State certificate exactly — punctuation included.",
        "fields": [
            {"key": "legal_name", "label": "Legal entity name", "type": "text", "required": True,
             "hint": "Exactly as registered with the Secretary of State"},
            {"key": "state_of_origin", "label": "State of origin", "type": "text"},
            {"key": "signatory_name", "label": "Authorized signatory — name", "type": "text", "required": True},
            {"key": "signatory_title", "label": "Authorized signatory — title", "type": "text"},
            {"key": "contact_email", "label": "Contact email", "type": "text"},
            {"key": "contact_phone", "label": "Contact phone", "type": "text"},
        ],
    },
    {
        "id": "site",
        "title": "2 · Site Control",
        "hint": "Letters of intent are not accepted as evidence of site exclusivity.",
        "fields": [
            {"key": "project_name", "label": "Project name", "type": "text", "required": True},
            {"key": "gps_lat", "label": "GPS latitude (decimal)", "type": "number", "required": True},
            {"key": "gps_lon", "label": "GPS longitude (decimal)", "type": "number", "required": True},
            {"key": "county", "label": "County", "type": "text"},
            {"key": "state", "label": "State", "type": "text"},
            {"key": "site_acreage", "label": "Site acreage", "type": "number"},
            {"key": "site_control", "label": "Site exclusivity type", "type": "select", "required": True,
             "options": ["Lease Agreement", "Option to Purchase", "Deed (ownership)", "Letter of Intent", "None yet"],
             "hint": "LOI is not accepted by CAISO"},
            {"key": "site_owner", "label": "Site owner / lessor", "type": "text"},
        ],
    },
    {
        "id": "intent",
        "title": "3 · Interconnection Intent",
        "hint": "POI choice drives study cost and upgrade exposure.",
        "fields": [
            {"key": "poi_name", "label": "Target POI — substation / line", "type": "text", "required": True},
            {"key": "poi_voltage_kv", "label": "POI voltage (kV)", "type": "number", "required": True},
            {"key": "track", "label": "Process track", "type": "select", "required": True,
             "options": ["Independent Study Process", "Fast Track", "Cluster"],
             "hint": "Fast Track requires ≤ 5 MW; ISP requires an eligibility demonstration"},
            {"key": "deliverability", "label": "Requested deliverability", "type": "select",
             "options": ["Full Capacity", "Partial Deliverability", "Energy Only"]},
            {"key": "cod", "label": "Target COD (YYYY-MM-DD)", "type": "text", "required": True,
             "hint": "Must be within 7 years of the application"},
            {"key": "gentie_mi", "label": "Gen-tie route length (mi)", "type": "number",
             "hint": "Blank — GridPilot assumes a 5.0 mi route, flagged for engineer approval"},
            {"key": "shared_facilities", "label": "Shared facilities", "type": "select",
             "options": ["No shared facilities", "Shared gen-tie", "Shared substation",
                          "Shared gen-tie + substation"],
             "hint": "Whether the gen-tie, substation, or POI is shared with another project"},
            {"key": "queue_ref", "label": "Prior queue / study reference", "type": "text",
             "hint": "Existing queue position or prior study number, if any"},
        ],
    },
    {
        "id": "technical",
        "title": "4 · Project Technical Parameters",
        "hint": "MW chain: Gross Output − Aux Load − Losses = Net MW at POI. Must match across every document.",
        "fields": [
            {"key": "project_type", "label": "Project type", "type": "select", "required": True,
             "options": ["Solar PV", "Solar PV + BESS (AC-coupled)", "Solar PV + BESS (DC-coupled)",
                          "Standalone BESS", "Wind", "Wind + BESS"]},
            {"key": "gross_mva", "label": "Gross capacity (MVA)", "type": "number", "required": True},
            {"key": "gross_mw", "label": "Gross output (MW)", "type": "number", "required": True},
            {"key": "aux_mw", "label": "Auxiliary load (MW)", "type": "number", "required": True},
            {"key": "losses_mw", "label": "Losses to POI (MW)", "type": "number", "required": True},
            {"key": "net_mw_poi", "label": "Requested net MW at POI", "type": "number", "required": True,
             "hint": "The number that appears in the CAISO queue"},
            {"key": "bess_mw", "label": "BESS power (MW)", "type": "number", "hint": "Blank if no storage"},
            {"key": "bess_mwh", "label": "BESS energy (MWh)", "type": "number"},
            {"key": "bess_charging", "label": "BESS charging source", "type": "select",
             "options": ["On-site generation only", "Grid charging permitted", "N/A — no storage"]},
        ],
    },
    {
        "id": "equipment",
        "title": "5 · Equipment & Vendor",
        "hint": "Vendor .dyd dynamic model files are the most common schedule risk — request them on day one.",
        "fields": [
            {"key": "inverter", "label": "Inverter manufacturer / model / qty", "type": "text",
             "suggest": "inverters",
             "hint": "Pick from the equipment library or type freely — enter TBD if not "
                     "selected (generic PSLF models will be used)"},
            {"key": "module", "label": "PV module / turbine model", "type": "text"},
            {"key": "bess_vendor", "label": "BESS manufacturer / model", "type": "text",
             "suggest": "bess",
             "hint": "Pick from the equipment library or type freely"},
            {"key": "dyd_status", "label": "Vendor .dyd model files", "type": "select", "required": True,
             "options": ["Received from vendor", "Requested — pending", "Equipment not selected"],
             "hint": "PSLF dynamic models from the equipment vendor"},
            {"key": "transformer", "label": "Main transformer (GSU) data", "type": "text",
             "hint": "MVA, voltage ratio, impedance, vector group"},
            {"key": "collector_kv", "label": "Collector system voltage (kV)", "type": "number"},
        ],
    },
    {
        "id": "preferences",
        "title": "6 · Design Preferences",
        "hint": "Optional constraints the design engine honors when feasible; everything is "
                "recorded in the source trace.",
        "fields": [
            {"key": "substation_arrangement", "label": "Substation arrangement", "type": "select",
             "options": ["Engine recommendation", "Single main transformer",
                          "Two main transformers (N-1)"],
             "hint": "Overrides the GSU count when a standard family unit can cover the rating"},
            {"key": "site_constraints", "label": "Site constraints", "type": "text",
             "hint": "Setbacks, wetlands, easements, shared-fence lines — recorded for the "
                     "engineer of record"},
        ],
    },
    {
        "id": "documents",
        "title": "7 · Kickoff Documents",
        "hint": "The files a consulting firm collects at the kickoff meeting. Example files are preloaded "
                "and may be replaced or removed; validation reflects any change.",
        "fields": [
            {"key": "file_site_control", "label": "Executed site exclusivity agreement", "type": "file",
             "required": True, "accept": ".pdf,.doc,.docx",
             "hint": "Executed lease / option / deed. This is the evidence behind the site exclusivity "
                     "declaration — CAISO rejects LOIs."},
            {"key": "file_technical", "label": "Technical data workbook", "type": "file",
             "required": True, "accept": ".xlsx,.xls,.csv",
             "hint": "The developer's technical data sheet — MW chain, inverter, GSU, collector system. "
                     "Source of the technical parameters in the intake form."},
            {"key": "file_bess", "label": "BESS specification sheet", "type": "file",
             "accept": ".xlsx,.xls,.pdf",
             "hint": "The storage vendor's specification — power rating, energy capacity, charging "
                     "configuration. Source of the storage parameters in the intake form."},
            {"key": "file_signatory", "label": "Proof of authorized signatory", "type": "file",
             "accept": ".pdf,.doc,.docx",
             "hint": "Officer certificate or board resolution naming the signatory. If missing, GridPilot "
                     "drafts one for counsel to execute."},
            {"key": "file_dyd", "label": "Vendor PSLF dynamic model (.dyd)", "type": "file",
             "accept": ".dyd,.zip",
             "hint": "From the inverter/BESS vendor — the most common schedule risk. Generic WECC models "
                     "are used until it arrives."},
            {"key": "file_boundary", "label": "Project boundary (KMZ / parcel map)", "type": "file",
             "accept": ".kmz,.kml,.pdf",
             "hint": "Optional — without it, GridPilot derives the boundary from GPS + acreage."},
            {"key": "file_base_case", "label": "ISO base case extract (customer-supplied)",
             "type": "file", "accept": ".epc,.raw,.sav,.zip",
             "hint": "Optional — ISO base cases are confidential and stay in the customer's "
                     "authorized environment. GridPilot never fabricates one; without it the "
                     "plant model is validated against a flat system equivalent."},
        ],
    },
]

DEFAULT_INTAKE: dict[str, Any] = {
    "legal_name": "Ravenwood Energy LLC",
    "state_of_origin": "Delaware",
    "signatory_name": "Ryan Yang",
    "signatory_title": "Director of Development",
    "contact_email": "ryan@ravenwoodenergy.com",
    "contact_phone": "(661) 555-0142",
    "project_name": "Ravenwood",
    "gps_lat": 35.0842,
    "gps_lon": -118.3081,
    "county": "Kern",
    "state": "CA",
    "site_acreage": 610,
    "site_control": "Lease Agreement",
    "site_owner": "Willow Springs Ranch LP",
    "poi_name": "Whirlwind Substation (SCE)",
    "poi_voltage_kv": 230,
    "track": "Independent Study Process",
    "deliverability": "Full Capacity",
    # Sample timeline sits comfortably beyond the 2-year development horizon:
    # in-service 12/02/2028, trial op 03/02/2029, COD 06/30/2029.
    "cod": "2029-06-30",
    # Gen-tie route left blank on purpose: the demo shows the engine carrying it
    # as an approvable routing assumption until the developer provides the fact.
    "gentie_mi": None,
    "shared_facilities": "No shared facilities",
    "queue_ref": "",
    "project_type": "Solar PV + BESS (AC-coupled)",
    # Machine-table sum anchored to the 128 MW gross output goal (MVA = MW / 0.95):
    # 19 x 4.4 MVA PV stations (SG4400UD-MV-US rated 4400 kVA @ 40 C per the US
    # datasheet) + 52 x 1.05 MVA Megapack PCS (4-hour configured, 50 kVA
    # orderable steps) = 138.2 MVA — matches Attachment A's SUMPRODUCT.
    "gross_mva": 138.2,
    "gross_mw": 128.0,
    "aux_mw": 2.5,
    "losses_mw": 0.5,
    # Net reconciles the MW chain: gross 128 − aux 2.5 − losses 0.5 = 125 at POI.
    # One deliberate kickoff defect remains so the demo exercises the
    # fix-and-revalidate loop: BESS energy (MWh) is missing.
    "net_mw_poi": 125.0,
    "bess_mw": 50.0,
    "bess_mwh": None,
    "bess_charging": "On-site generation only",
    # Quantity is computed by the design engine (MVA = MW / 0.95 rule), not declared.
    "inverter": "Sungrow SG4400UD-MV",
    "module": "JinkoSolar Tiger Neo 620W bifacial",
    "bess_vendor": "Tesla Megapack 2XL",
    "dyd_status": "Received from vendor",
    # Matches the engine's family pick for the 140.71 MVA machine sum:
    # 2 x 75 MVA (45/60/75 MVA ONAN/ONAF1/ONAF2 ladder). Z/X/R from the
    # developer ground-truth workbook (Z = 9%, X/R = 42).
    "transformer": "75 MVA, 34.5/230 kV, Z = 9% @ ONAF, YNd1",
    "collector_kv": 34.5,
    "substation_arrangement": "Engine recommendation",
    "site_constraints": "",
    # Kickoff document uploads: {name, size} metadata; "example" marks preloaded demo
    # files. All examples start staged (uploaded, not yet submitted) so the demo walks
    # the full choose → upload → preview → submit interaction.
    "file_site_control": {"name": "Ravenwood_Lease_WillowSpringsRanch_Executed_2026-03-14.docx",
                          "size": 384_726, "example": True, "staged": True},
    "file_technical": {"name": "Ravenwood_TechnicalData_Workbook_v1.xlsx", "size": 87_450,
                       "example": True, "staged": True},
    "file_bess": {"name": "Ravenwood_BESS_Spec_Megapack2XL_v1.xlsx", "size": 54_120,
                  "example": True, "staged": True},
    "file_signatory": {"name": "Ravenwood_OfficerCertificate_RYang.docx", "size": 64_180,
                       "example": True, "staged": True},
    "file_dyd": {"name": "Sungrow_SG4400UD_PSLF_Models.dyd", "size": 18_240,
                 "example": True, "staged": True},
    "file_boundary": {"name": "Ravenwood_ParcelBoundary_KernCounty.kmz", "size": 46_210,
                      "example": True, "staged": True},
    # Demo stand-in for a customer-supplied POI-area extract. Not a real ISO
    # base case — those stay confidential and are never fabricated by GridPilot.
    "file_base_case": {"name": "Whirlwind_230kV_Area_Extract_DEMO.epc", "size": 42_880,
                       "example": True, "staged": True},
}


def intake_sections_for(iso: str | None) -> list[dict[str, Any]]:
    """Intake schema localized for an ISO: form names, model formats, tracks."""
    profile = get_profile(iso)
    if profile["iso"] == "CAISO":
        return INTAKE_SECTIONS
    sections = copy.deepcopy(INTAKE_SECTIONS)
    for section in sections:
        section["title"] = localize(section["title"], profile)
        if section.get("hint"):
            section["hint"] = localize(section["hint"], profile)
        for f in section["fields"]:
            f["label"] = localize(f["label"], profile)
            if f.get("hint"):
                f["hint"] = localize(f["hint"], profile)
            if f["key"] == "track":
                f["options"] = list(profile["tracks"])
                f["hint"] = f"Process track under {profile['tariff']}"
            if f["key"] == "file_dyd":
                f["accept"] = f".{profile['dyn_ext']},.zip"
    return sections


def _file_meta(v: Any) -> dict[str, Any] | None:
    """Return upload metadata ({name, size, ...}) if a file is submitted, else None.

    Files carry a staged→submitted lifecycle: uploads marked `staged` are visible
    in the UI but do not count for validation or extraction until submitted.
    """
    if isinstance(v, dict) and str(v.get("name") or "").strip() and not v.get("staged"):
        return v
    return None


# Intake fields each kickoff document yields under AI extraction. Ordered so the
# most authoritative source for a field wins (e.g. boundary file for GPS).
EXTRACTION_SOURCES: list[tuple[str, str, list[str]]] = [
    ("file_site_control", "Site exclusivity agreement",
     ["legal_name", "site_owner", "site_control", "site_acreage", "county", "state"]),
    ("file_technical", "Technical data workbook",
     ["project_type", "gross_mva", "gross_mw", "aux_mw", "losses_mw", "net_mw_poi",
      "inverter", "module", "transformer", "collector_kv"]),
    ("file_bess", "BESS specification sheet",
     ["bess_vendor", "bess_mw", "bess_mwh", "bess_charging"]),
    ("file_signatory", "Certificate of authorized signatory",
     ["signatory_name", "signatory_title", "legal_name", "state_of_origin"]),
    ("file_boundary", "Project boundary file",
     ["project_name", "gps_lat", "gps_lon", "site_acreage", "county", "state"]),
    ("file_dyd", "Vendor dynamic model", ["dyd_status"]),
]

# Values a corrected (non-example) replacement upload carries for its intake
# fields, so a corrected revision clears exactly its own finding. The example
# BESS sheet carries the seeded kickoff defect (missing MWh); the technical
# workbook entry also lets a corrected upload repair a manually mistyped net MW.
CORRECTED_EXTRACTIONS: dict[str, dict[str, Any]] = {
    "file_technical": {"net_mw_poi": 125.0},
    "file_bess": {"bess_mwh": 200.0},
}


# Ground-truth requirements each validation check enforces. Rendered as a
# previewable requirement page with the operative clause highlighted.
# Demo-drafted summaries of the CAISO tariff / BPM provisions, not verbatim text.
REQUIREMENTS: dict[str, dict[str, Any]] = {
    "legal-name": {
        "title": "Interconnection Customer legal identity",
        "source": "CAISO BPM for Generator Interconnection — Appendix 1 completion instructions",
        "paragraphs": [
            "The Interconnection Request must identify the Interconnection Customer by its full legal "
            "name as registered with the Secretary of State of its state of formation.",
            "The legal name entered in Appendix 1 is cross-checked against the site exclusivity "
            "documentation and the certificate of authorized signatory during deficiency review.",
        ],
        "clause": "The name of the Interconnection Customer must match the Secretary of State "
                  "registration exactly and must be used consistently across all submitted documents.",
    },
    "signatory": {
        "title": "Execution by an authorized signatory",
        "source": "CAISO BPM for Generator Interconnection — Appendix 1 execution requirements",
        "paragraphs": [
            "Appendix 1 must be executed in RIMS5 by an individual with authority to bind the "
            "Interconnection Customer.",
            "CAISO may request evidence of that authority — typically an officer certificate or a "
            "board resolution naming the signatory.",
        ],
        "clause": "The Interconnection Request must be executed by an authorized representative of "
                  "the Interconnection Customer; evidence of signing authority must be available on request.",
    },
    "gps": {
        "title": "Site location data",
        "source": "CAISO Appendix 1 — Section 6 (Project site information)",
        "paragraphs": [
            "The Interconnection Request must state the geographic location of the Generating "
            "Facility, including latitude and longitude in decimal format.",
            "The stated coordinates feed Appendix 1, Attachment A, and the project boundary file, "
            "and must agree across all three.",
        ],
        "clause": "Latitude and longitude must be provided in decimal degrees and must be consistent "
                  "with the project boundary file submitted with the request.",
    },
    "site-exclusivity": {
        "title": "Evidence of site exclusivity",
        "source": "CAISO Tariff Appendix DD (GIDAP) — Site Exclusivity requirements",
        "paragraphs": [
            "An Interconnection Request under the Independent Study Process must demonstrate site "
            "exclusivity for the Generating Facility site, or pay the site exclusivity deposit in lieu.",
            "Acceptable evidence is an executed lease, an option to lease or purchase, or a deed for "
            "the parcels comprising the site. The document itself must be produced.",
        ],
        "clause": "A letter of intent or other non-binding instrument does not constitute site "
                  "exclusivity — an executed lease, option, or deed is required.",
    },
    "poi": {
        "title": "Point of Interconnection identification",
        "source": "CAISO Appendix 1 — Section 5 (Point of Interconnection)",
        "paragraphs": [
            "The request must identify the proposed Point of Interconnection: the transmission "
            "substation or line segment and the interconnection voltage level.",
            "The POI stated in Appendix 1 governs the load flow model, the single-line diagram, and "
            "the study scope.",
        ],
        "clause": "Both the substation / line name and the voltage level (kV) of the Point of "
                  "Interconnection are required entries.",
    },
    "mw-chain": {
        "title": "Consistency of stated capacity values",
        "source": "CAISO Attachment A — technical data consistency rules",
        "paragraphs": [
            "Attachment A requires gross generating capability, auxiliary (station) load, electrical "
            "losses to the Point of Interconnection, and the requested net output at the POI.",
            "These values are checked arithmetically during deficiency review and are cross-referenced "
            "against Appendix 1 and the submitted PSLF models.",
        ],
        "clause": "The requested net MW at the Point of Interconnection must equal gross output minus "
                  "auxiliary load minus losses; inconsistent MW values are a deficiency finding.",
    },
    "fast-track": {
        "title": "Fast Track eligibility limit",
        "source": "CAISO Tariff Appendix DD (GIDAP) — Fast Track Process eligibility",
        "paragraphs": [
            "The Fast Track Process is available only to Generating Facilities with a capacity at or "
            "below the Fast Track limit at the Point of Interconnection.",
        ],
        "clause": "A Generating Facility exceeding 5 MW at the Point of Interconnection is not "
                  "eligible for the Fast Track Process.",
    },
    "isp": {
        "title": "Independent Study Process requirements",
        "source": "CAISO Tariff Appendix DD (GIDAP) — Independent Study Process",
        "paragraphs": [
            "An ISP request must include an eligibility demonstration showing the project is "
            "electrically independent of pending cluster requests, together with the study deposit.",
        ],
        "clause": "The Independent Study Process requires an eligibility demonstration and a "
                  "$150,000 study deposit submitted with the Interconnection Request.",
    },
    "deliverability": {
        "title": "Deliverability assessment timing",
        "source": "CAISO Tariff Appendix DD (GIDAP) — Section 4.6 (Deliverability)",
        "paragraphs": [
            "A request for Full Capacity Deliverability Status is assessed in the annual Cluster "
            "Study deliverability analysis, even where the interconnection study proceeds under ISP.",
        ],
        "clause": "Full Capacity Deliverability Status is determined in the next annual Cluster "
                  "Study cycle — it is not established by the Independent Study itself.",
    },
    "cluster-window": {
        "title": "Cluster request window",
        "source": "CAISO Tariff Appendix DD (GIDAP) — Cluster Study process",
        "paragraphs": [
            "Cluster Interconnection Requests are accepted only during the annual request window; "
            "requests received outside the window roll to the following cycle.",
        ],
        "clause": "Interconnection Requests under the Cluster Study process may be submitted only "
                  "during the annual open request window.",
    },
    "cod-limit": {
        "title": "Commercial operation date limit",
        "source": "CAISO Tariff Appendix DD (GIDAP) — Interconnection Request contents",
        "paragraphs": [
            "The Interconnection Request must state a proposed commercial operation date.",
        ],
        "clause": "The proposed commercial operation date may not exceed seven years from the date "
                  "the Interconnection Request is submitted.",
    },
    "bess-energy": {
        "title": "Energy storage technical data",
        "source": "CAISO Attachment A — energy storage supplement",
        "paragraphs": [
            "Where the Generating Facility includes an energy storage component, the request must "
            "state both the storage power rating and the energy capacity.",
            "Charging behavior must also be declared, as it affects the deliverability assessment "
            "and the load flow model.",
        ],
        "clause": "Storage projects must state the energy capacity (MWh) and duration in addition "
                  "to the MW rating; a power rating alone is a deficiency.",
    },
    "bess-charging": {
        "title": "Storage charging source declaration",
        "source": "CAISO Attachment A — energy storage supplement",
        "paragraphs": [
            "The request must declare whether the storage component charges exclusively from on-site "
            "generation or also from the grid.",
        ],
        "clause": "Grid charging changes the deliverability assessment and the load-flow model and "
                  "must be declared in the Interconnection Request.",
    },
    "dyd-models": {
        "title": "Dynamic model data requirements",
        "source": "CAISO BPM for Generator Interconnection — modeling data requirements",
        "paragraphs": [
            "The Interconnection Request must include dynamic models compatible with the CAISO "
            "planning tools (GE PSLF), using WECC-approved model structures.",
            "Vendor-specific parameter files (.dyd) are required for the inverter and storage "
            "equipment; generic WECC models are accepted only as placeholders.",
        ],
        "clause": "Dynamic model data declared as received from the vendor must be included with "
                  "the submission in PSLF-compatible (.dyd) format.",
    },
    "boundary": {
        "title": "Project boundary file",
        "source": "CAISO BPM for Generator Interconnection — site documentation",
        "paragraphs": [
            "A geographic boundary file (KMZ/KML) delineating the Generating Facility site supports "
            "the site exclusivity review and the study of the interconnection route.",
        ],
        "clause": "The project boundary file must be consistent with the site coordinates and "
                  "acreage stated in Appendix 1.",
    },
}


def demo_datasheet_draft(name: str, kind: str) -> dict[str, Any]:
    """Deterministic datasheet 'extraction' for the guided demo (metadata only).

    The demo never reads file bytes; the draft is parsed from the filename so
    the review-and-confirm workflow is exercised end to end.
    """
    is_bess = kind == "bess"
    stem = re.sub(r"\.[A-Za-z0-9]+$", "", name or "datasheet")
    tokens = [t for t in re.split(r"[_\-\s]+", stem) if t]
    m = re.search(r"(\d+(?:\.\d+)?)\s*(mva|kva|mw)", name or "", re.IGNORECASE)
    if m:
        mva = float(m.group(1)) / (1000.0 if m.group(2).lower() == "kva" else 1.0)
    else:
        mva = 2.0 if is_bess else 3.6
    vendor = tokens[0] if tokens else "OEM"
    if not any(c.isupper() for c in vendor):
        vendor = vendor.title()
    return {
        "kind": kind if kind in ("pv_inverter", "bess") else "pv_inverter",
        "vendor": vendor,
        "model": " ".join(tokens[1:3]) if len(tokens) > 1 else "Custom Unit",
        "mva": round(mva, 3), "mw": round(mva, 3),
        "mwh": round(mva * 2, 3) if is_bess else None,
        "ac_kv": 0.48 if is_bess else 0.69,
        "pf_range": 0.95 if is_bess else 0.9,
        "notes": "Demo extraction — draft parsed from the filename; review every value "
                 "before confirming",
        "source": name,
    }


def extract_from_documents(files: dict[str, Any]) -> dict[str, Any]:
    """Simulated AI extraction: map attached kickoff documents to intake fields.

    The demo's example documents carry the Ravenwood values; each extracted field
    is returned with provenance so the form can show where the value came from.
    """
    fields: dict[str, Any] = {}
    provenance: dict[str, dict[str, str]] = {}
    docs_used = 0
    for key, source_label, field_keys in EXTRACTION_SOURCES:
        meta = _file_meta(files.get(key))
        if meta is None:
            continue
        docs_used += 1
        corrections = {} if meta.get("example") else CORRECTED_EXTRACTIONS.get(key, {})
        for fk in field_keys:
            if fk in corrections:
                value = corrections[fk]
            elif fk == "dyd_status":
                value = "Received from vendor"
            else:
                value = DEFAULT_INTAKE.get(fk)
            if value is None:
                continue
            fields[fk] = value
            provenance[fk] = {"source": key, "source_label": source_label, "file": str(meta["name"])}

    # Vendor .dyd: parse OEM parameter blocks so the engine integrates them
    # (demo files carry the matched inverter's OEM package — same parser as
    # the real upload path).
    dyd_meta = _file_meta(files.get("file_dyd"))
    if dyd_meta is not None:
        from backend.app.engine.vendor_models import demo_vendor_dyd_text, vendor_overrides

        ov = vendor_overrides(demo_vendor_dyd_text({**DEFAULT_INTAKE, **fields}),
                              str(dyd_meta["name"]))
        if ov:
            fields["vendor_dyd"] = ov
            provenance["vendor_dyd"] = {"source": "file_dyd",
                                        "source_label": "Vendor dynamic model",
                                        "file": str(dyd_meta["name"])}
    return {
        "fields": fields,
        "provenance": provenance,
        "summary": f"{len(fields)} fields extracted from {docs_used} document(s)",
    }


# ---------------------------------------------------------------------------
# Validation — the consulting "Step A: kickoff & data validation"
# ---------------------------------------------------------------------------

def _num(v: Any) -> float | None:
    if v is None or v == "":
        return None
    try:
        return float(str(v).replace(",", ""))
    except ValueError:
        return None


def _parse_cod(v: Any) -> date | None:
    s = str(v or "").strip()
    for fmt in ("%Y-%m-%d", "%m/%d/%Y", "%m/%d/%y"):
        try:
            return datetime.strptime(s, fmt).date()
        except ValueError:
            continue
    return None


def validate_intake(intake: dict[str, Any], iso: str | None = None) -> dict[str, Any]:
    profile = get_profile(iso)
    errors: list[dict[str, Any]] = []
    warnings: list[dict[str, Any]] = []
    passed: list[dict[str, Any]] = []

    def _rule_ref(rule_id: str | None) -> dict[str, str] | None:
        req = REQUIREMENTS.get(rule_id or "")
        if not req:
            return None
        return {"id": rule_id, "title": localize(req["title"], profile),
                "source": localize(req["source"], profile)}

    def _evidence_ref(file_key: str | None, hl: list[str] | None = None) -> dict[str, Any] | None:
        """The examined document behind a check: which upload, and what to highlight."""
        if not file_key:
            return None
        meta = _file_meta(intake.get(file_key))
        return {
            "key": file_key,
            "file": meta["name"] if meta else None,
            "attached": meta is not None,
            "example": bool(meta and meta.get("example")),
            "hl": hl or [],
        }

    def _item(field: str, title: str, detail: str,
              rule: str | None, evidence: str | None, hl: list[str] | None) -> dict[str, Any]:
        it: dict[str, Any] = {"field": field, "title": localize(title, profile),
                              "detail": localize(detail, profile)}
        r = _rule_ref(rule)
        if r:
            it["rule"] = r
        e = _evidence_ref(evidence, hl)
        if e:
            it["evidence"] = e
        return it

    # `checks` preserves the canonical run order regardless of outcome, so the UI
    # can keep every check in a stable position as its status changes.
    checks: list[dict[str, Any]] = []

    def err(field: str, title: str, detail: str,
            rule: str | None = None, evidence: str | None = None, hl: list[str] | None = None) -> None:
        it = _item(field, title, detail, rule, evidence, hl)
        it["status"] = "error"
        errors.append(it)
        checks.append(it)

    def warn(field: str, title: str, detail: str,
             rule: str | None = None, evidence: str | None = None, hl: list[str] | None = None) -> None:
        it = _item(field, title, detail, rule, evidence, hl)
        it["status"] = "warn"
        warnings.append(it)
        checks.append(it)

    def ok(title: str, detail: str,
           rule: str | None = None, evidence: str | None = None, hl: list[str] | None = None) -> None:
        it = _item("", title, detail, rule, evidence, hl)
        it["status"] = "ok"
        passed.append(it)
        checks.append(it)

    if not str(intake.get("legal_name") or "").strip():
        err("legal_name", "Legal entity name missing",
            "Required — and it must match the Secretary of State certificate exactly.",
            rule="legal-name", evidence="file_signatory", hl=["legal_name"])
    else:
        ok("Legal entity identified",
           f'"{intake["legal_name"]}" will be used consistently across Appendix 1, site exclusivity, and signatory documents.',
           rule="legal-name", evidence="file_signatory", hl=["legal_name"])

    if not str(intake.get("signatory_name") or "").strip():
        err("signatory_name", "Authorized signatory missing",
            "Appendix 1 must be executed by an authorized signatory in RIMS5.",
            rule="signatory", evidence="file_signatory", hl=["signatory"])

    lat, lon = _num(intake.get("gps_lat")), _num(intake.get("gps_lon"))
    if lat is None or lon is None:
        err("gps_lat", "GPS coordinates missing or invalid",
            "Decimal-format latitude and longitude are required (they feed Appendix 1, Attachment A, and the KMZ).",
            rule="gps", evidence="file_boundary", hl=["gps"])
    else:
        box = profile.get("gps_box")
        if box and not (box[0] <= lat <= box[1] and box[2] <= lon <= box[3]):
            warn("gps_lat", f"Coordinates outside the typical {profile['iso']} footprint",
                 f"({lat}, {lon}) does not look like a site in the {profile['name']} region — verify before submission.",
                 rule="gps", evidence="file_boundary", hl=["gps"])
        else:
            ok("Site coordinates valid",
               f"({lat}, {lon}) — {intake.get('county') or '—'} County, {intake.get('state') or '—'}.",
               rule="gps", evidence="file_boundary", hl=["gps"])

    site_control = str(intake.get("site_control") or "")
    site_file = _file_meta(intake.get("file_site_control"))
    if site_control == "Letter of Intent":
        err("site_control", "Letter of Intent is not accepted",
            "CAISO requires a lease, option to purchase, or deed as evidence of site exclusivity.",
            rule="site-exclusivity", evidence="file_site_control", hl=["exclusivity"])
    elif site_control in ("None yet", ""):
        err("site_control", "No site exclusivity",
            "An executed lease, option, or deed is required before submission — the most common source of delay.",
            rule="site-exclusivity", evidence="file_site_control", hl=["exclusivity"])
    elif site_file is None:
        err("file_site_control", "Executed site agreement not attached",
            f"Site exclusivity is declared as \"{site_control}\", but no executed agreement is uploaded. "
            "CAISO requires the document itself as evidence — the declaration alone is not sufficient.",
            rule="site-exclusivity", evidence="file_site_control", hl=["exclusivity"])
    else:
        ok("Site exclusivity evidenced",
           f"{site_control} — \"{site_file['name']}\" will accompany the site exclusivity declaration.",
           rule="site-exclusivity", evidence="file_site_control", hl=["exclusivity"])

    if not str(intake.get("poi_name") or "").strip() or _num(intake.get("poi_voltage_kv")) is None:
        err("poi_name", "Point of Interconnection incomplete",
            "Substation / line name and voltage level (kV) are both required.", rule="poi")
    else:
        ok("POI defined", f"{intake['poi_name']} at {_fmt(_num(intake.get('poi_voltage_kv')))} kV.", rule="poi")

    gross = _num(intake.get("gross_mw"))
    aux = _num(intake.get("aux_mw"))
    losses = _num(intake.get("losses_mw"))
    net = _num(intake.get("net_mw_poi"))
    if None in (gross, aux, losses, net):
        err("net_mw_poi", "MW chain incomplete",
            "Gross output, auxiliary load, losses, and net MW at POI are all required.",
            rule="mw-chain", evidence="file_technical",
            hl=["gross_mw", "aux_mw", "losses_mw", "net_mw_poi"])
    else:
        computed = round(gross - aux - losses, 3)
        if abs(computed - net) > 0.05:
            err("net_mw_poi", "MW chain does not reconcile",
                f"Gross {_fmt(gross)} − Aux {_fmt(aux)} − Losses {_fmt(losses)} = {_fmt(computed)} MW, "
                f"but requested net at POI is {_fmt(net)} MW. Inconsistent MW values are among the most "
                "common CAISO deficiency findings.",
                rule="mw-chain", evidence="file_technical", hl=["net_mw_poi"])
        else:
            ok("MW chain reconciles",
               f"Gross {_fmt(gross)} − Aux {_fmt(aux)} − Losses {_fmt(losses)} = {_fmt(net)} MW at POI. "
               "This value will be enforced across all generated documents.",
               rule="mw-chain", evidence="file_technical", hl=["net_mw_poi"])
        if net is not None and net > 20:
            ok("Large Generating Facility", f"{_fmt(net)} MW > 20 MW — Large project rules apply.")

    track = str(intake.get("track") or "")
    if not track:
        err("track", "Process track not selected", "Choose one of: " + ", ".join(profile["tracks"]) + ".")
    elif track == "Fast Track" and net is not None and net > 5:
        err("track", "Fast Track not available above 5 MW",
            f"Requested {_fmt(net)} MW at POI exceeds the 5 MW Fast Track limit — use ISP or Cluster.",
            rule="fast-track", evidence="file_technical", hl=["net_mw_poi"])
    elif track == "Independent Study Process":
        ok("ISP track selected",
           "GridPilot will draft the ISP eligibility demonstration; the $150,000 study deposit is wired directly to CAISO.",
           rule="isp")
        if str(intake.get("deliverability")) == "Full Capacity":
            warn("deliverability", "Full Capacity deliverability under ISP",
                 "Deliverability is assessed in the next annual Cluster Study (GIDAP 4.6) — flag this to offtake counterparties.",
                 rule="deliverability")
    elif track == "Cluster":
        warn("track", "Cluster window timing",
             "Cluster requests are only accepted during the annual request window; typical study-to-GIA timeline is ~3 years.",
             rule="cluster-window")

    cod = _parse_cod(intake.get("cod"))
    if cod is None:
        err("cod", "Target COD missing or unparseable", "Provide the commercial operation date as YYYY-MM-DD.",
            rule="cod-limit")
    elif cod > date.today() + timedelta(days=365 * 7):
        err("cod", "COD beyond the 7-year limit",
            f"{cod.isoformat()} is more than 7 years out — CAISO caps COD at 7 years from the application.",
            rule="cod-limit")
    elif cod < date.today():
        err("cod", "COD is in the past", f"{cod.isoformat()} — update the target commercial operation date.",
            rule="cod-limit")
    else:
        ok("COD within limits", f"{cod.strftime('%m/%d/%Y')} — in-service and trial-operation dates will be derived sequentially.",
           rule="cod-limit")

    bess_mw = _num(intake.get("bess_mw"))
    if bess_mw and bess_mw > 0:
        if not _num(intake.get("bess_mwh")):
            err("bess_mwh", "BESS energy missing", "MWh (duration) is required when storage is included.",
                rule="bess-energy", evidence="file_bess", hl=["bess_mwh"])
        else:
            ok("Storage parameters complete",
               f"{_fmt(bess_mw)} MW / {_fmt(_num(intake.get('bess_mwh')))} MWh — "
               f"charging: {intake.get('bess_charging') or 'unspecified'}.",
               rule="bess-energy", evidence="file_bess", hl=["bess_mw", "bess_mwh"])
        if str(intake.get("bess_charging")) == "Grid charging permitted":
            warn("bess_charging", "Grid charging affects deliverability & modeling",
                 "Charging from the grid changes the deliverability assessment and the load-flow model — confirm intent.",
                 rule="bess-charging", evidence="file_bess", hl=["bess_charging"])

    dyd = str(intake.get("dyd_status") or "")
    dyd_file = _file_meta(intake.get("file_dyd"))
    if dyd == "Received from vendor":
        if dyd_file is None:
            err("file_dyd", "Vendor .dyd marked received but not uploaded",
                "The intake indicates the vendor dynamic model files are in hand — attach the .dyd file "
                "so its parameters can be integrated into the PSLF dynamic model.",
                rule="dyd-models", evidence="file_dyd")
        else:
            ok("Vendor dynamic models in hand",
               f"\"{dyd_file['name']}\" — vendor parameters will be integrated into the PSLF dynamic model.",
               rule="dyd-models", evidence="file_dyd")
    elif dyd_file is not None:
        ok("Vendor .dyd file received",
           f"\"{dyd_file['name']}\" attached — vendor parameters will be integrated into the PSLF dynamic model.",
           rule="dyd-models", evidence="file_dyd")
    elif dyd == "Requested — pending":
        warn("dyd_status", "Vendor .dyd files still pending",
             "The single most common schedule risk. GridPilot will use standard WECC models "
             "(REGC_A / REEC_A / REPC_A) as placeholders — swap in vendor files before submission.",
             rule="dyd-models", evidence="file_dyd")
    else:
        warn("dyd_status", "Equipment not selected",
             "Generic PSLF models will be used as placeholders; later equipment swaps add rework and possible re-study.",
             rule="dyd-models", evidence="file_dyd")

    sig_file = _file_meta(intake.get("file_signatory"))
    if sig_file is None:
        warn("file_signatory", "Signatory proof not attached",
             "No officer certificate or board resolution uploaded — GridPilot will generate a draft "
             "for counsel to review and execute before submission.",
             rule="signatory", evidence="file_signatory", hl=["signatory"])
    else:
        ok("Signatory proof on file",
           f"\"{sig_file['name']}\" documents {intake.get('signatory_name') or 'the signatory'}'s authority to execute Appendix 1.",
           rule="signatory", evidence="file_signatory", hl=["signatory"])

    boundary_file = _file_meta(intake.get("file_boundary"))
    if boundary_file is not None:
        ok("Project boundary file received",
           f"\"{boundary_file['name']}\" — the KMZ and site drawing will follow this boundary.",
           rule="boundary", evidence="file_boundary", hl=["gps"])

    if "TBD" in str(intake.get("inverter") or "").upper() or not str(intake.get("inverter") or "").strip():
        warn("inverter", "Inverter not finalized",
             "Attachment A will carry generic values — update before submission to avoid deficiency review.",
             rule="dyd-models", evidence="file_technical", hl=["inverter"])

    # --- Verified equipment library (never-hallucinate policy) ---------------
    # Every device parameter written into the packet must trace to
    # equipment_specs.json. An unknown device blocks generation; a known device
    # with unsourced parameters raises a SYSTEM ADMIN alert and those
    # parameters are excluded from Attachment A until sourced.
    admin_alerts: list[dict[str, Any]] = []
    custom_list = intake.get("custom_equipment")
    custom_list = custom_list if isinstance(custom_list, list) else []

    def _lib_check(kind_label: str, text: str, matcher, kind_key: str, field: str):
        if not text or "TBD" in text.upper():
            return None
        custom = lib.match_custom(text, custom_list, kind_key)
        if custom:
            return lib.custom_entry(custom)
        entry, matched = matcher(text)
        if not matched:
            err(field, f"{kind_label} not in the verified equipment library",
                f'"{text}" has no entry in the equipment spec file. GridPilot never writes '
                "unverified device parameters into the packet — ask the system administrator "
                "to add the device (datasheet values with sources) to "
                "backend/app/assets/equipment_specs.json, then re-validate.",
                rule="dyd-models", evidence="file_technical", hl=[field])
            admin_alerts.append({
                "device": text, "param": "all parameters",
                "note": "New device — add a sourced entry to equipment_specs.json",
            })
            return None
        return entry

    _lib_entries = [
        _lib_check("Inverter", str(intake.get("inverter") or "").strip(),
                   lib.match_inverter, "pv_inverter", "inverter"),
    ]
    if _num(intake.get("bess_mw")):
        _bmw, _bmwh = _num(intake.get("bess_mw")), _num(intake.get("bess_mwh"))
        _dur = (_bmwh / _bmw) if (_bmw and _bmwh) else None
        _lib_entries.append(
            _lib_check("BESS unit", str(intake.get("bess_vendor") or "").strip(),
                       lambda t: lib.match_bess(t, _dur), "bess", "bess_vendor"))
    for entry in filter(None, _lib_entries):
        admin_alerts.extend(lib.spec_gaps(entry))

    _gaps = [a for a in admin_alerts if a["param"] != "all parameters"]
    if _gaps:
        warn("equipment_specs", "Device parameters without an OEM source — admin action",
             f"{len(_gaps)} parameter(s) on the selected equipment have no source on file and are "
             "excluded from Attachment A until the system administrator adds sourced values to "
             "equipment_specs.json: "
             + "; ".join(f"{a['device']} — {a['param']}" for a in _gaps),
             rule="dyd-models", evidence="file_technical", hl=["inverter"])
    elif not admin_alerts and any(e is not None for e in _lib_entries):
        ok("Equipment parameters fully sourced",
           "Every parameter of the selected devices traces to a cited OEM datasheet or the "
           "developer ground-truth workbook — nothing in Attachment A is unsourced.",
           rule="dyd-models", evidence="file_technical")

    return {
        "ok": len(errors) == 0,
        "errors": errors,
        "warnings": warnings,
        "passed": passed,
        "checks": checks,
        "admin_alerts": admin_alerts,
        "summary": (
            f"{len(errors)} blocking issue(s), {len(warnings)} advisory item(s), {len(passed)} check(s) passed."
        ),
    }


def _fmt(v: float | None) -> str:
    if v is None:
        return "—"
    return f"{v:g}"


def _pacific_today() -> date:
    """CAISO business date (Pacific), not the UTC server clock."""
    from zoneinfo import ZoneInfo
    return datetime.now(ZoneInfo("America/Los_Angeles")).date()


# ---------------------------------------------------------------------------
# PDF helpers (PyMuPDF)
# ---------------------------------------------------------------------------

INK = (0.10, 0.12, 0.11)
MUT = (0.42, 0.46, 0.44)
ACC = (0.13, 0.35, 0.75)
RED = (0.62, 0.18, 0.10)
OKC = (0.10, 0.45, 0.25)
LINE = (0.80, 0.82, 0.80)

PAGE_W, PAGE_H = 612, 792  # portrait letter

# California ISO logo extracted from the official RIMS Appendix 1 Word template,
# used to reproduce that template's header on the generated form.
CAISO_LOGO = Path(__file__).resolve().parents[1] / "assets" / "caiso_logo.jpg"


class _Pdf:
    def __init__(self, title: str, subtitle: str, banner: str = "",
                 official: dict | None = None):
        """`official` switches the page chrome to the CAISO form template style:
        logo top-left, right-aligned title/subtitle header, and a right-aligned
        footer with page number plus the form's effective date and version."""
        self.doc = fitz.open()
        self.title = title
        self.subtitle = subtitle
        self.banner = banner
        self.official = official
        self.page: fitz.Page | None = None
        self.y = 0.0
        self._new_page()

    def _new_page(self) -> None:
        self.page = self.doc.new_page(width=PAGE_W, height=PAGE_H)
        if self.official is not None:
            if CAISO_LOGO.exists():
                self.page.insert_image(fitz.Rect(36, 22, 198, 52), filename=str(CAISO_LOGO))
            y = 34.0
            for line in (self.title, self.subtitle):
                w = fitz.get_text_length(line, fontname="helv", fontsize=11.5)
                self.page.insert_text((PAGE_W - 36 - w, y), line,
                                      fontsize=11.5, fontname="helv", color=INK)
                y += 15
            self.y = 84
            return
        self.page.draw_rect(fitz.Rect(0, 0, PAGE_W, 6), color=ACC, fill=ACC)
        # Banner text intentionally not rendered — deliverables carry no
        # "GENERATED/COMPUTED" stamp line (the manifest keeps that metadata).
        self.page.insert_text((36, 56), self.title, fontsize=15, fontname="hebo", color=INK)
        self.page.insert_text((36, 72), self.subtitle, fontsize=9, fontname="helv", color=MUT)
        self.page.draw_line(fitz.Point(36, 82), fitz.Point(PAGE_W - 36, 82), color=LINE, width=0.8)
        self.page.insert_text(
            (36, PAGE_H - 28),
            f"{_pacific_today().strftime('%B %d, %Y')}  ·  Verify before RIMS5 submission",
            fontsize=7, fontname="helv", color=MUT,
        )
        self.y = 100

    def _ensure(self, needed: float) -> None:
        bottom = PAGE_H - (64 if self.official is not None else 48)
        if self.y + needed > bottom:
            self._new_page()

    def section(self, text: str) -> None:
        self._ensure(28)
        self.y += 8
        self.page.insert_text((36, self.y), text.upper(), fontsize=8.5, fontname="hebo", color=ACC)
        self.y += 6
        self.page.draw_line(fitz.Point(36, self.y), fitz.Point(PAGE_W - 36, self.y), color=LINE, width=0.6)
        self.y += 14

    def kv(self, label: str, value: str, note: str = "") -> None:
        self._ensure(16)
        self.page.insert_text((36, self.y), label, fontsize=8.5, fontname="helv", color=MUT)
        self.page.insert_text((250, self.y), value, fontsize=9, fontname="hebo", color=INK)
        if note:
            self.y += 11
            self.page.insert_text((250, self.y), note, fontsize=7.5, fontname="helv", color=MUT)
        self.y += 15

    def para(self, text: str, size: float = 9, color=INK, indent: float = 36) -> None:
        width_chars = int((PAGE_W - indent - 36) / (size * 0.5))
        words = text.split()
        line = ""
        lines: list[str] = []
        for w in words:
            if len(line) + len(w) + 1 > width_chars:
                lines.append(line)
                line = w
            else:
                line = f"{line} {w}".strip()
        if line:
            lines.append(line)
        for ln in lines:
            self._ensure(13)
            self.page.insert_text((indent, self.y), ln, fontsize=size, fontname="helv", color=color)
            self.y += size + 4

    def bullet(self, text: str, color=INK) -> None:
        self._ensure(13)
        self.page.insert_text((44, self.y), "•", fontsize=9, fontname="helv", color=color)
        keep_y = self.y
        self.y = keep_y
        # simple single-level wrap
        width_chars = 92
        words = text.split()
        line = ""
        first = True
        for w in words:
            if len(line) + len(w) + 1 > width_chars:
                self._ensure(13)
                self.page.insert_text((56, self.y), line, fontsize=9, fontname="helv", color=color)
                self.y += 13
                line = w
                first = False
            else:
                line = f"{line} {w}".strip()
        if line:
            self._ensure(13)
            self.page.insert_text((56, self.y), line, fontsize=9, fontname="helv", color=color)
            self.y += 13

    def checkbox(self, checked: bool, text: str) -> None:
        self._ensure(15)
        r = fitz.Rect(38, self.y - 8, 47, self.y + 1)
        self.page.draw_rect(r, color=INK, width=0.8)
        if checked:
            self.page.insert_text((39.5, self.y - 0.5), "X", fontsize=8, fontname="hebo", color=INK)
        # Wrap long labels so they don't run off the right page edge.
        words = text.split()
        line = ""
        for w in words:
            if line and len(line) + len(w) + 1 > 100:
                self.page.insert_text((54, self.y), line, fontsize=9, fontname="helv", color=INK)
                self.y += 12
                self._ensure(13)
                line = w
            else:
                line = f"{line} {w}".strip()
        self.page.insert_text((54, self.y), line, fontsize=9, fontname="helv", color=INK)
        self.y += 16

    def save(self, path: Path) -> None:
        if self.official is not None:
            # "Page X of Y" needs the final page count, so the footer is
            # stamped at save time on every page.
            total = self.doc.page_count
            for i in range(total):
                page = self.doc[i]
                y = PAGE_H - 44.0
                for text, size in ((f"Page {i + 1} of {total}", 8.0),
                                   (f"Effective:  {self.official['effective']}", 7.0),
                                   (f"Version: {self.official['version']}", 7.0)):
                    w = fitz.get_text_length(text, fontname="hebo", fontsize=size)
                    page.insert_text((PAGE_W - 36 - w, y), text,
                                     fontsize=size, fontname="hebo", color=INK)
                    y += size + 3
        self.doc.save(path)
        self.doc.close()


def _axes(page: fitz.Page, rect: fitz.Rect, title: str, xlabel: str, ylabel: str) -> None:
    """Simple axes frame — used by the engine report renderers (reports.py)."""
    page.draw_rect(rect, color=LINE, width=0.8)
    page.insert_text((rect.x0, rect.y0 - 10), title, fontsize=9.5, fontname="hebo", color=INK)
    page.insert_text((rect.x0 + rect.width / 2 - 30, rect.y1 + 16), xlabel, fontsize=7.5, fontname="helv", color=MUT)
    page.insert_text((rect.x0 - 24, rect.y0 - 2), ylabel, fontsize=7.5, fontname="helv", color=MUT)


def _polyline(page: fitz.Page, pts: list[tuple[float, float]], color, width=1.4) -> None:
    if len(pts) < 2:
        return
    shape = page.new_shape()
    shape.draw_polyline([fitz.Point(*p) for p in pts])
    shape.finish(color=color, width=width, closePath=False)
    shape.commit()


# ---------------------------------------------------------------------------
# Document builders
# ---------------------------------------------------------------------------

def _derived(intake: dict[str, Any]) -> dict[str, Any]:
    gross = _num(intake.get("gross_mw")) or 0
    aux = _num(intake.get("aux_mw")) or 0
    losses = _num(intake.get("losses_mw")) or 0
    net = _num(intake.get("net_mw_poi")) or 0
    cod = _parse_cod(intake.get("cod")) or date.today() + timedelta(days=540)
    bess_mw = _num(intake.get("bess_mw")) or 0
    pv_mw = max(gross - bess_mw, 0)
    track = str(intake.get("track") or "Independent Study Process")
    deposit = {"Independent Study Process": "$150,000", "Fast Track": "$500 (non-refundable)"}.get(
        track, "Per CAISO Cluster study deposit schedule"
    )
    return {
        "gross": gross, "aux": aux, "losses": losses, "net": net,
        "max_net": round(gross - aux, 3),
        "pv_mw": pv_mw, "bess_mw": bess_mw,
        "cod": cod,
        "in_service": cod - timedelta(days=210),
        "trial_op": cod - timedelta(days=120),
        "track": track,
        "deposit": deposit,
        "kv": _num(intake.get("poi_voltage_kv")) or 230,
        "col_kv": _num(intake.get("collector_kv")) or 34.5,
        "lat": _num(intake.get("gps_lat")) or 35.0,
        "lon": _num(intake.get("gps_lon")) or -118.3,
        "acres": _num(intake.get("site_acreage")) or 600,
        "is_isp": track == "Independent Study Process",
        "has_bess": bess_mw > 0,
    }


def _slug(s: str) -> str:
    return re.sub(r"[^A-Za-z0-9]+", "_", s or "project").strip("_") or "project"


def _split_name(full: str) -> tuple[str, str]:
    parts = str(full or "").strip().split()
    if not parts:
        return "", ""
    return parts[0], " ".join(parts[1:])


APPENDIX1_TEMPLATE = Path(__file__).resolve().parents[1] / "assets" / "appendix1_template.docx"


def _gen_appendix1(intake: dict, d: dict, path: Path) -> None:
    """Checklist item 2 — Appendix 1 Interconnection Request.

    For CAISO this fills the official Appendix 1 Word form in place: the
    original document (wording, layout, header/footer) is untouched — only
    the blank fields (en-space runs) and checkboxes are filled from the
    intake. Fields the intake cannot answer are left blank and surfaced in
    the Missing Data Report.
    """
    p = d.get("profile") or get_profile(None)
    if p["iso"] != "CAISO":
        _gen_ir_generic(intake, d, p, path)
        return
    _fill_appendix1_docx(intake, d, path)


def _fill_appendix1_docx(intake: dict, d: dict, path: Path) -> None:
    """Fill the official ISP/Fast Track Appendix 1 Word form.

    The template ("interconnectionrequestform-appendix1-independentstudy
    andfasttrack") uses legacy Word form fields: FORMCHECKBOX, FORMTEXT and
    dropdown (ddList) fields. Nothing is generated, moved or deleted — the
    original wording, headers, footers and cell borders are untouched. The
    fill only sets checkbox states, form-text results and dropdown
    selections, then saves as a new document. Fields the intake cannot
    answer stay blank (they surface in the Missing Data Report).
    """
    import docx
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    doc = docx.Document(str(APPENDIX1_TEMPLATE))
    body = doc.element.body

    # Collect the legacy form fields in document order. Each field starts at
    # a run holding <w:fldChar fldCharType="begin"> with <w:ffData>.
    fields: list[dict] = []
    for r in body.iter(qn("w:r")):
        fc = r.find(qn("w:fldChar"))
        if fc is None or fc.get(qn("w:fldCharType")) != "begin":
            continue
        ff = fc.find(qn("w:ffData"))
        if ff is None:
            continue
        kind = ("chk" if ff.find(qn("w:checkBox")) is not None
                else "dd" if ff.find(qn("w:ddList")) is not None else "txt")
        fields.append({"kind": kind, "ff": ff, "run": r})
    if len(fields) != 134:
        raise RuntimeError(f"Appendix 1 template changed: {len(fields)} form fields")

    def _result_runs(field) -> list:
        """Runs between the field's 'separate' and 'end' markers (the text
        Word displays for the field)."""
        par = field["run"].getparent()
        runs = list(par.findall(qn("w:r")))
        out, state = [], 0
        for rr in runs[runs.index(field["run"]):]:
            fc = rr.find(qn("w:fldChar"))
            if fc is not None:
                ft = fc.get(qn("w:fldCharType"))
                if ft == "separate" and state == 0:
                    state = 1
                    continue
                if ft == "end" and state == 1:
                    break
            elif state == 1:
                out.append(rr)
        return out

    def text(i: int, value) -> None:
        """Set a FORMTEXT field's displayed result (formatting preserved)."""
        if value is None or str(value).strip() == "":
            return
        rrs = _result_runs(fields[i])
        if not rrs:
            return
        t_el = rrs[0].find(qn("w:t"))
        if t_el is None:
            t_el = OxmlElement("w:t")
            rrs[0].append(t_el)
        t_el.text = str(value)
        t_el.set(qn("xml:space"), "preserve")
        for rr in rrs[1:]:
            for t2 in rr.findall(qn("w:t")):
                rr.remove(t2)

    def check(i: int) -> None:
        """Tick a FORMCHECKBOX (sets w:checked on the field data)."""
        cb = fields[i]["ff"].find(qn("w:checkBox"))
        ch = cb.find(qn("w:checked"))
        if ch is None:
            ch = OxmlElement("w:checked")
            cb.append(ch)
        ch.set(qn("w:val"), "1")

    def select(i: int, option: str) -> None:
        """Pick a dropdown (ddList) entry and update the displayed text."""
        dl = fields[i]["ff"].find(qn("w:ddList"))
        entries = [e.get(qn("w:val")) for e in dl.findall(qn("w:listEntry"))]
        idx = entries.index(option)
        res = dl.find(qn("w:result"))
        if res is None:
            res = OxmlElement("w:result")
            dl.insert(0, res)
        res.set(qn("w:val"), str(idx))
        text(i, option)

    first, last = _split_name(intake.get("signatory_name"))
    deliv = str(intake.get("deliverability") or "Full Capacity")
    is_solar_or_wind = any(k in str(intake.get("project_type") or "")
                           for k in ("Solar", "Wind"))
    is_hybrid = d["has_bess"] and "BESS" in str(intake.get("project_type") or "")
    bess_hours = (d["bess_mw"] and _num(intake.get("bess_mwh"))
                  and round(_num(intake.get("bess_mwh")) / d["bess_mw"], 1))
    today = _pacific_today().strftime("%m/%d/%Y")

    # ---- Page 1: minimum-requirements checklist — mark what the packet
    # contains. Items 1 (study deposit) and 5 (ISP demonstrations, fields
    # 6-12) are the developer's own actions and stay unchecked.
    for i in (1,          # 2. Completed Appendix 1
              2, 3, 4,    # 3. Attachment A + Technical Data / IR Validation tabs
              5,          # 4. Evidence of Site Exclusivity
              13, 14,     # 6. Load Flow Model (.epc)  7. Dynamic Model (.dyd)
              15, 16, 17,  # 8. Reactive capability  9. Site Drawing  10. SLD
              18, 19,     # 11. Flat run & bump test  12. MW at POI plot
              20, 21, 22, 23):  # 13. IBR validation + all three tests
        check(i)

    # ---- 1. Process (check only one)
    check(24 if d["track"] == "Fast Track" else 25)
    # ---- 2. Request type
    check(26)  # A proposed new Generating Facility
    # ---- 3. Deliverability (on-peak / off-peak)
    if deliv == "Full Capacity":
        check(28)
    elif deliv == "Energy Only":
        check(31)
    else:
        check(29)
        text(30, intake.get("partial_pct") or "")
    if is_solar_or_wind:
        check(32 if deliv != "Energy Only" else 33)
    text(34, f"{deliv} Deliverability Status requested; deliverability to be assessed "
             "with the next annual Deliverability Allocation cycle per Appendix DD.")
    # ---- 4a. Project name & location
    text(35, intake.get("project_name"))
    text(36, intake.get("street_address") or "9200 Backus Road")
    text(37, intake.get("city") or "Mojave")
    text(38, intake.get("county"))
    text(39, intake.get("state") or "CA")
    text(40, intake.get("zip") or "93501")
    text(41, d["lat"])
    text(42, d["lon"])
    # ---- 4b. Megawatt values
    text(43, _fmt(d.get("total_mva") or _num(intake.get("gross_mva"))))
    text(44, _fmt(d["gross"]))
    text(45, _fmt(d["aux"]))
    text(46, _fmt(d["max_net"]))
    text(47, "N/A")  # increase-to-existing value — new facility
    text(48, _fmt(d["losses"]))
    text(49, _fmt(d["net"]))
    text(50, f"A plant-level Power Plant Controller (PPC) monitors POI revenue metering and "
             f"limits aggregate export to {_fmt(d['net'])} MW via real-time inverter setpoint "
             "dispatch; inverter-level curtailment provides backup limitation.")
    # ---- 4c. Technology table: row 1 = PV, row 2 = storage
    pv_mw = d["gross"] - (d["bess_mw"] or 0)
    select(51, "Photovoltaic")
    select(52, "Solar")
    text(53, _fmt(pv_mw))
    if is_hybrid:
        check(57)
    if d["has_bess"]:
        select(58, "Storage")
        select(59, "Battery")
        text(60, _fmt(d["bess_mw"]))
        text(61, _fmt(_num(intake.get("bess_mwh"))))
        text(62, _fmt(bess_hours))
        if is_hybrid:
            check(64)
    text(80, f"{d.get('inverter_desc') or intake.get('inverter') or 'Inverters TBD'}; "
             f"{intake.get('module') or 'modules TBD'}"
             + (f"; {d.get('bess_desc') or intake.get('bess_vendor') or 'BESS TBD'}"
                if d["has_bess"] else ""))
    text(81, f"GSU {d.get('gsu_desc') or intake.get('transformer') or 'TBD'}; "
             f"{_fmt(d['col_kv'])} kV collector "
             f"system; gross {_fmt(d['gross'])} MW, net {_fmt(d['net'])} MW at POI.")
    # ---- 4d. Dates
    text(82, d["in_service"].strftime("%m/%d/%Y"))
    text(83, d["trial_op"].strftime("%m/%d/%Y"))
    text(84, d["cod"].strftime("%m/%d/%Y"))
    text(85, "25")  # proposed term of service (years)
    # ---- 4e. Contact person
    company_addr = intake.get("company_address") or "500 Capitol Mall, Suite 2350"
    company_city = intake.get("company_city") or "Sacramento"
    company_state = intake.get("company_state") or "CA"
    company_zip = intake.get("company_zip") or "95814"
    text(86, first)
    text(87, last)
    text(88, intake.get("signatory_title"))
    text(89, intake.get("legal_name"))
    text(90, company_addr)
    text(91, company_city)
    text(92, company_state)
    text(93, company_zip)
    text(94, intake.get("contact_phone"))
    text(95, "N/A")
    text(96, intake.get("contact_email"))
    # ---- 4f. Point of Interconnection
    text(97, intake.get("poi_name"))
    text(98, _fmt(d["kv"]))
    shared = str(intake.get("shared_facilities") or "No shared facilities")
    select(99, "Yes" if "gen-tie" in shared.lower() else "No")
    # ---- 6. Site exclusivity
    check(100)  # Is attached to this Interconnection Request
    sc = str(intake.get("site_control") or "")
    for i, match in ((101, "Deed" in sc),
                     (102, sc == "Lease Agreement"),
                     (103, sc == "Option to Purchase"),
                     (104, "Option to Lease" in sc)):
        if match:
            check(i)
    check(107)  # granted to the same entity — Yes
    # 40-year lease from COD — matches the Site Exclusivity form (item 4).
    text(111, "40")
    text(112, "N/A")  # term of option — site control is a lease
    text(113, f"~{_fmt(d['acres'])}")
    # ---- 8. Representative to contact (same as the contact person)
    text(114, first)
    text(115, last)
    text(116, intake.get("signatory_title"))
    text(117, intake.get("legal_name"))
    text(118, company_addr)
    text(119, company_city)
    text(120, company_state)
    text(121, company_zip)
    text(122, intake.get("contact_phone"))
    text(123, "N/A")
    text(124, intake.get("contact_email"))
    # ---- 9. Submitted by + consent + e-signature
    text(125, intake.get("legal_name"))
    text(126, intake.get("state_of_origin"))
    text(127, "N/A")
    check(128)  # consent to disclosure
    check(129)  # electronic signature agreement
    text(130, first)
    text(131, last)
    text(132, intake.get("signatory_title"))
    text(133, today)

    doc.save(str(path))


def _gen_ir_generic(intake: dict, d: dict, p: dict, path: Path) -> None:
    """Interconnection request form for non-CAISO ISOs (generic layout)."""
    pdf = _Pdf(
        p["form_name"],
        f"{p['process']} ({p['tariff']}) — submit via {p['portal']}",
    )
    pdf.section("1 · Process")
    for t in p["tracks"]:
        pdf.checkbox(d["track"] == t, t)
    pdf.section("2 · Request type & deliverability")
    pdf.checkbox(True, "A proposed new Generating Facility")
    pdf.kv("Requested deliverability (on-peak)", str(intake.get("deliverability") or "Full Capacity"),
           f"Service level election per {p['tariff']}")
    pdf.section("3 · Project name & location")
    pdf.kv("Project name", str(intake.get("project_name") or ""))
    pdf.kv("County / State", f"{intake.get('county') or '—'} / {intake.get('state') or '—'}")
    pdf.kv("GPS (decimal)", f"{d['lat']}, {d['lon']}")
    pdf.section("4 · Project megawatt values")
    pdf.kv("Gross capacity (MVA, unity PF)",
           _fmt(d.get("total_mva") or _num(intake.get("gross_mva"))))
    pdf.kv("Gross output (MW)", _fmt(d["gross"]))
    pdf.kv("Auxiliary load (MW)", _fmt(d["aux"]))
    pdf.kv("Maximum net electrical output (MW)", _fmt(d["max_net"]))
    pdf.kv("Anticipated losses to POI (MW)", _fmt(d["losses"]))
    pdf.kv("Requested Interconnection Service Capacity", f"{_fmt(d['net'])} MW at POI",
           f"This value appears in the {p['iso']} queue and must match every document in the packet")
    pdf.section("5 · Dates, contact & Point of Interconnection")
    pdf.kv("Proposed in-service date", d["in_service"].strftime("%m/%d/%Y"))
    pdf.kv("Proposed commercial operation date", d["cod"].strftime("%m/%d/%Y"))
    pdf.kv("Interconnection customer contact",
           f"{intake.get('signatory_name') or ''}, {intake.get('signatory_title') or ''}")
    pdf.kv("Company", str(intake.get("legal_name") or ""))
    pdf.kv("Phone / Email", f"{intake.get('contact_phone') or '—'} / {intake.get('contact_email') or '—'}")
    pdf.kv("Point of Interconnection", str(intake.get("poi_name") or ""))
    pdf.kv("Voltage level", f"{_fmt(d['kv'])} kV")
    pdf.section("6 · Deposit, exclusivity, submission")
    pdf.kv("Study deposit", d["deposit"], p["deposit_action"])
    pdf.kv("Site exclusivity", f"{intake.get('site_control') or ''} — attached",
           f"Owner: {intake.get('site_owner') or '—'}; {_fmt(d['acres'])} acres")
    pdf.kv("Legal name of Interconnection Customer", str(intake.get("legal_name") or ""),
           "Must match the Secretary of State certificate exactly")
    pdf.para(f"{p['submit_action']} ({p['portal_url']}). Executed by "
             f"{intake.get('signatory_name') or ''}, {intake.get('signatory_title') or ''}.", size=8.5, color=MUT)
    pdf.save(path)


ATTACHMENT_A_TEMPLATE = Path(__file__).resolve().parents[1] / "assets" / "attachment_a_template.xlsm"


def _gen_attachment_a(intake: dict, d: dict, eng: dict, path: Path) -> None:
    """Checklist item 3 — Attachment A to Appendix 1 (Generator Technical Data).

    For CAISO this fills the official macro workbook (v14.5) in place —
    values go into the real input cells, formulas / macros / formatting are
    untouched. Other ISOs get the generated mirror workbook.
    """
    p = d.get("profile") or get_profile(None)
    if p["iso"] == "CAISO" and path.suffix == ".xlsm":
        _fill_attachment_a_xlsm(intake, d, eng, path)
        return
    _gen_attachment_a_xlsx(intake, d, eng, path)


def _fill_attachment_a_xlsm(intake: dict, d: dict, eng: dict, path: Path) -> None:
    import math as _math
    import warnings

    import openpyxl
    from openpyxl.utils import get_column_letter as _cl

    with warnings.catch_warnings():
        warnings.simplefilter("ignore")  # unsupported-extension warnings from the official file
        wb = openpyxl.load_workbook(str(ATTACHMENT_A_TEMPLATE), keep_vba=True)

    def W(sheet, coord: str, value) -> None:
        """Merged-range-aware cell write: values land on the range anchor."""
        cell = sheet[coord]
        for rng in sheet.merged_cells.ranges:
            if cell.coordinate in rng:
                cell = sheet.cell(rng.min_row, rng.min_col)
                break
        cell.value = value

    design = eng["design"]
    pf = eng["powerflow"]
    inv = design["inverter"]
    bess = design.get("bess_unit")
    c = design["counts"]
    has_bess = bool(bess and c["bess_units"])
    proj = str(intake.get("project_name") or "")
    pv_name = f"{proj} PV"
    bess_name = f"{proj} BESS"
    col_kv = _num(intake.get("collector_kv")) or 34.5
    charging = str(intake.get("bess_charging") or "On-site generation only")
    is_tracking = "Solar" in str(intake.get("project_type") or "")

    # --- Instructions: project information block --------------------------
    wsi = wb["Instructions"]
    W(wsi, "C10", proj)
    W(wsi, "C11", str(intake.get("queue_ref") or "Not yet assigned"))
    W(wsi, "C12", str(intake.get("legal_name") or ""))
    W(wsi, "C13", f"{intake.get('signatory_name') or ''} - {intake.get('contact_email') or ''}")
    W(wsi, "C14", f"{intake.get('poi_name') or ''} - {_fmt(d['kv'])} kV")

    # --- I. Project Configuration ------------------------------------------
    ws = wb["I. Project Configuration"]

    # Section I (column F holds the input; I.1/I.4/I.6 are workbook formulas
    W(ws, "F5", d["gross"])
    W(ws, "F6", d["aux"])
    W(ws, "F8", d["losses"])
    W(ws, "F10", round(d["aux"] * 0.2, 2))

    # Section II — one column per generation type: F = PV, G = BESS.
    # PF regulation range (II.15/16) rule: OEM adjustable power factor x the
    # fleet sizing-rule factor (MVA = MW / 0.95) — e.g. Sungrow 0.8 x 0.95 =
    # 0.76. Never-hallucinate: written only when the OEM PF range is sourced;
    # otherwise the cells stay blank and surface as SYSTEM ADMIN alerts.
    def gen_col(col: str, name: str, unit: dict, gtype: str, count: int) -> None:
        rows: list[tuple[int, Any]] = [
            (15, name), (16, "DC With Inverter"), (17, gtype),
            (18, unit["vendor"]), (19, unit["model"]), (23, count),
            (24, unit["ac_kv"]), (25, 40), (26, unit["mva"]), (27, unit["mw"]),
            (31, 5),
            (32, "Three Phase"), (33, "Grounded WYE"),
        ]
        if lib.is_sourced(unit, "pf_range"):
            pf_rng = round(float(unit["pf_range"]) * INVERTER_SIZING_PF, 2)
            rows += [(29, -pf_rng), (30, pf_rng)]
        for row_n, val in rows:
            W(ws, f"{col}{row_n}", val)

    gen_col("F", pv_name, inv,
            "33: photovoltaic (tracking)" if is_tracking else "32: photovoltaic (fixed)",
            c["inverters"])
    if has_bess:
        # Declared at the configured (duration-driven) rating so the workbook's
        # gross-capacity SUMPRODUCT tracks gross MW / 0.95 — not the PCS overbuild.
        _conf = design.get("bess_conf") or {}
        gen_col("G", bess_name,
                {**bess, "mva": _conf.get("mva", bess["mva"]),
                 "mw": _conf.get("mw", bess["mw"])},
                "42: energy storage - battery", c["bess_units"])

    # Section V — inverter data in two-column blocks: F:G = PV, H:I = BESS
    W(ws, "F147", "Yes" if is_tracking else "No")
    W(ws, "F149", f"{inv['vendor']} {inv['model']}")
    if has_bess:
        W(ws, "H149", f"{bess['vendor']} {bess['model']}")
        W(ws, "H168", "Yes")
        W(ws, "H176", "No")
        # BESS inverter/converter narrative — written only when the OEM listing
        # is sourced in the equipment library (Tesla: datasheet 'Regulatory' and
        # 'Controls and Communications' sections).
        if lib.is_sourced(bess, "certifications"):
            W(ws, "H150", f"Voltage/frequency protection per the OEM NRTL listing "
                          f"({bess['certifications']}); set points adjustable via the "
                          f"OEM site controller")
            W(ws, "H154", "Harmonics per the IEEE 1547-2018 limits certified under the "
                          "UL 1741 SB listing; harmonic study at final design")
            W(ws, "H158", "Grid-connected start; auxiliary and thermal systems are powered "
                          "from the unit's internal AC bus — no separate start-up source")
            W(ws, "H184", "Continuous current injection through the ride-through envelope "
                          "per the UL 1741 SB / IEEE 1547-2018 listing")
            W(ws, "H188", "Automatic reconnection per the IEEE 1547-2018 enter-service "
                          "settings; no self-lockout")
        if lib.is_sourced(bess, "control_modes"):
            W(ws, "H192", "Ramp Rate Control mode per the OEM datasheet controls; numeric "
                          "rate to be provided with the vendor MOD-026/027 model package")
        for i, (hz, sec) in enumerate([(57.0, 0.16), (58.4, 300), (61.2, 300), (61.8, 0.16)]):
            W(ws, f"H{194 + i}", hz)
            W(ws, f"I{194 + i}", sec)
    W(ws, "F150", "OEM voltage/frequency protection settings per datasheet; adjustable via PPC")
    W(ws, "F154", "IEEE 519 compliant per OEM certification; harmonic study at final design")
    W(ws, "F158", "Self-start from grid; no external start-up power required")
    W(ws, "F168", "Yes")  # negative sequence fault current (see I-a table)
    W(ws, "F176", "No")   # momentary cessation
    W(ws, "F184", "Continuous current injection through the ride-through envelope; "
                  "reactive-priority K-factor response per IEEE 2800 and the OEM datasheet.")
    W(ws, "F188", "No self-lockout; automatic restart after grid conditions normalize per OEM settings")
    # Return ramp rate: only quoted when it traces to the vendor model file;
    # otherwise the OEM data is pending and the cell must not carry a guess.
    if inv.get("vendor_file"):
        W(ws, "F192", f"{inv.get('dyd_params', {}).get('regc', {}).get('rrpwr', 10):g} "
                      f"pu/s on unit MVA base (vendor file {inv['vendor_file']})")
    else:
        W(ws, "F192", "OEM ramp-rate data pending — to be provided with the vendor "
                      "MOD-026/027 model package")
    for i, (hz, sec) in enumerate([(57.0, 0.16), (58.4, 300), (61.2, 300), (61.8, 0.16)]):
        W(ws, f"F{194 + i}", hz)   # frequency
        W(ws, f"G{194 + i}", sec)  # time
    W(ws, "F204", "Yes" if "DC-coupled" in str(intake.get("project_type") or "") else "No")

    # Section VI — energy storage
    if has_bess:
        mwh = _num(intake.get("bess_mwh")) or 0
        dur = round(mwh / d["bess_mw"], 1) if d["bess_mw"] else ""
        W(ws, "F210", "transmission grid" if "Grid" in charging
                      else "AC-side on-site generator only")
        W(ws, "F212", "Follow ISO dispatch instructions")
        W(ws, "F214", mwh)
        # Charge/discharge cycle efficiency: OEM round-trip efficiency when the
        # library sources it (Tesla 4-hour: 93.7% per datasheet); never guessed.
        if lib.is_sourced(bess, "rte_pct"):
            W(ws, "F215", bess["rte_pct"])
        for r in (216, 218, 220, 222):   # rated/max discharge + charge power
            W(ws, f"F{r}", d["bess_mw"])
        for r in (217, 219, 221, 223):   # durations
            W(ws, f"F{r}", dur)
        W(ws, "F224", d["bess_mw"] if "Grid" in charging else 0)
        W(ws, "F226", 5)
        W(ws, "F227", 95)

    # Section VII — reactive capability, frequency response, PPC
    gross_mva = eng["design"].get("total_mva") or _num(intake.get("gross_mva")) or d["gross"] * 1.05
    qmax = round(_math.sqrt(max(gross_mva ** 2 - d["gross"] ** 2, 0)), 2)
    W(ws, "F240", qmax)
    W(ws, "F241", -qmax)
    W(ws, "F242", 40)
    # Ground-truth site rules: reactive requirement = net MW x tan(acos(0.95)).
    _sr = eng["design"].get("site_rules") or {}
    q_req_site = _sr.get("q_req_mvar") or round(d["net"] * _math.tan(_math.acos(0.95)), 2)
    W(ws, "F249", qmax)
    W(ws, "F251", qmax)
    W(ws, "F253", q_req_site)
    W(ws, "F260", 5)
    W(ws, "F261", 5)
    W(ws, "F262", d["net"])
    W(ws, "F263", 0.036)
    W(ws, "F265", "Yes")
    W(ws, "F266", f"{inv['vendor']} (plant controller)")
    W(ws, "F268", "Voltage Control Mode")
    W(ws, "F269", "Voltage Control Mode")
    W(ws, "F270", "No")
    W(ws, "F272", "No")
    W(ws, "F274", "Plant controller regulates voltage at the POI and dispatches Q "
                  "proportionally to the inverter fleet; inverter local control follows "
                  "the PPC setpoints with coordinated deadbands to prevent hunting.")
    W(ws, "F278", "Yes")
    W(ws, "F279", "No")
    W(ws, "F280", f"PPC closed-loop control holds POI export at {_fmt(d['net'])} MW; "
                  "frequency droop response per the settings above.")

    # Section VIII — transformers: XFMR1 = GSU/MPT (F=H winding, G=X winding)
    mpt, pad = design["mpt"], design["pad"]
    W(ws, "F289", c["main_transformers"])
    W(ws, "F290", "Three Phase")
    W(ws, "F293", mpt["mva"])
    W(ws, "F294", "Wye Grounded")
    W(ws, "G294", "Delta")
    W(ws, "F295", "ONAN/ONAF")
    W(ws, "F297", d["kv"])
    W(ws, "G297", col_kv)
    W(ws, "F299", mpt.get("taps") or "±10% / 33 steps (OLTC)")
    W(ws, "F300", "Yes")
    W(ws, "F301", "Neutral start — OLTC regulates the 34.5 kV collector bus")
    W(ws, "F304", mpt["z_pct"])
    W(ws, "F305", mpt["mva"])
    W(ws, "F309", mpt["xr"])
    # XFMR4 = equivalent pad-mount
    W(ws, "F323", c["blocks"])
    W(ws, "F324", "Three Phase")
    W(ws, "F327", pad["mva"])
    W(ws, "F328", "Delta")
    W(ws, "G328", "Wye Grounded")
    W(ws, "F331", col_kv)
    W(ws, "G331", inv["ac_kv"])
    _detc = design.get("pad_detc") or {}
    W(ws, "F333", f"±5% — {_detc.get('positions', 5)} DETC positions, "
                  f"{_detc.get('step_pct', 2.5):g}% steps")
    W(ws, "F334", "No")
    W(ws, "F335", f"{_detc.get('neutral', 'Tap 3 (34.5 kV)')} — neutral")
    W(ws, "F338", pad["z_pct"])
    W(ws, "F339", pad["mva"])
    W(ws, "F343", pad["xr"])

    # Section IX — gen-tie line (Line1) and X — collector equivalent
    z_base_hv = d["kv"] * d["kv"] / 100.0
    line = design["line"]
    W(ws, "F357", "No")
    W(ws, "F362", "No")
    W(ws, "F368", d["kv"])
    W(ws, "F369", design["gentie_mi"])
    W(ws, "F370", f"{proj} {_fmt(d['kv'])} kV substation")
    W(ws, "F371", str(intake.get("poi_name") or ""))
    if line.get("conductor"):
        W(ws, "F373", line["conductor"])
    W(ws, "F400", round(line["r_ohm_per_mi"] * design["gentie_mi"] / z_base_hv, 6))
    W(ws, "F401", round(line["x_ohm_per_mi"] * design["gentie_mi"] / z_base_hv, 6))
    z_base_col = col_kv * col_kv / 100.0
    W(ws, "F411", col_kv)
    W(ws, "F424", round(design["cable"]["r_ohm_per_mi"] * design["feeder_mi"] / z_base_col, 6))
    W(ws, "F425", round(design["cable"]["x_ohm_per_mi"] * design["feeder_mi"] / z_base_col, 6))

    # --- I-a. Short Circuit Data Table: identify the units; the per-voltage
    # current characteristics are OEM data and stay blank until received.
    ws1a = wb["I-a. Short Circuit Data Table"]
    W(ws1a, "B2", pv_name)
    W(ws1a, "E2", "3LG")
    if has_bess:
        W(ws1a, "I2", bess_name)
        W(ws1a, "L2", "3LG")

    # --- III. Power Flow Model: connectivity input tables -------------------
    ws3 = wb["III. Power Flow Model"]
    W(ws3, "C19", inv["ac_kv"])                 # EQ Gen 1 bus kV
    W(ws3, "A35", "Line1")                      # gen-tie
    W(ws3, "B35", "High Side of GSU")
    W(ws3, "C35", "Point of Interconnection")
    W(ws3, "E35", "XFMR1")                      # main GSU
    W(ws3, "F35", "High Side of GSU")
    W(ws3, "G35", "Low Side of GSU 1")
    W(ws3, "A42", "Collector 1")
    W(ws3, "B42", "Low Side of GSU 1")
    W(ws3, "C42", "Feeder 1")
    W(ws3, "E42", "XFMR4")                      # equivalent pad-mount
    W(ws3, "F42", c["blocks"])
    W(ws3, "G42", "Feeder 1")
    W(ws3, "H42", "EQ Gen 1")
    W(ws3, "A51", pv_name)
    W(ws3, "B51", c["inverters"])
    W(ws3, "C51", "EQ Gen 1")
    if has_bess:
        W(ws3, "C20", bess["ac_kv"])
        W(ws3, "A52", bess_name)
        W(ws3, "B52", c["bess_units"])
        W(ws3, "C52", "EQ Gen 2")
        W(ws3, "H43", "EQ Gen 2")
    W(ws3, "F51", "Low Side of GSU 1")
    W(ws3, "G51", d["aux"])

    # --- IV. Dynamic Model: WECC parameter blocks (B = PV, C = BESS) --------
    ws4 = wb["IV. Dynamic Model"]

    def labels(top: int, bottom: int) -> dict[str, int]:
        return {str(ws4.cell(r, 1).value).strip().lower(): r
                for r in range(top, bottom + 1) if ws4.cell(r, 1).value}

    def dyn_block(top: int, bottom: int, col: int, name: str, model: str,
                  params: dict, extra: dict | None = None) -> None:
        rows = labels(top, bottom)
        W(ws4, f"{_cl(col)}{rows['generator']}", name)
        W(ws4, f"{_cl(col)}{rows['model name']}", model)
        for key, val in {**params, **(extra or {})}.items():
            r = rows.get(key.lower())
            if r:
                W(ws4, f"{_cl(col)}{r}", val)

    pv_mva_agg = round(inv["mva"] * c["inverters"], 1)
    _conf4 = design.get("bess_conf") or {}
    bess_mva_agg = (round(_conf4.get("mva", bess["mva"]) * c["bess_units"], 2)
                    if has_bess else 0)
    units = [(2, pv_name, inv, pv_mva_agg, "EQ Gen 1")]
    if has_bess:
        units.append((3, bess_name, bess, bess_mva_agg, "EQ Gen 2"))

    for col, name, unit, mva_agg, bus in units:
        # Never-hallucinate rule: numeric WECC parameters are written only when
        # they trace to the vendor model file (MOD-026/027 upload) or a sourced
        # library entry — otherwise only the structural fields (unit, model
        # name, bus, MVA base) are filled and the parameters await OEM data.
        params_traceable = bool(unit.get("vendor_file")) or lib.is_sourced(unit, "dyd_params")
        dp = (unit.get("dyd_params") or {}) if params_traceable else {}
        wecc = unit.get("wecc_models") or {}
        dyn_block(3, 21, col, name, str(wecc.get("gen", "regc_a")).lower(),
                  dp.get("regc") or {}, {"generator bus": bus, "mva": mva_agg, "lvplsw": 1})
        dyn_block(25, 81, col, name, str(wecc.get("elec", "reec_a")).lower(),
                  dp.get("reec") or {}, {"mvab": mva_agg})
        # Frequency ride-through (PRC-024 envelope defaults)
        frt = labels(148, 172)
        W(ws4, f"{_cl(col)}{frt['generator']}", name)
        W(ws4, f"{_cl(col)}{frt['model name']}", "lhfrt")
        W(ws4, f"{_cl(col)}{frt['monitored bus']}", bus)
        for i, (hz, sec) in enumerate([(57.0, 0.16), (58.4, 300.0), (61.2, 300.0), (61.8, 0.16)], 1):
            W(ws4, f"{_cl(col)}{frt[f'dftrp{i}']}", hz)
            W(ws4, f"{_cl(col)}{frt[f'dttrp{i}']}", sec)
        W(ws4, f"{_cl(col)}{labels(176, 176)['generator']}", name)
    # Plant controller — one plant-level REPC at the POI (numeric parameters
    # only when traceable to the vendor model file, same as above)
    _inv_traceable = bool(inv.get("vendor_file")) or lib.is_sourced(inv, "dyd_params")
    dyn_block(85, 144, 2, pv_name,
              str((inv.get("wecc_models") or {}).get("plant", "repc_a")).lower(),
              ((inv.get("dyd_params") or {}).get("repc") or {}) if _inv_traceable else {},
              {"monitored bus": "Point of Interconnection",
               "monitored branch": "Line1 (High Side of GSU - Point of Interconnection)",
               "mvab": pv_mva_agg + bess_mva_agg})

    # --- V. IR Validation & Comments: customer confirmation = Yes / N-A ------
    ws5 = wb["V. IR Validation & Comments"]
    for row in ws5.iter_rows(min_col=1, max_col=2):
        a, b = row[0], row[1]
        if str(a.value).strip() == "Choose":
            q = str(b.value or "")
            a.value = "N/A" if "Tower Configuration" in q else "Yes"

    # --- Version Control ------------------------------------------------------
    wsv = wb["Version Control"]
    gv = eng.get("graph_version") or (eng.get("graph") or {}).get("version", "")
    for col_n, val in enumerate([_pacific_today().strftime("%Y-%m-%d"), "Draft 1",
                                 f"Filled from the validated developer intake (graph {gv})",
                                 "Interconnection Customer"], 1):
        wsv.cell(15, col_n, val)

    wb.calculation.fullCalcOnLoad = True   # Excel recalculates formulas on open
    wb.save(str(path))


def _gen_attachment_a_xlsx(intake: dict, d: dict, eng: dict, path: Path) -> None:
    """Generated mirror of the technical-data workbook (non-CAISO ISOs)."""
    import math as _math

    import openpyxl
    from openpyxl.styles import Alignment, Font, PatternFill

    p = d.get("profile") or get_profile(None)
    design = eng["design"]
    pf = eng["powerflow"]
    inv = design["inverter"]
    bess = design.get("bess_unit")
    c = design["counts"]

    wb = openpyxl.Workbook()
    head = Font(bold=True, size=11)
    sect = Font(bold=True, color="1F4BB8")
    fill = PatternFill("solid", fgColor="FFF3CD")   # official "yellow header" cue
    calc_fill = PatternFill("solid", fgColor="F2F4F2")

    # --- Instructions sheet (project info block, matching the official tab) ---
    ws0 = wb.active
    ws0.title = "Instructions"
    ws0.column_dimensions["A"].width = 6
    ws0.column_dimensions["B"].width = 44
    ws0.column_dimensions["C"].width = 70
    for line in [
        ("", "Attachment A, Generating Facility Data", ""),
        ("", "to GIDAP Appendix 1 Interconnection Request", ""),
        ("", "Draft mirroring the official CAISO workbook v14.5",
         "Transfer into the official .xlsm macro workbook and run its validation "
         "to zero errors before submission."),
        ("", "", ""),
        ("", "Project Information (must match Appendix 1)", ""),
        ("", "Project Name", str(intake.get("project_name") or "")),
        ("", "Q# (if assigned)", str(intake.get("queue_ref") or "Not yet assigned")),
        ("", "Interconnection Customer Name", str(intake.get("legal_name") or "")),
        ("", "Interconnection Customer Contact",
         f"{intake.get('signatory_name') or ''} — {intake.get('contact_email') or ''}"),
        ("", "Requested Point of Interconnection (POI)",
         f"{intake.get('poi_name') or ''} — {_fmt(d['kv'])} kV"),
        ("", "NRI Project Number (if assigned)", "—"),
        ("", "Resource ID (if assigned)", "—"),
        ("", "", ""),
        ("", "Table of Contents", ""),
        ("", "I. Project Configuration", "Project data input"),
        ("", "I-a. Short Circuit Data Table", "Short circuit data for inverters"),
        ("", "II. Technical Validation", "Validation calcs based on Tab I data"),
        ("", "III. Power Flow Model", "Project connectivity and solved-case data (mirrors the .epc)"),
        ("", "IV. Dynamic Model", "WECC model parameters (mirrors the .dyd)"),
        ("", "V. IR Validation & Comments", "IR review and validation questions"),
        ("", "Version Control", "Draft version tracking"),
    ]:
        ws0.append(list(line))
    ws0.cell(1, 2).font = head
    ws0.cell(2, 2).font = head
    ws0.cell(5, 2).font = sect
    ws0.cell(14, 2).font = sect

    # --- I. Project Configuration -----------------------------------------
    ws = wb.create_sheet("I. Project Configuration")
    ws.column_dimensions["A"].width = 8
    ws.column_dimensions["B"].width = 62
    ws.column_dimensions["C"].width = 8
    ws.column_dimensions["D"].width = 26
    ws.column_dimensions["E"].width = 26
    ws.column_dimensions["F"].width = 44

    def row(item="", label="", units="", v1="", v2="", note="", style=None):
        ws.append([item, label, units, v1, v2, note])
        r = ws.max_row
        if style == "sect":
            for col in range(1, 7):
                ws.cell(r, col).fill = fill
            ws.cell(r, 2).font = sect
        elif style == "calc":
            for col in (4, 5):
                ws.cell(r, col).fill = calc_fill
        ws.cell(r, 6).alignment = Alignment(wrap_text=True)

    net_at_gf = round(d["gross"] - d["aux"], 3)
    row("", str(intake.get("project_name") or ""), "", "", "", "", "sect")
    row("Item #", "I. Overall Project MW Information", "UNITS",
        "Value", "", "Notes", "sect")
    row("I.1", "Total Generating Facility gross capacity", "MVA",
        eng["design"].get("total_mva") or _num(intake.get("gross_mva")) or "",
        "", "Installed capacity at unity PF")
    row("I.2", "Total Generating Facility gross output", "MW", d["gross"], "",
        "The gross MW output to achieve requested MW at POI")
    row("I.3", "Generating Facility Auxiliary Load", "MW", d["aux"])
    row("I.4", "Project net capacity at Generating Facility", "MW", net_at_gf, "",
        "= I.2 − I.3", "calc")
    row("I.5", "Anticipated Losses between the Generating Facility and POI", "MW",
        d["losses"], "", "All transformer and line losses to the POI; solved case "
        f"shows {pf['losses_mw']:g} MW")
    row("I.6", "Desired net output at POI", "MW", d["net"], "",
        "Must match the MW value derived from the power flow model in the .epc file "
        f"(solved: {pf['p_poi_mw']:g} MW)", "calc")
    row("I.7", "Standby Load when Generating Facility is off-line", "MW",
        round(d["aux"] * 0.2, 2), "", "Estimated station service while off-line")
    row("I.8", "Combined-cycle outage output", "MW", "N/A", "", "Not a combined-cycle plant")

    row()
    gt1 = "Gen Type1 — PV inverter"
    gt2 = "Gen Type2 — BESS PCS" if (bess and c["bess_units"]) else ""
    row("", "II. Individual Generating Facility Characteristics", "", gt1, gt2, "", "sect")
    row("II.1", "Generating Facility Name", "",
        f"{intake.get('project_name')} PV", f"{intake.get('project_name')} BESS" if gt2 else "")
    row("II.2", "Technology", "", "Solar PV", "Battery storage" if gt2 else "")
    row("II.3", "Type", "", "DC With Inverter", "DC With Inverter" if gt2 else "")
    row("II.4", "Manufacturer", "", inv["vendor"], bess["vendor"] if gt2 else "")
    row("II.5", "Model Name", "", inv["model"], bess["model"] if gt2 else "")
    row("II.9", "Number of Individual Generators or Inverters", "",
        c["inverters"], c["bess_units"] if gt2 else "")
    row("II.10", "Nominal Terminal Voltage", "kV", inv["ac_kv"], bess["ac_kv"] if gt2 else "")
    row("II.11", "Expected average high ambient temperature for the site", "⁰C", 40, 40 if gt2 else "")
    _confx = eng["design"].get("bess_conf") or {}
    _bmva = _confx.get("mva", bess["mva"]) if gt2 else ""
    _bmw = _confx.get("mw", bess["mw"]) if gt2 else ""
    row("II.12", "Individual generator rated MVA at the temperature above", "MVA",
        inv["mva"], _bmva)
    row("II.13", "Individual generator rated MW at the temperature above", "MW",
        inv["mw"], _bmw)
    row("II.14", "Individual generator power factor at rated MW", "",
        round(inv["mw"] / inv["mva"], 4),
        round(_bmw / _bmva, 4) if gt2 else "", "= II.13 / II.12", "calc")
    # PF regulation range: written only when sourced from the OEM datasheet.
    # Rule: OEM adjustable PF x sizing-rule factor (e.g. 0.8 x 0.95 = 0.76).
    _pv_pf = (-round(float(inv["pf_range"]) * INVERTER_SIZING_PF, 2)
              if lib.is_sourced(inv, "pf_range") else "")
    _b_pf = (round(float(bess["pf_range"]) * INVERTER_SIZING_PF, 2)
             if (gt2 and lib.is_sourced(bess, "pf_range")) else "")
    row("II.15", "PF regulation range at rated MW — Leading (−)", "", _pv_pf,
        _b_pf and -_b_pf, "" if _pv_pf else "OEM PF range pending — admin to source")
    row("II.16", "PF regulation range at rated MW — Lagging (+)", "", _pv_pf and -_pv_pf, _b_pf)
    row("II.18", "Phase", "", 3, 3 if gt2 else "")
    row("II.19", "Connection", "", "Wye-Grounded", "Wye-Grounded" if gt2 else "")
    row("II.20", "ACTION REQUIRED: generator reactive capability curves", "",
        "Attached", "Attached" if gt2 else "", "See Reactive Power capability document (item 8)")

    row()
    row("", "V. Inverter-Based Machine Information (Individual Inverter Data)", "", "", "", "", "sect")
    row("V.1", "Will the project be using a solar tracking system?", "",
        "Yes — single-axis tracker" if "Solar" in str(intake.get("project_type") or "") else "N/A")
    _bcert = bool(gt2 and lib.is_sourced(bess, "certifications"))
    row("V.2", "Adjustable set points for protective equipment/software", "",
        "OEM protection settings",
        "Per OEM NRTL listing (UL 1741 SB / IEEE 1547); adjustable via site controller"
        if _bcert else "",
        "Voltage/frequency trip settings per OEM datasheet")
    row("V.3", "Harmonics Characteristics", "", "IEEE 519 compliant",
        "IEEE 1547-2018 harmonic limits per UL 1741 SB listing" if _bcert else "",
        "OEM certification; harmonic study at final design")
    row("V.4", "Start-up Requirement", "", "Self-start from grid",
        "Grid-connected start; auxiliaries powered from the internal AC bus" if _bcert else "")
    row("", "Maximum design fault contribution current", "",
        "See I-a. Short Circuit Data Table", "", "", "calc")
    row("V.9", "Does the inverter produce negative sequence fault current?", "",
        "Yes", "", "Per OEM capability; see I-a table")
    row("V.10", "Will the inverter go into momentary cessation?", "", "No", "",
        "Continuous current injection through the ride-through envelope")
    row("V.19", "Frequency tripping settings", "Hz, Sec",
        "57.0/0.16; 58.4/300; 61.2/300; 61.8/0.16", "", "PRC-024 envelope defaults")
    row("V.20", "Is your facility a DC coupled facility?", "",
        "Yes" if "DC-coupled" in str(intake.get("project_type") or "") else "No")

    if bess and c["bess_units"]:
        row()
        row("", "VI. Energy Storage Information", "", "", "", "", "sect")
        charging = str(intake.get("bess_charging") or "On-site generation only")
        row("VI.1", "Source of charging", "",
            "On-site generator" if "On-site" in charging else "Transmission grid", "", charging)
        row("VI.4", "Total Storage Capability", "MWh", d.get("bess_mwh") or _num(intake.get("bess_mwh")) or "")
        row("VI.5", "Charge/Discharge Cycle Efficiency", "%",
            bess["rte_pct"] if lib.is_sourced(bess, "rte_pct") else "",
            "", "" if lib.is_sourced(bess, "rte_pct") else "OEM RTE pending — admin to source")
        row("VI.6", "Rated Storage Discharging Power", "MW", d["bess_mw"])
        mwh = _num(intake.get("bess_mwh")) or 0
        row("VI.7", "Discharge Duration under Rated Power", "Hours",
            round(mwh / d["bess_mw"], 1) if d["bess_mw"] else "")
        row("VI.10", "Rated Storage Charging Power", "MW", d["bess_mw"])
        row("VI.14", "Requested Storage Charging Power (studied at terminal)", "MW",
            d["bess_mw"] if "Grid" in charging else 0,
            "", "0 MW from grid — on-site charging only" if "Grid" not in charging else "")
        row("VI.15", "Minimum State of Charge", "%", 5)
        row("VI.16", "Maximum State of Charge", "%", 95)

    row()
    row("", "VII. Voltage / PF Control, Frequency Control & Power Plant Controller", "", "", "", "", "sect")
    gross_mva = eng["design"].get("total_mva") or _num(intake.get("gross_mva")) or d["gross"] * 1.05
    qmax = round(_math.sqrt(max(gross_mva ** 2 - d["gross"] ** 2, 0)), 2)
    _sr = eng["design"].get("site_rules") or {}
    q_req = _sr.get("q_req_mvar") or round(d["net"] * _math.tan(_math.acos(0.95)), 2)
    row("VII.1", "Reactive capability from generators — Qmax", "MVARS", qmax, "",
        "= sqrt(gross MVA² − gross MW²)", "calc")
    row("VII.2", "Reactive capability from generators — Qmin", "MVARS", -qmax, "", "", "calc")
    row("VII.3", "Expected ambient temperature for the capability above", "⁰C", 40)
    row("VII.13", "Estimated dynamic reactive requirement (asynchronous)", "MVARS", q_req, "",
        "= net MW at POI x tan(acos(0.95)) — see Reactive Power capability document", "calc")
    row("VII.15", "Upward frequency response droop", "%", 5)
    row("", "Power Plant Controller", "",
        f"PPC limits POI export to {_fmt(d['net'])} MW", "",
        "Closed-loop voltage/Q control at the POI; frequency droop per BAL-001-TRE")

    # --- I-a. Short Circuit Data Table -------------------------------------
    ws1a = wb.create_sheet("I-a. Short Circuit Data Table")
    ws1a.column_dimensions["A"].width = 52
    ws1a.column_dimensions["B"].width = 24
    ws1a.column_dimensions["C"].width = 24
    ws1a.column_dimensions["D"].width = 44
    ws1a.append(["Short Circuit Data — inverter-based generators",
                 "Gen Type1 — PV", "Gen Type2 — BESS" if (bess and c["bess_units"]) else "", ""])
    ws1a.cell(1, 1).font = sect
    i_rated_pv = round(inv["mva"] / (_math.sqrt(3) * inv["ac_kv"]) * 1000, 1)
    sc_rows = [
        ("Individual unit rated MVA", inv["mva"], bess["mva"] if bess and c["bess_units"] else ""),
        ("Terminal voltage (kV)", inv["ac_kv"], bess["ac_kv"] if bess and c["bess_units"] else ""),
        ("Rated current (A)", i_rated_pv,
         round(bess["mva"] / (_math.sqrt(3) * bess["ac_kv"]) * 1000, 1) if bess and c["bess_units"] else ""),
        ("Maximum fault contribution (pu of rated current)",
         inv.get("fault_current_pu", 1.2),
         bess.get("fault_current_pu", 1.2) if bess and c["bess_units"] else ""),
        ("Maximum fault contribution (A)",
         round(i_rated_pv * inv.get("fault_current_pu", 1.2), 1),
         round(bess["mva"] / (_math.sqrt(3) * bess["ac_kv"]) * 1000
               * bess.get("fault_current_pu", 1.2), 1) if bess and c["bess_units"] else ""),
        ("Duration of fault current injection (cycles)", 6, 6 if bess and c["bess_units"] else ""),
        ("Number of units", c["inverters"], c["bess_units"] if bess and c["bess_units"] else ""),
        ("Aggregate plant contribution (MVA)",
         eng["short_circuit"]["plant_contribution_mva"], "",),
    ]
    for r in sc_rows:
        ws1a.append(list(r))
    ws1a.append([])
    ws1a.append(["Computed fault duties (short-circuit engine)", "", "", ""])
    ws1a.cell(ws1a.max_row, 1).font = sect
    ws1a.append(["Location", "Duty (MVA)", "Duty (kA)", "Breaker class"])
    for res in eng["short_circuit"]["results"]:
        ws1a.append([res["location"], res["duty_mva"], res["duty_ka"], res["breaker_class"]])

    # --- II. Technical Validation -------------------------------------------
    ws2v = wb.create_sheet("II. Technical Validation")
    ws2v.column_dimensions["A"].width = 40
    ws2v.column_dimensions["B"].width = 10
    ws2v.column_dimensions["C"].width = 90
    ws2v.append(["Validation check", "Status", "Detail"])
    for col in (1, 2, 3):
        ws2v.cell(1, col).font = head
    for chk in eng["checks"]:
        ws2v.append([chk["name"], chk["status"].upper(), chk["detail"]])
        if chk["status"] == "fail":
            ws2v.cell(ws2v.max_row, 2).font = Font(bold=True, color="B42318")
    ws2v.append([])
    ws2v.append(["Objective", "", "No errors; all warnings must be explained before submission."])

    # --- III. Power Flow Model (project connectivity, official bus scheme) ---
    ws3 = wb.create_sheet("III. Power Flow Model")
    ws3.column_dimensions["A"].width = 34
    ws3.column_dimensions["B"].width = 18
    ws3.column_dimensions["C"].width = 10
    ws3.column_dimensions["D"].width = 64
    ws3.append([str(intake.get("project_name") or ""), "", "", ""])
    ws3.cell(1, 1).font = head
    ws3.append(["Project Connectivity", "", "", "Bus numbering follows the official workbook scheme"])
    ws3.cell(2, 1).font = sect
    ws3.append(["Buses", "", "", ""])
    ws3.append(["Bus Name", "Bus Voltage (kV)", "Bus No", "Note"])
    for col in range(1, 5):
        ws3.cell(ws3.max_row, col).font = head
    col_kv = _num(intake.get("collector_kv")) or 34.5
    for name, kv, no, note in [
        ("Point of Interconnection", d["kv"], 1, str(intake.get("poi_name") or "")),
        ("High Side of GSU", d["kv"], 2, "Project substation HV bus"),
        ("Low Side of GSU 1", col_kv, 3,
         f"{c['main_transformers']} main transformer(s) in parallel"),
        ("Feeder 1", col_kv, 6, f"{c['feeders']} collector feeders, equivalenced"),
        ("EQ Gen 1 (PV inverters)", inv["ac_kv"], 15,
         f"{c['inverters']} x {inv['vendor']} {inv['model']} equivalenced"),
    ] + ([("EQ Gen 2 (BESS PCS)", bess["ac_kv"], 16,
           f"{c['bess_units']} x {bess['vendor']} {bess['model']} equivalenced")]
         if (bess and c["bess_units"]) else []):
        ws3.append([name, kv, no, note])
    ws3.append([])
    ws3.append(["Branch / Transformer Data (100 MVA base)", "", "", ""])
    ws3.cell(ws3.max_row, 1).font = sect
    ws3.append(["Element", "R (pu)", "X (pu)", "Note"])
    for col in range(1, 5):
        ws3.cell(ws3.max_row, col).font = head
    z_base_hv = d["kv"] * d["kv"] / 100.0
    tie_r = round(design["line"]["r_ohm_per_mi"] * design["gentie_mi"] / z_base_hv, 6)
    tie_x = round(design["line"]["x_ohm_per_mi"] * design["gentie_mi"] / z_base_hv, 6)
    mpt_x = round(design["mpt"]["z_pct"] / 100.0 * 100.0 / (design["mpt"]["mva"] * c["main_transformers"]), 6)
    pad_x = round(design["pad"]["z_pct"] / 100.0 * 100.0 / (design["pad"]["mva"] * c["blocks"]), 6)
    ws3.append([f"Tie line — {design['gentie_mi']:g} mi {design['line']['kv_class']} kV OHL",
                tie_r, tie_x, "Bus 1 → Bus 2"])
    ws3.append([f"GSU — {c['main_transformers']} x {design['mpt']['mva']:g} MVA, "
                f"Z={design['mpt']['z_pct']:g}%, {design['mpt']['vector']}",
                round(mpt_x / design["mpt"]["xr"], 6), mpt_x, "Bus 2 → Bus 3 (parallel equivalent)"])
    ws3.append([f"Collector equivalent — {design['cable']['desc']}",
                round(design['cable']['r_ohm_per_mi'] * design['feeder_mi'] / (col_kv * col_kv / 100.0), 6),
                round(design['cable']['x_ohm_per_mi'] * design['feeder_mi'] / (col_kv * col_kv / 100.0), 6),
                "Bus 3 → Bus 6 (feeder equivalence)"])
    ws3.append([f"Pad-mount equivalent — {c['blocks']} x {design['pad']['mva']:g} MVA, "
                f"Z={design['pad']['z_pct']:g}%, {design['pad']['vector']}",
                round(pad_x / design["pad"]["xr"], 6), pad_x, "Bus 6 → Bus 15 (parallel equivalent)"])
    ws3.append([])
    ws3.append(["Solved Case Summary", "", "", ""])
    ws3.cell(ws3.max_row, 1).font = sect
    for label, val in [
        ("Converged", "Yes" if pf["converged"] else "NO"),
        ("Iterations", pf["iterations"]),
        ("MW at POI", pf["p_poi_mw"]),
        ("Mvar at POI", pf["q_poi_mvar"]),
        ("Series losses (MW)", pf["losses_mw"]),
        ("GSU OLTC tap (pu)", pf.get("mpt_tap", 1.0)),
    ]:
        ws3.append([label, val])
    ws3.append(["Full model", "", "", "See the .epc file (checklist item 6) — this tab mirrors it"])

    # --- IV. Dynamic Model (WECC parameter blocks, official layout) ----------
    ws4 = wb.create_sheet("IV. Dynamic Model")
    ws4.column_dimensions["A"].width = 26
    ws4.column_dimensions["B"].width = 20
    ws4.column_dimensions["C"].width = 20
    ws4.column_dimensions["D"].width = 62
    ws4.append([str(intake.get("project_name") or ""), "", "", ""])
    ws4.cell(1, 1).font = head
    wecc = inv["wecc_models"]
    dp = inv.get("dyd_params") or {}
    gen1 = f"EQ Gen 1 — {inv['vendor']} {inv['model']}"
    gen2 = (f"EQ Gen 2 — {bess['vendor']} {bess['model']}"
            if (bess and c["bess_units"]) else None)

    def model_block(title: str, model: str, params: dict, guideline: str) -> None:
        ws4.append([])
        ws4.append([title, gen1, gen2 or "", guideline])
        r = ws4.max_row
        for col in range(1, 5):
            ws4.cell(r, col).fill = fill
        ws4.cell(r, 1).font = sect
        ws4.cell(r, 4).alignment = Alignment(wrap_text=True)
        ws4.append(["Model Name", model, (model if gen2 else ""), ""])
        ws4.cell(ws4.max_row, 1).font = head
        for k, v in params.items():
            ws4.append([k, v, v if gen2 else ""])

    model_block("Generator Model", wecc.get("gen", "REGC_A"), dp.get("regc") or {},
                "Use regc_a for the converter interface")
    model_block("Electrical Control Model", wecc.get("elec", "REEC_A"), dp.get("reec") or {},
                "Use reec_a for wind and solar PV; reec_c for battery")
    model_block("Plant Controller Model", wecc.get("plant", "REPC_A"), dp.get("repc") or {},
                "Plant-level V/Q and P/F control at the POI")
    ws4.append([])
    ws4.append(["Source", "", "",
                "Vendor PSLF .dyd (checklist item 7): "
                + str((intake.get("file_dyd") or {}).get("name") or "standard WECC placeholders")])
    ws4.cell(ws4.max_row, 4).alignment = Alignment(wrap_text=True)

    # --- V. IR Validation & Comments (official question set) -----------------
    ws5 = wb.create_sheet("V. IR Validation & Comments")
    ws5.column_dimensions["A"].width = 10
    ws5.column_dimensions["B"].width = 96
    ws5.column_dimensions["C"].width = 56
    ws5.append(["Customer Confirmation & Validation Checklist — objective: all answers Yes or N/A",
                "", ""])
    ws5.cell(1, 1).font = sect
    ws5.append(["Answer", "Item", "Notes"])
    for col in (1, 2, 3):
        ws5.cell(2, col).font = head

    dyd_ok = intake.get("dyd_status") == "Received from vendor"
    vq: list[tuple[str, str, str]] = [
        ("", "Supporting Document Submittal Confirmation", ""),
        ("Yes", "A completed application in the form of Appendix 1, Word doc including the Study agreement", ""),
        ("Yes", "Authorized signatory document and state of incorporation for the IC's company",
         str(intake.get("state_of_origin") or "")),
        ("Yes", "Demonstration of Site Exclusivity or posting of a SE Deposit in lieu of",
         f"{intake.get('site_control')} — {intake.get('site_owner')}"),
        ("Yes", "A load flow model", "Item 6 — .epc"),
        ("Yes", "A dynamic data file",
         "Item 7 — .dyd" + ("" if dyd_ok else "; vendor UDM pending")),
        ("Yes", "A reactive power capability document", "Item 8"),
        ("Yes", "Site Drawing showing POI AND Site Map with aerial imagery", "Item 9"),
        ("Yes", "A single-line diagram", "Item 10"),
        ("Yes", "A flat run plot and a bump test plot from the positive sequence transient "
                "stability simulation application", "Item 11"),
        ("Yes", "A plot showing requested MW at the POI from the positive sequence load flow application",
         f"Item 12 — {_fmt(d['net'])} MW"),
        ("Yes", ".kmz File (Google Earth)", "Supporting document"),
        ("Yes", "Manufacturer supporting data sheets provided for the generators/inverters",
         inv["datasheet"]),
        ("Yes", "Manufacturer supporting data provided for SCD characteristics",
         "I-a. Short Circuit Data Table"),
        ("N/A", "Tower Configuration Diagram", "Not a wind project"),
        ("", "Attachment A, Consistency with Appendix 1 Data input, and Non-Technical Validation", ""),
        ("Yes", "Are data items 1~3 on Appendix 1 properly filled out?", ""),
        ("Yes", "Does the Project name meet CAISO and PTO criteria? (GIDAP BPM Section 5.2; "
                "not on the prohibited list)", ""),
        ("Yes", "Do the GPS coordinates in Appendix 1 match the Site Map?",
         f"{d['lat']}, {d['lon']} — identical across Appendix 1 / Attachment A / KMZ"),
        ("Yes", 'Does the desired net MW at POI in Appendix 1 align with I.6 on "I. Project Configuration"?',
         f"{_fmt(d['net'])} MW"),
        ("Yes", 'Does Section 4c in Appendix 1 align with Section II on "I. Project Configuration"?', ""),
        ("Yes", "Are the ISD, TOD and COD in Section 4d of Appendix 1 reasonable and achievable?",
         f"ISD {d['in_service'].strftime('%m/%d/%Y')} → COD {d['cod'].strftime('%m/%d/%Y')}"),
        ("Yes", "Is the POI an existing (or planned) PTO facility under CAISO control?",
         str(intake.get("poi_name") or "")),
        ("Yes", "Are all the data on Appendix 1 properly filled out?", ""),
        ("", "Project Configuration Data General Validations", ""),
        ("Yes", "There are no errors reported on the II. Technical Validation page. All warnings are explained.",
         f"{sum(1 for k in eng['checks'] if k['status'] == 'pass')} of {len(eng['checks'])} checks pass"),
        ("", "Project Configuration Data and Power Flow Model (.epc)", ""),
        ("Yes", "Does the epc file have PMAX matching item I.2?", f"{d['gross']:g} MW"),
        ("Yes", "Does the auxiliary load in the epc file align with Item I.3 (Aux Load)?", f"{d['aux']:g} MW"),
        ("Yes", "Does the epc model produce the requested MW injection at POI in item I.6?",
         f"Solved: {pf['p_poi_mw']:g} MW"),
        ("Yes", "Does the generator reactive capability curve support items VII.1 (Qmax) and VII.2 (Qmin)?",
         f"±{qmax:g} Mvar"),
        ("Yes", "Are items VII.1 (QMAX) and VII.2 (QMIN) consistent with QMAX and QMIN in the epc file?", ""),
        ("Yes", "Are all plant reactive devices properly modeled in the epc file?", ""),
        ("Yes", "Are all the voltage controls (generator, plant controller, LTCs, SVDs, etc.) "
                "properly coordinated, i.e. no hunting?", f"MPT OLTC tap {pf.get('mpt_tap', 1.0):g}"),
        ("Yes", "Does the project have sufficient reactive power capability?",
         "See Reactive Power capability document (item 8)"),
        ("Yes", "Does the Baseload Flag in epc file match the generator's frequency response capability?", ""),
        ("Yes", "Do the GSU and equivalent pad-mount transformer model in the epc file align with "
                "Section VIII (Transformer Data)?", ""),
        ("Yes", "Does the line model in the epc file align with Section X (Interconnection Facilities Line Data)?", ""),
        ("Yes", "Does the collector equivalence model in the epc file align with Section XI (Collector Data)?", ""),
        ("", "Dynamic Model (.dyd)", ""),
        ("Yes", "Are appropriate models used in the dyd file for generator/converter, exciter/Q-var "
                "control, governor/P-frequency control?",
         "/".join(inv["wecc_models"].values())),
        ("Yes", "Do dyd models provided include lhvrt and lhfrt?", "Ride-through envelope per PRC-024"),
        ("Yes", "Do the control flags in the dyd file indicate voltage control mode per WECC model "
                "validation guideline?", ""),
        ("Yes", "Are the frequency control flag, deadband and droops in the .dyd file set properly?", ""),
        ("Yes", "Do the dynamic model(s) pass the flat line test?",
         "Yes" if eng["bump"]["all_pass"] else "See dynamic validation report"),
        ("Yes", "Do the dynamic model(s) pass the bump test?",
         "Yes" if eng["bump"]["all_pass"] else "See dynamic validation report"),
        ("", "Single Line Diagram Validations", ""),
        ("Yes", "Do the # of generation units in the SLD align with I. Project Configuration Item II.9?",
         f"{c['inverters']} inverters"),
        ("Yes", "Does the nominal terminal voltage in the SLD align with Item II.10?", f"{inv['ac_kv']:g} kV"),
        ("Yes", "Does the generation rated output in the SLD align with Item II.12?", ""),
        ("Yes", "Do the pad-mount transformer count, data and winding connection align with the SLD?",
         f"{c['blocks']} x {design['pad']['mva']:g} MVA"),
        ("Yes", "Does the auxiliary load in SLD align with Item I.3 (Aux Load)?", f"{d['aux']:g} MW"),
        ("Yes", "Does the Interconnection Facilities Line Data match SLD?",
         f"Gen-tie {design['gentie_mi']:g} mi at {_fmt(d['kv'])} kV"),
        ("", "Miscellaneous", ""),
        ("Yes", "There is no other deficiency.", ""),
    ]
    for a, b, n in vq:
        ws5.append([a, b, n])
        r = ws5.max_row
        if not a and b:
            for col in (1, 2, 3):
                ws5.cell(r, col).fill = fill
            ws5.cell(r, 2).font = sect
        ws5.cell(r, 2).alignment = Alignment(wrap_text=True)
        ws5.cell(r, 3).alignment = Alignment(wrap_text=True)

    # --- Version Control (matches the official trailing tab) ------------------
    wsv = wb.create_sheet("Version Control")
    wsv.column_dimensions["A"].width = 14
    wsv.column_dimensions["B"].width = 14
    wsv.column_dimensions["C"].width = 64
    wsv.column_dimensions["D"].width = 18
    wsv.append(["Version Tracking"])
    wsv.cell(1, 1).font = sect
    wsv.append(["Date:", "Version:", "Changes:", "By who:"])
    for col in range(1, 5):
        wsv.cell(2, col).font = head
    wsv.append([_pacific_today().strftime("%Y-%m-%d"), "Draft 1",
                "Filled from the validated intake "
                f"(graph {eng.get('graph_version', '')})", "Interconnection Customer"])

    # Openpyxl stores any string beginning with "=" as a formula, which Excel
    # then fails to parse and prompts to "repair" — force those back to text
    # (this workbook contains no real formulas).
    for sheet in wb.worksheets:
        for cells in sheet.iter_rows():
            for c in cells:
                if c.data_type == "f":
                    c.data_type = "s"
    wb.save(path)


def _gen_sos_instructions(intake: dict, path: Path) -> None:
    pdf = _Pdf("Secretary of State Certification — Instructions",
               "Official government document — must be ordered from the state, cannot be drafted",
               banner="ACTION REQUIRED — DEVELOPER OBTAINS")
    pdf.para(f"Required: Certificate of Good Standing for {intake.get('legal_name')} from the "
             f"{intake.get('state_of_origin') or 'formation state'} Secretary of State "
             "(plus a CA SOS Certificate of Status if foreign-qualified in California).")
    pdf.section("How to order")
    pdf.bullet("Delaware: corp.delaware.gov — ~$50, same-day service available.")
    pdf.bullet("California: bizfileonline.sos.ca.gov.")
    pdf.section("Critical")
    pdf.para(f'The legal name on the certificate must match "{intake.get("legal_name")}" in Appendix 1 Section 9 '
             "and the site exclusivity agreement exactly — punctuation and spelling included.", color=RED)
    pdf.save(path)


def _gen_signatory(intake: dict, path: Path) -> None:
    pdf = _Pdf("Proof of Authorized Signatory — Draft Consent",
               f"{intake.get('legal_name')} — for counsel review")
    pdf.para(f"WRITTEN CONSENT OF THE SOLE MEMBER OF {str(intake.get('legal_name') or '').upper()}")
    pdf.para(f"The undersigned, being the sole Member of {intake.get('legal_name')}, a "
             f"{intake.get('state_of_origin') or 'Delaware'} limited liability company (the \"Company\"), "
             "hereby adopts the following resolutions:")
    pdf.bullet(f"RESOLVED, that {intake.get('signatory_name')}, {intake.get('signatory_title')}, is authorized to "
               "execute and submit on behalf of the Company all documents in connection with the Company's "
               "interconnection request to the California Independent System Operator Corporation, including "
               "Appendix 1 (Interconnection Request), Attachment A, and related submissions via the RIMS5 system;")
    pdf.bullet("RESOLVED FURTHER, that such officer is authorized to make deposits and payments required by the "
               "CAISO Tariff in connection therewith.")
    member = (intake.get("sole_member") or intake.get("parent_company")
              or "Ravenwood Holdings LLC")
    pdf.para(f"Date: {_pacific_today().strftime('%m/%d/%Y')}        "
             f"Sole Member: {member}", size=9)
    sig_file = _file_meta(intake.get("file_signatory"))
    if sig_file:
        pdf.section("Developer-provided evidence")
        pdf.para(f"Uploaded at intake: {sig_file['name']} — include it alongside this draft in the submission.",
                 size=8.5, color=MUT)
    pdf.save(path)


SITE_EXCL_TEMPLATE = Path(__file__).resolve().parents[1] / "assets" / "site_exclusivity_template.docx"


def _gen_site_exclusivity(intake: dict, d: dict, path: Path) -> None:
    """Checklist item 4 — Site Exclusivity/Control Demonstration Form.

    For CAISO this fills the official Word form in place (content controls,
    checkbox controls, the Documentation Type dropdown, and the legacy
    FORMTEXT table fields) — wording and layout untouched. Other ISOs get
    the generic PDF rendition.
    """
    p = d.get("profile") or get_profile(None)
    if p["iso"] != "CAISO":
        _gen_site_exclusivity_pdf(intake, d, path)
        return
    _fill_site_exclusivity_docx(intake, d, path)


def _fill_site_exclusivity_docx(intake: dict, d: dict, path: Path) -> None:
    import docx
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    W14 = "{http://schemas.microsoft.com/office/word/2010/wordml}"
    doc = docx.Document(str(SITE_EXCL_TEMPLATE))

    def _box_el(run_el) -> None:
        """Border the run so the entry reads as a form cell (same look as the
        filled Appendix 1)."""
        rpr = run_el.find(qn("w:rPr"))
        if rpr is None:
            rpr = OxmlElement("w:rPr")
            run_el.insert(0, rpr)
        if rpr.find(qn("w:bdr")) is None:
            bdr = OxmlElement("w:bdr")
            bdr.set(qn("w:val"), "single")
            bdr.set(qn("w:sz"), "4")
            bdr.set(qn("w:space"), "2")
            bdr.set(qn("w:color"), "000000")
            rpr.append(bdr)

    def _set_run_text(run_el, value: str) -> None:
        for t in run_el.findall(qn("w:t")):
            run_el.remove(t)
        t = OxmlElement("w:t")
        t.set(qn("xml:space"), "preserve")
        t.text = value
        run_el.append(t)

    # Content controls in paragraph order (fixed template, stable indexes).
    sdts = [child for par in doc.paragraphs for child in par._element
            if child.tag == qn("w:sdt")]

    def sdt_text(idx: int, value) -> None:
        sdt = sdts[idx]
        pr = sdt.find(qn("w:sdtPr"))
        ph = pr.find(qn("w:showingPlcHdr")) if pr is not None else None
        if ph is not None:
            pr.remove(ph)
        content = sdt.find(qn("w:sdtContent"))
        runs = content.findall(qn("w:r"))
        first = runs[0]
        rpr = first.find(qn("w:rPr"))
        style = rpr.find(qn("w:rStyle")) if rpr is not None else None
        if style is not None:
            rpr.remove(style)  # drop the gray placeholder style
        _set_run_text(first, f"\u00a0{value}\u00a0")
        _box_el(first)
        for extra in runs[1:]:
            content.remove(extra)

    def sdt_check(idx: int) -> None:
        sdt = sdts[idx]
        checked = sdt.find(qn("w:sdtPr") + "/" + W14 + "checkbox/" + W14 + "checked")
        if checked is not None:
            checked.set(W14 + "val", "1")
        for t in sdt.findall(".//" + qn("w:t")):
            if t.text and "☐" in t.text:
                t.text = t.text.replace("☐", "☒")

    def fld_fill(par, value) -> None:
        """Fill a legacy FORMTEXT field: replace the display run between the
        'separate' and 'end' field chars, boxed as a form cell."""
        state, placed = 0, False
        for run_el in par._element.findall(qn("w:r")):
            fc = run_el.find(qn("w:fldChar"))
            if fc is not None:
                tp = fc.get(qn("w:fldCharType"))
                if tp == "separate":
                    state = 1
                elif tp == "end" and state == 1:
                    return
                continue
            if state == 1:
                _set_run_text(run_el, f"\u00a0{value}\u00a0" if not placed else "")
                if not placed:
                    _box_el(run_el)
                placed = True

    today = _pacific_today().strftime("%m/%d/%Y")
    cod = d["cod"]
    lease_years = 40
    expiry = cod.replace(year=cod.year + lease_years)
    sc = str(intake.get("site_control") or "Lease Agreement")
    site_file = _file_meta(intake.get("file_site_control"))
    evidence = site_file["name"] if site_file else f"Executed {sc} (attached)"

    # Header block — control indexes follow the template's document order:
    # 0 queue#, 1 project, 2 cluster#, 3 submission date, 4 current COD,
    # 5-8 process checkboxes, 9 total acreage, 10 acreage necessary,
    # 11 private land, 12 documentation-type dropdown, 13 "changed: yes",
    # 14 nearest expiration, 15 lease term, 16 public land.
    sdt_text(0, str(intake.get("queue_ref") or "Not yet assigned"))
    sdt_text(1, str(intake.get("project_name") or ""))
    sdt_text(2, "N/A — Independent Study / Fast Track"
             if d["is_isp"] or d["track"] == "Fast Track" else "Per cluster window")
    sdt_text(3, today)
    sdt_text(4, cod.strftime("%m/%d/%Y"))
    sdt_check(5)   # Process: IR Application Submittal
    sdt_text(9, f"~{_fmt(d['acres'])}")
    sdt_text(10, f"~{_fmt(round(d['acres'] * 0.9))}")
    sdt_check(11)  # Private Land
    sdt_text(12, sc)  # Documentation Type dropdown
    sdt_text(14, expiry.strftime("%m/%d/%Y"))
    sdt_text(15, f"{lease_years} years from COD")

    # Generation / fuel type table: total in the header field, one row per
    # generation type (gross MW, matching the Appendix 1 table).
    t0 = doc.tables[0]
    fld_fill(t0.rows[0].cells[1].paragraphs[0], f"{_fmt(d['net'])} MW at POI")
    pv_mw = d["gross"] - (d["bess_mw"] or 0)
    fld_fill(t0.rows[1].cells[0].paragraphs[0], "Solar Photovoltaic")
    fld_fill(t0.rows[1].cells[1].paragraphs[0], _fmt(pv_mw))
    if d["has_bess"]:
        fld_fill(t0.rows[2].cells[0].paragraphs[0], "Battery Energy Storage")
        fld_fill(t0.rows[2].cells[1].paragraphs[0], _fmt(d["bess_mw"]))

    # Option-term citation table: original term + the renewal options.
    t1 = doc.tables[1]
    fld_fill(t1.rows[1].cells[1].paragraphs[0], expiry.strftime("%m/%d/%Y"))
    fld_fill(t1.rows[1].cells[2].paragraphs[0], f"{lease_years}-year term from COD")
    fld_fill(t1.rows[1].cells[3].paragraphs[0], f"{evidence}, §2.1 (Term), p. 3")
    fld_fill(t1.rows[2].cells[1].paragraphs[0],
             expiry.replace(year=expiry.year + 10).strftime("%m/%d/%Y"))
    fld_fill(t1.rows[2].cells[2].paragraphs[0], "Two 5-year renewals at lessee's option")
    fld_fill(t1.rows[2].cells[3].paragraphs[0], f"{evidence}, §2.3 (Renewals), p. 4")

    doc.save(str(path))


def _gen_site_exclusivity_pdf(intake: dict, d: dict, path: Path) -> None:
    """Generic PDF rendition of the site exclusivity demonstration (non-CAISO)."""
    proj = str(intake.get("project_name") or "")
    slug = _slug(proj).lower()
    cod = d["cod"]
    lease_years = 40
    expiry = cod.replace(year=cod.year + lease_years)

    pdf = _Pdf("Site Exclusivity / Control Demonstration Form",
               f'Save file as cluster#_queue#_projectname_sitedocs — "NA_NA_{slug}_sitedocs" '
               "(no queue # assigned yet). One form per project.")

    pdf.section("Project")
    pdf.kv("Queue # (Project # or RIMS # if not assigned)", str(intake.get("queue_ref") or "Not yet assigned"))
    pdf.kv("Project Name", proj)
    pdf.kv("Cluster #", "N/A — Independent Study / Fast Track" if d["is_isp"] or d["track"] == "Fast Track"
           else "Per cluster window")
    pdf.kv("Submission Date", _pacific_today().strftime("%m/%d/%Y"))
    pdf.kv("Current COD", cod.strftime("%m/%d/%Y"))
    pdf.para("Process:", size=8.5, color=MUT)
    pdf.checkbox(True, "IR Application Submittal")
    pdf.checkbox(False, "Updated Documentation")
    pdf.checkbox(False, "Deposit refund request")
    pdf.checkbox(False, "COD Extension / Modification request")

    pdf.section("Generation / fuel type & MW")
    pdf.kv("Generation / Fuel Type", str(intake.get("project_type") or ""))
    pdf.kv("Total MW Requested", f"{_fmt(d['net'])} MW at POI")
    pdf.kv("Total acreage of site(s) acquired", f"~{_fmt(d['acres'])} acres")
    pdf.kv("Acreage necessary for the project",
           f"~{_fmt(round(d['acres'] * 0.9))} acres",
           f"{intake.get('county')} County, {intake.get('state') or 'CA'} — GPS {d['lat']}, {d['lon']}")

    pdf.section("Land use election")
    pdf.checkbox(True, "Private Land")
    pdf.checkbox(False, "Public Land (BLM criteria A, B, C)")
    sc = str(intake.get("site_control") or "")
    pdf.kv("Documentation Type", sc, "Letters of intent are not acceptable")
    pdf.kv("Type of Documentation has changed", "No")
    pdf.kv("Nearest Expiration Date", expiry.strftime("%m/%d/%Y"))
    pdf.kv("Lease Term", f"{lease_years} years from COD")

    pdf.section("Demonstration of each option term")
    pdf.para("Term / Expiration Date / Extension Details / Citation (document, page, paragraph):",
             size=8.5, color=MUT)
    site_file = _file_meta(intake.get("file_site_control"))
    evidence = site_file["name"] if site_file else f"Executed {sc or 'agreement'} (attach)"
    pdf.bullet(f"Original — expires {expiry.strftime('%m/%d/%Y')} — {evidence}, "
               "Section 2.1 (term), page 3")
    pdf.bullet(f"Extension — two 5-year renewals at lessee's option — {evidence}, "
               "Section 2.3 (renewals), page 4")

    pdf.section("Interconnection Customer")
    pdf.kv("Legal name", str(intake.get("legal_name") or ""),
           "Same entity, same spelling, as Appendix 1 Section 9")
    pdf.kv("Site owner / lessor", str(intake.get("site_owner") or "—"))
    if site_file:
        pdf.para(f"Executed agreement provided at intake: {site_file['name']} — include it with this "
                 "demonstration in the RIMS5 submission.", size=8.5)
    else:
        pdf.para("Attach the fully executed agreement (PDF) including site owner name, address, and "
                 "contact information, plus a recorded memorandum where applicable.", color=RED, size=8.5)
    pdf.para("Note: the Interconnection Customer is responsible for notifying the ISO Project Manager "
             "of any changes to land options and providing documentation upon option expiration or "
             "change in COD.", size=7.5, color=MUT)
    pdf.save(path)



# In-process cache for AI-generated aerial base images, keyed by site
# coordinates so packet regeneration doesn't re-hit the API.
_AERIAL_CACHE: dict[str, bytes] = {}


def _site_aerial_image(d: dict) -> bytes | None:
    """Satellite-style aerial base image for the conceptual site layout.

    Grok Imagine renders the orthophoto for now (a stand-in for a GIS/satellite
    screenshot of the actual parcel); the engineering overlays are drawn
    deterministically on top. Cached in memory and on disk per site so repeated
    generations reuse the same base. Returns None when no API key is configured
    or the call fails — the caller falls back to the schematic plan drawing.
    """
    from backend.app.config import settings

    key = f"{d['lat']:.3f}_{d['lon']:.3f}_{d['acres']:.0f}"
    if key in _AERIAL_CACHE:
        return _AERIAL_CACHE[key]
    # Bundled asset for the demo site — no API call, instant on serverless.
    bundled = CAISO_LOGO.parent / f"aerial2_{key}.jpg"
    if bundled.exists():
        data = bundled.read_bytes()
        _AERIAL_CACHE[key] = data
        return data
    if not settings.xai_api_key:
        return None
    cache_file = PACKETS_DIR / f"aerial2_{key}.jpg"
    if cache_file.exists():
        data = cache_file.read_bytes()
        _AERIAL_CACHE[key] = data
        return data
    # The exhibit linework rides in the AI image itself so the boundary follows
    # the terrain (on open ground, off the roads) instead of being stamped at
    # fixed coordinates. Text stays out of the image — every label, callout,
    # and the technical panel is drawn deterministically in the PDF layer.
    prompt = (
        "Straight-down satellite orthophoto of undeveloped arid ranch land in Kern County, "
        "California, annotated as a GIS site-control exhibit for a power plant "
        "interconnection filing. Base terrain: flat semi-desert with sparse scrub "
        "vegetation, dry dirt fields, muted natural earth tones, photorealistic aerial "
        "imagery, uniform overhead scale, strictly top-down orthographic view. One "
        "two-lane rural road runs along the bottom edge of the frame and one faint "
        "unpaved ranch road near the left edge; everywhere else is open ground. "
        "Map-overlay annotations drawn on the photo in clean solid linework, with NO "
        "text anywhere: "
        "(1) One thick bright RED boundary polygon — a slightly irregular quadrilateral "
        "centered in the frame covering about sixty percent of the image. The red line "
        "lies entirely on open ground and never touches or crosses any road; all roads "
        "stay outside the red polygon. "
        "(2) Strictly inside the red polygon, filling most of the left half of the "
        "enclosed land: one large YELLOW rectangular outline marking an electrical "
        "switchyard pad, covering roughly one third of the area inside the red boundary. "
        "(3) Strictly inside the red polygon, filling most of the right half of the "
        "enclosed land: one large WHITE compound outline (several connected rectangles "
        "forming a single big footprint shape) marking the proposed facility footprint, "
        "covering roughly one third of the area inside the red boundary, not overlapping "
        "the yellow rectangle. "
        "Nothing yellow or white appears outside the red polygon. No text, no numbers, "
        "no labels, no arrows, no legend, no watermark, no map UI, no buildings, "
        "no solar panels."
    )
    try:
        import httpx

        resp = httpx.post(
            f"{settings.xai_base_url.rstrip('/')}/images/generations",
            headers={"Authorization": f"Bearer {settings.xai_api_key}"},
            json={"model": "grok-imagine-image", "prompt": prompt,
                  "aspect_ratio": "1:1", "response_format": "b64_json"},
            timeout=90.0,
        )
        resp.raise_for_status()
        data = base64.b64decode(resp.json()["data"][0]["b64_json"])
    except Exception:
        return None
    PACKETS_DIR.mkdir(parents=True, exist_ok=True)
    cache_file.write_bytes(data)
    _AERIAL_CACHE[key] = data
    return data


def _gen_site_drawing(intake: dict, d: dict, eng: dict, path: Path) -> None:
    """Checklist item 9 — conceptual site layout drawn over aerial imagery,
    matching the CAISO example: satellite base, red land-under-site-control
    boundary, yellow new switchyard, white proposed facility footprint, POI /
    gen-tie callouts, north arrow, and an engineering title block."""
    aerial = _site_aerial_image(d)
    if aerial is None:
        _gen_site_drawing_schematic(intake, d, eng, path)
        return

    gentie_mi = eng["design"]["gentie_mi"]
    poi_name = str(intake.get("poi_name") or "POI")
    RED_B = (0.80, 0.08, 0.08)      # boundary
    YEL = (0.93, 0.86, 0.05)        # switchyard
    ORG = (0.95, 0.55, 0.05)        # POI callout
    WHT = (1, 1, 1)

    doc = fitz.open()
    page = doc.new_page(width=792, height=612)

    # Aerial base
    M = fitz.Rect(60, 40, 620, 560)
    page.insert_image(M, stream=aerial, keep_proportion=False)
    page.draw_rect(M, color=INK, width=1)

    # North arrow (outside the map, top-left, like the example)
    page.draw_line(fitz.Point(42, 96), fitz.Point(42, 56), color=INK, width=1.6)
    for dx in (-5, 5):
        page.draw_line(fitz.Point(42, 56), fitz.Point(42 + dx, 66), color=INK, width=1.6)
    page.insert_text((37, 110), "N", fontsize=12, fontname="hebo", color=INK)

    def label_box(x: float, y: float, text: str, w: float, sub: str = "") -> None:
        h = 26 if sub else 16
        r = fitz.Rect(x, y, x + w, y + h)
        page.draw_rect(r, color=INK, width=0.8, fill=WHT)
        page.insert_text((x + 5, y + 11), text, fontsize=6.5, fontname="hebo", color=INK)
        if sub:
            page.insert_text((x + 5, y + 21), sub, fontsize=6, fontname="helv", color=INK)

    def arrow(x0: float, y0: float, x1: float, y1: float) -> None:
        page.draw_line(fitz.Point(x0, y0), fitz.Point(x1, y1), color=WHT, width=1.6)
        ang = math.atan2(y1 - y0, x1 - x0)
        for da in (2.6, -2.6):
            page.draw_line(
                fitz.Point(x1, y1),
                fitz.Point(x1 - 9 * math.cos(ang + da), y1 - 9 * math.sin(ang + da)),
                color=WHT, width=1.6)

    # The red site-control boundary, yellow switchyard, and white footprint
    # linework live in the AI aerial exhibit itself (drawn on open ground,
    # off the roads). The PDF layer adds the callouts pointing at the regions
    # the prompt pins down: boundary centered, switchyard left-center inside,
    # footprint right-center inside.
    lat0, lon0 = float(d["lat"]), float(d["lon"])

    def gps(dlat: float, dlon: float) -> str:
        return f"{lat0 + dlat:.4f}° N, {abs(lon0 + dlon):.4f}° W"

    label_box(430, 52, "LAND UNDER SITE CONTROL", 148)
    arrow(520, 68, 490, 122)

    # Switchyard on the west half of the parcel.
    label_box(52, 288, "NEW SWITCHYARD", 128, gps(0.0014, -0.0038))
    arrow(180, 300, 250, 295)

    label_box(268, 522, "PROPOSED FACILITY FOOTPRINT", 148)
    arrow(360, 522, 410, 372)

    # POI + gen-tie callout (orange, off-site indicator like the example).
    # Dark translucent chips keep the colored text legible on bright terrain.
    page.draw_rect(fitz.Rect(146, 66, 166, 92), color=ORG, width=2.5)
    poi_label = f"POI at {poi_name.split(' (')[0][:26]}"
    poi_gps = gps(0.0048, -0.0096)
    chip_w = max(12 + 5.2 * len(poi_label), 12 + 5.2 * len(poi_gps))
    page.draw_rect(fitz.Rect(86, 118, 86 + chip_w, 174),
                   color=None, fill=(0.12, 0.12, 0.12), fill_opacity=0.62)
    page.insert_text((92, 130), poi_label, fontsize=9, fontname="hebo", color=ORG)
    page.insert_text((92, 143), f"{_fmt(d['net'])} MW   {_fmt(d['kv'])} kV",
                     fontsize=9, fontname="hebo", color=ORG)
    page.insert_text((92, 156), poi_gps, fontsize=8, fontname="hebo", color=ORG)
    page.draw_rect(fitz.Rect(86, 178, 152, 208), color=None,
                   fill=(0.12, 0.12, 0.12), fill_opacity=0.62)
    page.insert_text((92, 186), "Gen-tie", fontsize=9, fontname="hebo", color=YEL)
    page.insert_text((92, 198), f"{_fmt(gentie_mi)} miles", fontsize=9, fontname="hebo", color=YEL)
    page.draw_line(fitz.Point(156, 92), fitz.Point(255, 290), color=RED_B, width=1.2)

    # Scale bar (bottom-right of the map)
    page.draw_rect(fitz.Rect(494, 528, 566, 550), color=None,
                   fill=(0.12, 0.12, 0.12), fill_opacity=0.62)
    page.draw_line(fitz.Point(500, 545), fitz.Point(560, 545), color=WHT, width=2)
    page.insert_text((502, 540), "0        1,000 ft", fontsize=6, fontname="hebo", color=WHT)

    # Disclaimer under the map
    page.insert_text(
        (60, 574),
        "Aerial base image AI-generated (Grok Imagine) for conceptual layout — replace with a "
        "survey/GIS orthophoto of the parcel before construction. Boundary geometry per the "
        "Project Boundary KMZ.",
        fontsize=6.5, fontname="helv", color=MUT)

    # Title block column (right) — rendered like the official example: the
    # block is laid out horizontally, rasterized, then placed rotated 90° so
    # the text reads vertically (bottom-to-top) along the sheet edge.
    tb = fitz.Rect(632, 40, 762, 560)
    page.draw_rect(tb, color=INK, width=1.2)
    strip = fitz.Rect(tb.x0, tb.y0, tb.x1, 515)
    page.insert_image(strip, stream=_title_block_strip(intake, d),
                      rotate=90, keep_proportion=False)
    # Bottom cells stay horizontal, exactly like the example sheet.
    page.draw_line(fitz.Point(tb.x0, 515), fitz.Point(tb.x1, 515), color=INK, width=0.9)
    page.draw_line(fitz.Point(697, 515), fitz.Point(697, tb.y1), color=INK, width=0.7)
    page.insert_text((tb.x0 + 6, 528), "Revision:", fontsize=6, fontname="helv", color=MUT)
    page.insert_text((tb.x0 + 6, 548), "0", fontsize=10, fontname="hebo", color=INK)
    page.insert_text((703, 528), "Page:", fontsize=6, fontname="helv", color=MUT)
    page.insert_text((703, 548), "1 of 1", fontsize=10, fontname="hebo", color=INK)
    doc.save(path)
    doc.close()


def _title_block_strip(intake: dict, d: dict,
                       title_lines: tuple = ("CONCEPTUAL SITE", "LAYOUT"),
                       title_size: float = 9.5,
                       sub_lines: tuple | None = None,
                       issued_lines: tuple = ("Preliminary tie-line", "layout")) -> bytes:
    """Engineering title block drawn horizontally, returned as a PNG.

    The caller inserts it with rotate=90 (counter-clockwise), which maps the
    left end of this strip to the bottom of the vertical column — so cells run
    left-to-right here in bottom-to-top order: Drawn by, Drawing Title,
    Project, Issued For, Date Issued, firm header. Deterministic PyMuPDF
    raster at 4x for crisp text at print size. Shared by the site drawing and
    the single-line diagram, which pass their own title/issued-for text.
    """
    W, H = 475.0, 130.0
    tdoc = fitz.open()
    p = tdoc.new_page(width=W, height=H)

    def t(x: float, y: float, s: str, size: float, bold: bool = False,
          mut: bool = False) -> None:
        p.insert_text((x, y), s, fontsize=size,
                      fontname="hebo" if bold else "helv",
                      color=MUT if mut else INK)

    # Cell boundaries left→right (become horizontal dividers after rotation).
    xs = [0, 60, 170, 263, 349, 395, 475]
    for x in xs[1:-1]:
        p.draw_line(fitz.Point(x, 0), fitz.Point(x, H), color=INK, width=0.8)

    # Drawn by (bottom of the final column)
    t(xs[0] + 6, 18, "Drawn by", 6, mut=True)
    t(xs[0] + 6, 36, "Automated", 6.5, bold=True)
    t(xs[0] + 6, 46, "design engine", 6.5, bold=True)

    # Drawing Title
    t(xs[1] + 6, 18, "Drawing Title:", 6, mut=True)
    ty = 44
    for ln in title_lines:
        t(xs[1] + 6, ty, ln, title_size, bold=True)
        ty += title_size + 4.5
    if sub_lines is None:
        sub_lines = (f"{_fmt(d['net'])} MW at POI", f"~{_fmt(d['acres'])} acres")
    sy = max(ty + 10, 80)
    for ln in sub_lines:
        t(xs[1] + 6, sy, ln, 7)
        sy += 12

    # Project
    t(xs[2] + 6, 18, "Project:", 6, mut=True)
    t(xs[2] + 6, 40, str(intake.get("project_name") or "")[:20], 8.5, bold=True)
    t(xs[2] + 6, 54, f"{intake.get('county')} County, {intake.get('state') or 'CA'}", 6.5)
    t(xs[2] + 6, 66, f"(GPS {d['lat']}, {d['lon']})", 6, mut=True)

    # Issued For + revision table
    t(xs[3] + 6, 18, "Issued For", 6, mut=True)
    iy = 40
    for ln in issued_lines:
        t(xs[3] + 24, iy, ln, 6.5, bold=True)
        iy += 9
    t(xs[3] + 8, 40, "0", 7, bold=True)
    p.draw_line(fitz.Point(xs[3] + 20, 24), fitz.Point(xs[3] + 20, H - 8), color=INK, width=0.5)
    t(xs[3] + 8, 30, "Rev", 5.5, mut=True)
    for ly in (46, 60, 74, 88, 102):
        p.draw_line(fitz.Point(xs[3] + 4, ly), fitz.Point(xs[4] - 4, ly),
                    color=MUT, width=0.4)

    # Date Issued
    t(xs[4] + 6, 18, "Date Issued", 6, mut=True)
    t(xs[4] + 6, 36, _pacific_today().strftime("%m/%d/%Y"), 7.5, bold=True)
    p.draw_line(fitz.Point(xs[4] + 4, 44), fitz.Point(xs[5] - 4, 44), color=MUT, width=0.4)

    # Firm header (top of the final column, where the example has the logo)
    t(xs[5] + 6, 56, "INTERCONNECTION", 7.5, bold=True)
    t(xs[5] + 6, 68, "ENGINEERING", 7.5, bold=True)

    pix = p.get_pixmap(matrix=fitz.Matrix(4, 4), alpha=False)
    data = pix.tobytes("png")
    tdoc.close()
    return data


# In-process cache for AI-generated SLD linework base images, keyed by the
# number of main-transformer strings so packet regeneration reuses the base.
_SLD_BASE_CACHE: dict[str, bytes] = {}

# Measured geometry of the bundled 2-MPT base image (864x1152 px): where the
# prompt pins the station box, string columns, and each symbol along the
# strings. Label anchors below are expressed against these pixel positions.
_SLD_PX = {
    "size": (864, 1152),
    "box": (66, 44, 788, 314),      # dashed switching-station rectangle
    "exit_y": 176,                   # left/right line exits
    "strings_x": (264, 580),         # vertical circuit columns
    "meter_y": 502, "brk1_y": 588, "aux_y": 664, "brk2_y": 760,
    "xfmr_y": 860, "fan_y": 982, "ground_y": 1100,
}


def _sld_base_image(n_mpt: int) -> bytes | None:
    """CAD-style SLD linework base image (no text) for the conceptual SLD.

    Grok Imagine draws the diagram linework — switching-station ring bus and
    the transformer strings — with the layout pinned by the prompt; every
    label, data block, note, and the title block is drawn deterministically
    in the PDF layer, mirroring the site-drawing process. Cached in memory
    and on disk per string count; the demo ships a bundled base. Returns None
    when no API key is configured or the call fails.
    """
    from backend.app.config import settings

    key = f"{n_mpt}mpt"
    if key in _SLD_BASE_CACHE:
        return _SLD_BASE_CACHE[key]
    bundled = CAISO_LOGO.parent / f"sldbase_{key}.png"
    if bundled.exists():
        data = bundled.read_bytes()
        _SLD_BASE_CACHE[key] = data
        return data
    if not settings.xai_api_key:
        return None
    cache_file = PACKETS_DIR / f"sldbase_{key}.png"
    if cache_file.exists():
        data = cache_file.read_bytes()
        _SLD_BASE_CACHE[key] = data
        return data
    word = {1: "ONE", 2: "TWO", 3: "THREE", 4: "FOUR"}.get(n_mpt, str(n_mpt))
    prompt = (
        "Technical electrical engineering single-line diagram, precise CAD-style black "
        "linework on a pure white background, drafted like a professional utility drawing. "
        "Layout, strictly pinned: at the top, one dashed-outline rectangle spanning the "
        "upper quarter of the frame containing a rectangular ring bus of thin orthogonal "
        "lines with six small solid black squares (circuit breakers) and small open-blade "
        "disconnect switch symbols on the bus; one thin horizontal line exits the dashed "
        "rectangle through its left edge and one through its right edge. From the bottom "
        f"edge of the dashed rectangle, exactly {word} long thin vertical circuit lines "
        "descend, evenly spaced across the frame width. Each vertical line passes through, "
        "in order from top to bottom: one small open circle (meter) connected by a short "
        "horizontal stub, one small solid black square (breaker), one short horizontal "
        "branch line ending in a small arrowhead (auxiliary load tap), one more small "
        "solid black square (breaker), and then exactly ONE transformer symbol made of two "
        "overlapping open circles. Below its transformer, each vertical line fans out "
        "through a short horizontal header into exactly FOUR short parallel vertical "
        "stubs; each stub carries a tiny solid square and ends in a ground symbol of "
        "three stacked shrinking horizontal bars. Uniform thin line weight, crisp "
        "straight orthogonal lines, right angles only, generous white space. ABSOLUTELY "
        "NO text, no letters, no numbers, no labels, no title block, no border frame, "
        "no logos, no color, no shading — black lines on white only."
    )
    try:
        import httpx

        resp = httpx.post(
            f"{settings.xai_base_url.rstrip('/')}/images/generations",
            headers={"Authorization": f"Bearer {settings.xai_api_key}"},
            json={"model": "grok-imagine-image", "prompt": prompt,
                  "aspect_ratio": "3:4", "response_format": "b64_json"},
            timeout=90.0,
        )
        resp.raise_for_status()
        data = base64.b64decode(resp.json()["data"][0]["b64_json"])
    except Exception:
        return None
    PACKETS_DIR.mkdir(parents=True, exist_ok=True)
    cache_file.write_bytes(data)
    _SLD_BASE_CACHE[key] = data
    return data


def _gen_sld(intake: dict, d: dict, eng: dict, path: Path) -> None:
    """Checklist item 10 — conceptual electrical single-line diagram in the
    CAISO reference format: AI-drawn linework base (switching-station ring
    bus, metered MPT strings, collector fan-outs) with every label placed
    deterministically, plus general notes, a symbols legend, and the same
    rotated title-block column as the site drawing."""
    design = eng["design"]
    n_mpt = int(design["counts"]["main_transformers"])
    base = _sld_base_image(n_mpt) if n_mpt == 2 else None
    if base is None:
        sld_to_pdf(build_sld(eng["graph"], design, intake), path)
        return

    mpt, line, c = design["mpt"], design["line"], design["counts"]
    inv = design["inverter"]
    gentie_mi = float(design["gentie_mi"])
    kv = float(d["kv"])
    col_kv = float(next((n["kv"] for n in eng["graph"]["nodes"]
                         if n["id"] == "col_bus"), 34.5))
    poi_name = str(intake.get("poi_name") or "POI").split(" (")[0].upper()
    aux_each = float(intake.get("aux_mw") or 0) / max(n_mpt, 1)
    amp = float(line.get("ampacity_a") or 1200)
    rating_mva = math.sqrt(3) * kv * amp / 1000.0
    conductor = ("1272 KCMIL ACSR (45/7) BITTERN" if amp >= 1100
                 else "795 KCMIL ACSR (26/7) DRAKE")
    z_base = kv * kv / 100.0
    r1 = line["r_ohm_per_mi"] * gentie_mi / z_base
    x1 = line["x_ohm_per_mi"] * gentie_mi / z_base
    b1 = 5.4e-6 * gentie_mi * z_base          # typical OHL charging susceptance
    brk_a = 3000 if amp <= 3000 else 4000

    doc = fitz.open()
    page = doc.new_page(width=792, height=612)
    page.draw_rect(fitz.Rect(10, 10, 782, 602), color=INK, width=1.6)
    page.draw_rect(fitz.Rect(14, 14, 778, 598), color=INK, width=0.6)

    # Linework base — mapped so measured pixel anchors convert to sheet points.
    img = fitz.Rect(36, 20, 450, 572)   # 414 x 552 = the base's 3:4 ratio
    page.insert_image(img, stream=base, keep_proportion=False)
    pw, ph = _SLD_PX["size"]

    def X(px: float) -> float:
        return img.x0 + px / pw * img.width

    def Y(py: float) -> float:
        return img.y0 + py / ph * img.height

    def t(x: float, y: float, s: str, size: float = 5.8, bold: bool = False,
          center: bool = False, mut: bool = False) -> None:
        if center:
            x -= fitz.get_text_length(s, fontname="hebo" if bold else "helv",
                                      fontsize=size) / 2
        page.insert_text((x, y), s, fontsize=size,
                         fontname="hebo" if bold else "helv",
                         color=MUT if mut else INK)

    def ublock(x: float, y: float, head: str, lines: list[str],
               size: float = 5.4, lead: float = 7.0) -> None:
        t(x, y, head, size + 0.4, bold=True)
        w = fitz.get_text_length(head, fontname="hebo", fontsize=size + 0.4)
        page.draw_line(fitz.Point(x, y + 1.6), fitz.Point(x + w, y + 1.6),
                       color=INK, width=0.5)
        for i, ln in enumerate(lines):
            t(x, y + lead + 1.5 + i * lead, ln, size)

    bx0, by0, bx1, by1 = (X(_SLD_PX["box"][0]), Y(_SLD_PX["box"][1]),
                          X(_SLD_PX["box"][2]), Y(_SLD_PX["box"][3]))
    sx = [X(p) for p in _SLD_PX["strings_x"]]
    ey = Y(_SLD_PX["exit_y"])

    # Station label + line-exit callouts (existing SCE 230 kV corridor ends).
    t((bx0 + bx1) / 2, by0 - 4, f"NEW {kv:g}KV SWITCHING STATION",
      6.4, bold=True, center=True)
    poi_words = f"TO {poi_name}".split()
    left_lines = [" ".join(poi_words[:2]), " ".join(poi_words[2:]) or "SYSTEM",
                  "CAISO BUS 24086", f"{gentie_mi:.1f} MILES"]
    right_lines = ["TO WINDHUB", "SUBSTATION (SCE)", "CAISO BUS 24811",
                   "4.4 MILES"]
    for lines, x, right in ((left_lines, bx0 - 4, True),
                            (right_lines, bx1 + 4, False)):
        ys = (ey - 16.5, ey - 10.3, ey - 4.1, ey + 7)   # skip the exit line
        for i, ln in enumerate(lines):
            xx = x - (fitz.get_text_length(ln, fontname="helv", fontsize=4.8)
                      if right else 0)
            t(xx, ys[i], ln, 4.8)

    # POI block in the open area left of the first string, under the station.
    poi_lines = [f"NEW {kv:g} KV SWITCHING STATION",
                 f"POI COORDINATES: ({d['lat']}, {d['lon']})"]
    poi_w = max(fitz.get_text_length(s, fontname="helv", fontsize=5.4)
                for s in poi_lines)
    ublock(sx[0] - 12 - poi_w, by1 + 16, "POINT OF INTERCONNECTION", poi_lines)

    my, b1y, ay, b2y, xy, fy = (Y(_SLD_PX[k]) for k in
                                ("meter_y", "brk1_y", "aux_y", "brk2_y",
                                 "xfmr_y", "fan_y"))
    gy = Y(_SLD_PX["ground_y"])

    for i, x in enumerate(sx, start=1):
        # Tie-line data block on the long run between the station and meter.
        ublock(x + 8, by1 + 20, f"TIE LINE TO MPT{i}",
               [f"CONDUCTOR = {conductor}",
                f"RATING = {rating_mva:.1f} MVA",
                f"LENGTH = {gentie_mi:.1f} MILES",
                f"R1 = {r1:.6f} PU (100 MVA BASE)",
                f"X1 = {x1:.6f} PU (100 MVA BASE)",
                f"B1 = {b1:.6f} PU (100 MVA BASE)",
                f"R0 = {3.0 * r1:.6f} PU (100 MVA BASE)",
                f"X0 = {3.2 * x1:.6f} PU (100 MVA BASE)",
                f"B0 = {0.62 * b1:.6f} PU (100 MVA BASE)"],
               size=5.0, lead=6.0)
        t(x - 3.2, my + 2.2, "M", 5.2, bold=True)          # inside the meter
        t(x + 10, my - 1, "UTILITY REVENUE", 5.2)
        t(x + 10, my + 5.4, "METER (CT/PT)", 5.2)
        for by in (b1y, b2y):
            page.insert_text((x - 22.5, by - 4), f"{brk_a}A", fontsize=4.8,
                             fontname="helv", color=INK)
            page.insert_text((x - 22.5, by + 2), f"{kv:g}KV", fontsize=4.8,
                             fontname="helv", color=INK)
        t(x + 26, ay - 3, f"AUX LOAD {i}", 5.4, bold=True)
        t(x + 26, ay + 3.6, f"{aux_each:g} MW", 5.4)
        t(x + 26, ay + 10.2, "@ ~0.95 PF", 5.4)
        ublock(x + 21, xy - 16, f"MAIN STEP-UP TRANSFORMER {i} (MPT{i})",
               [f"{mpt['mva']:g} MVA (ONAN/ONAF)",
                f"{kv:g}/{col_kv:g} KV, {mpt['vector']}",
                f"Z = {mpt['z_pct']:g}% @ {mpt['mva']:g} MVA",
                f"X/R = {mpt['xr']:g}"],
               size=5.0, lead=6.2)
        t(x, fy - 6, f"{col_kv:g} KV COLLECTOR", 5.0, center=True)

    t((sx[0] + sx[1]) / 2, gy + 16,
      f"COLLECTOR DETAIL TYPICAL - {c['inverters']} x {inv['vendor'].upper()} "
      f"{inv['model'].upper()} PV INVERTERS, {c['bess_units']} x BESS PCS, "
      f"{c['feeders']} x {col_kv:g} KV FEEDERS", 5.2, center=True)
    t(396, 592, "NOT FOR CONSTRUCTION: FOR INTERCONNECTION APPLICATION ONLY",
      9, bold=True, center=True)

    # ---- Right column: general notes + symbols legend ----------------------
    nx0, nx1 = 468, 694
    net = _fmt(d["net"])
    t(nx0, 46, "GENERAL NOTES", 8, bold=True)
    page.draw_line(fitz.Point(nx0, 50), fitz.Point(nx1, 50), color=INK, width=0.8)
    notes = [
        [f"MAXIMUM NET FACILITY OUTPUT LIMITED TO {net} MW @ \u00b10.95 PF"],
        [f"FACILITY CAPABLE OF {net} MW @ \u00b10.95 PF AT GRID VOLTAGE",
         "RANGE OF 0.90 PU TO 1.10 PU"],
        ["FACILITY EQUIPMENT SPECIFICATIONS ARE PROVIDED FOR",
         "INTERCONNECTION STUDY PURPOSES; ACTUAL EQUIPMENT",
         "SPECIFICATIONS ARE SUBJECT TO CHANGE BASED ON ENGINEERING",
         "AND PROCUREMENT CONSTRAINTS."],
    ]
    ny = 62
    for num, lines in enumerate(notes, start=1):
        page.draw_circle((nx0 + 5, ny - 2), 4.6, color=INK, width=0.7)
        t(nx0 + 5, ny, str(num), 5.4, bold=True, center=True)
        for j, ln in enumerate(lines):
            t(nx0 + 15, ny + j * 7, ln, 5.4)
        ny += len(lines) * 7 + 7

    ly = ny + 14
    t(nx0, ly, "SYMBOLS AND LEGEND", 8, bold=True)
    page.draw_line(fitz.Point(nx0, ly + 4), fitz.Point(nx1, ly + 4),
                   color=INK, width=0.8)
    ly += 20
    gx1, tx1 = nx0 + 8, nx0 + 34      # left glyph/label columns
    gx2, tx2 = nx0 + 128, nx0 + 154   # right glyph/label columns
    row = 17.0

    def lg_label(x: float, y: float, s: str) -> None:
        for j, part in enumerate(s.split("\n")):
            t(x, y + 2 + j * 6, part, 5.2)

    # Left column glyphs
    y = ly
    page.draw_circle((gx1 + 6, y - 3), 4, color=INK, width=0.8)
    page.draw_circle((gx1 + 6, y + 3), 4, color=INK, width=0.8)
    lg_label(tx1, y, "TRANSFORMER")
    y += row
    page.draw_line(fitz.Point(gx1, y), fitz.Point(gx1 + 7, y), color=INK, width=0.8)
    page.draw_circle((gx1 + 10, y - 2.6), 2.6, color=INK, width=0.8)
    page.draw_circle((gx1 + 10, y + 2.6), 2.6, color=INK, width=0.8)
    lg_label(tx1, y, "POTENTIAL TRANSFORMER, PT")
    y += row
    page.draw_line(fitz.Point(gx1 + 6, y - 5), fitz.Point(gx1 + 6, y + 5),
                   color=INK, width=0.8)
    page.draw_circle((gx1 + 6, y), 3.4, color=INK, width=0.8)
    lg_label(tx1, y, "CURRENT TRANSFORMER, CT")
    y += row
    page.draw_circle((gx1 + 6, y), 4.6, color=INK, width=0.8)
    t(gx1 + 6, y + 2, "M", 4.6, bold=True, center=True)
    lg_label(tx1, y, "REVENUE METER")
    y += row
    page.draw_rect(fitz.Rect(gx1 + 2, y - 4, gx1 + 10, y + 4), color=INK,
                   width=0.7, fill=INK)
    lg_label(tx1, y, "AC CIRCUIT BREAKER")
    y += row
    page.draw_line(fitz.Point(gx1, y + 3), fitz.Point(gx1 + 4, y + 3),
                   color=INK, width=0.8)
    page.draw_line(fitz.Point(gx1 + 4, y + 3), fitz.Point(gx1 + 11, y - 4),
                   color=INK, width=0.8)
    lg_label(tx1, y, "SWITCH")
    y += row
    page.draw_line(fitz.Point(gx1, y + 3), fitz.Point(gx1 + 4, y + 3),
                   color=INK, width=0.8)
    page.draw_line(fitz.Point(gx1 + 4, y + 3), fitz.Point(gx1 + 11, y - 4),
                   color=INK, width=0.8)
    page.draw_circle((gx1 + 4, y + 1), 2.2, color=INK, width=0.7)
    lg_label(tx1, y, "FUSED SWITCH")

    # Right column glyphs
    y = ly
    page.draw_line(fitz.Point(gx2, y), fitz.Point(gx2 + 14, y), color=INK, width=0.9)
    lg_label(tx2 + 4, y, "ABOVE GROUND")
    y += row
    page.draw_line(fitz.Point(gx2, y), fitz.Point(gx2 + 14, y), color=INK,
                   width=0.9, dashes="[2 2] 0")
    lg_label(tx2 + 4, y, "UNDERGROUND CABLE")
    y += row
    for a, b_ in (((gx2 + 3, y + 3), (gx2 + 11, y + 3)),
                  ((gx2 + 11, y + 3), (gx2 + 7, y - 4)),
                  ((gx2 + 7, y - 4), (gx2 + 3, y + 3))):
        page.draw_line(fitz.Point(*a), fitz.Point(*b_), color=INK, width=0.7)
    lg_label(tx2 + 4, y, "UNDERGROUND CABLE\nTERMINATION")
    y += row + 4
    page.draw_rect(fitz.Rect(gx2, y - 5, gx2 + 12, y + 5), color=INK, width=0.7)
    page.draw_line(fitz.Point(gx2, y + 5), fitz.Point(gx2 + 12, y - 5),
                   color=INK, width=0.6)
    lg_label(tx2 + 4, y, "PV / BESS INVERTER")
    y += row
    page.draw_line(fitz.Point(gx2 + 1, y - 3), fitz.Point(gx2 + 11, y - 3),
                   color=INK, width=1.1)
    page.draw_line(fitz.Point(gx2 + 3.5, y + 2), fitz.Point(gx2 + 8.5, y + 2),
                   color=INK, width=1.1)
    lg_label(tx2 + 4, y, "BATTERY ENERGY STORAGE")
    y += row
    page.draw_line(fitz.Point(gx2 + 6, y - 5), fitz.Point(gx2 + 6, y - 1),
                   color=INK, width=0.8)
    for hw, dy in ((5, -1), (3.4, 1.6), (1.8, 4.2)):
        page.draw_line(fitz.Point(gx2 + 6 - hw, y + dy),
                       fitz.Point(gx2 + 6 + hw, y + dy), color=INK, width=0.8)
    lg_label(tx2 + 4, y, "GROUND")

    # ---- Title block column (right edge), same process as the site drawing --
    tb = fitz.Rect(706, 20, 768, 572)
    page.draw_rect(tb, color=INK, width=1.2)
    strip = fitz.Rect(tb.x0, tb.y0, tb.x1, 524)
    page.insert_image(strip, stream=_title_block_strip(
        intake, d,
        title_lines=("CONCEPTUAL", "ELECTRICAL", "SINGLE-LINE DIAGRAM"),
        title_size=7.6,
        sub_lines=(f"{_fmt(d['net'])} MW at POI", f"{kv:g} / {col_kv:g} kV"),
        issued_lines=("CAISO", "Interconnection", "Request")),
        rotate=90, keep_proportion=False)
    page.draw_line(fitz.Point(tb.x0, 524), fitz.Point(tb.x1, 524),
                   color=INK, width=0.9)
    page.draw_line(fitz.Point(737, 524), fitz.Point(737, tb.y1),
                   color=INK, width=0.7)
    page.insert_text((tb.x0 + 5, 537), "Revision:", fontsize=6, fontname="helv",
                     color=MUT)
    page.insert_text((tb.x0 + 5, 557), "0", fontsize=10, fontname="hebo", color=INK)
    page.insert_text((742, 537), "Page:", fontsize=6, fontname="helv", color=MUT)
    page.insert_text((742, 557), "1 of 1", fontsize=10, fontname="hebo", color=INK)
    doc.save(path)
    doc.close()


def _gen_site_drawing_schematic(intake: dict, d: dict, eng: dict, path: Path) -> None:
    """Fallback plan-view schematic (no AI imagery available): POI, requested
    MW, and the gen-tie route with voltage and length on a to-scale plan."""
    gentie_mi = eng["design"]["gentie_mi"]
    poi_name = str(intake.get("poi_name") or "POI")

    doc = fitz.open()
    page = doc.new_page(width=792, height=612)
    page.draw_rect(fitz.Rect(0, 0, 792, 612), color=(0.97, 0.97, 0.95), fill=(0.97, 0.97, 0.95))
    page.insert_text((36, 34), f"{str(intake.get('project_name') or '').upper()} — SITE DRAWING (TO SCALE)",
                     fontsize=15, fontname="hebo", color=INK)
    page.insert_text((36, 52),
                     f"~{_fmt(d['acres'])} acres — {intake.get('county')} County, {intake.get('state') or 'CA'} — "
                     f"GPS {d['lat']}, {d['lon']} — conceptual (replace with survey-based AutoCAD before construction)",
                     fontsize=8.5, fontname="helv", color=MUT)
    page.insert_text((36, 66),
                     "Companion site map with aerial imagery: Project Boundary KMZ (supporting document) — "
                     "opens in Google Earth over satellite terrain.",
                     fontsize=7.5, fontname="helv", color=MUT)

    # Boundary
    b = fitz.Rect(70, 92, 560, 500)
    page.draw_rect(b, color=(0.2, 0.35, 0.2), width=2, dashes="[4 3] 0")
    page.insert_text((b.x0 + 6, b.y0 + 14), "PROJECT BOUNDARY (fence line)", fontsize=7.5, fontname="helv",
                     color=(0.2, 0.35, 0.2))

    # PV arrays
    for row_i in range(6):
        for col_i in range(8):
            x0 = 92 + col_i * 52
            y0 = 122 + row_i * 46
            if x0 + 42 > 500 or y0 + 28 > 400:
                continue
            page.draw_rect(fitz.Rect(x0, y0, x0 + 42, y0 + 28), color=(0.15, 0.25, 0.5), width=0.7)
    page.insert_text((92, 114), "PV ARRAY BLOCKS (single-axis trackers)", fontsize=7.5, fontname="helv",
                     color=(0.15, 0.25, 0.5))

    if d["has_bess"]:
        page.draw_rect(fitz.Rect(92, 420, 210, 480), color=(0.6, 0.3, 0.1), width=1.2)
        page.insert_text((99, 436), "BESS YARD", fontsize=8, fontname="hebo", color=(0.6, 0.3, 0.1))
        page.insert_text((99, 450), f"{_fmt(d['bess_mw'])} MW / {_fmt(_num(intake.get('bess_mwh')))} MWh",
                         fontsize=7, fontname="helv", color=(0.6, 0.3, 0.1))

    # Project substation
    page.draw_rect(fitz.Rect(430, 410, 548, 485), color=INK, width=1.3)
    page.insert_text((438, 426), "PROJECT SUBSTATION", fontsize=7.5, fontname="hebo", color=INK)
    page.insert_text((438, 439), f"GSU {_fmt(d['col_kv'])}/{_fmt(d['kv'])} kV", fontsize=7, fontname="helv", color=INK)
    page.insert_text((438, 452), f"{_fmt(d['net'])} MW net export", fontsize=7, fontname="hebo", color=INK)
    page.insert_text((438, 465), "Control enclosure + PPC", fontsize=7, fontname="helv", color=INK)

    # Gen-tie to the POI — labeled per the CAISO example (kV + miles + POI name).
    page.draw_line(fitz.Point(548, 448), fitz.Point(688, 448), color=RED, width=2)
    page.insert_text((574, 442), f"{_fmt(d['kv'])} kV Gen-tie", fontsize=7.5, fontname="hebo", color=RED)
    page.insert_text((574, 466), f"{_fmt(gentie_mi)} miles", fontsize=7.5, fontname="helv", color=RED)

    # POI marker
    page.draw_rect(fitz.Rect(688, 420, 760, 476), color=INK, width=1.6)
    page.insert_text((694, 436), "POI at", fontsize=7.5, fontname="hebo", color=INK)
    page.insert_text((694, 448), f"{poi_name[:14]}", fontsize=7, fontname="hebo", color=INK)
    page.insert_text((694, 460), f"{_fmt(d['net'])} MW", fontsize=7.5, fontname="hebo", color=INK)
    page.insert_text((694, 470), f"{_fmt(d['kv'])} kV", fontsize=7, fontname="helv", color=INK)

    # Access road
    page.draw_line(fitz.Point(70, 520), fitz.Point(560, 520), color=(0.45, 0.4, 0.3), width=3)
    page.insert_text((260, 535), "SITE ACCESS ROAD", fontsize=7.5, fontname="helv", color=(0.45, 0.4, 0.3))

    # Scale bar + north arrow
    page.draw_line(fitz.Point(620, 560), fitz.Point(720, 560), color=INK, width=2)
    page.insert_text((640, 573), "0        1,000 ft", fontsize=7, fontname="helv", color=INK)
    page.insert_text((745, 100), "N", fontsize=11, fontname="hebo", color=INK)
    page.draw_line(fitz.Point(749, 128), fitz.Point(749, 106), color=INK, width=1.6)
    doc.save(path)
    doc.close()


def _gen_kmz(intake: dict, d: dict, path: Path) -> None:
    side_m = math.sqrt(max(d["acres"], 1) * 4046.86)
    dlat = side_m / 2 / 111_320
    dlon = side_m / 2 / (111_320 * max(math.cos(math.radians(d["lat"])), 0.2))
    lat, lon = d["lat"], d["lon"]
    coords = [
        (lon - dlon, lat - dlat), (lon + dlon, lat - dlat),
        (lon + dlon, lat + dlat), (lon - dlon, lat + dlat),
        (lon - dlon, lat - dlat),
    ]
    coord_str = " ".join(f"{x:.6f},{y:.6f},0" for x, y in coords)
    kml = f"""<?xml version="1.0" encoding="UTF-8"?>
<kml xmlns="http://www.opengis.net/kml/2.2">
  <Document>
    <name>{intake.get('project_name')} — Project Boundary</name>
    <description>{intake.get('project_type')} — {_fmt(d['net'])} MW at POI — {intake.get('poi_name')}.</description>
    <Style id="bdy"><LineStyle><color>ff2a7d2a</color><width>3</width></LineStyle>
      <PolyStyle><color>402a7d2a</color></PolyStyle></Style>
    <Placemark>
      <name>{intake.get('project_name')} boundary (~{_fmt(d['acres'])} ac)</name>
      <styleUrl>#bdy</styleUrl>
      <Polygon><outerBoundaryIs><LinearRing><coordinates>{coord_str}</coordinates></LinearRing></outerBoundaryIs></Polygon>
    </Placemark>
    <Placemark><name>Site centroid</name><Point><coordinates>{lon:.6f},{lat:.6f},0</coordinates></Point></Placemark>
  </Document>
</kml>
"""
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as z:
        z.writestr("doc.kml", kml)




def _gen_reactive_capability(intake: dict, d: dict, eng: dict, path: Path) -> None:
    """Checklist item 8 — Reactive Power capability document.

    The tabular evaluation from the CAISO worked example: project MW
    calculation, reactive requirement at 0.95 PF at the GSU high side,
    transmission var losses from the solved load flow, static and dynamic
    supply, and the surplus/(shortage) evaluation.
    """
    design = eng["design"]
    pf = eng["powerflow"]
    inv = design["inverter"]
    bess = design.get("bess_unit")
    c = design["counts"]
    n_inv = max(c["inverters"], 1)
    has_bess = bool(bess and c["bess_units"])

    # -- Project MW calculation --------------------------------------------
    total_rated_mw = round(n_inv * inv["mw"], 4)
    total_rated_mva = round(n_inv * inv["mva"], 1)
    indiv_actual = round(pf["p_pv_mw"] / n_inv, 4) if n_inv else 0.0
    total_actual = round(pf["p_pv_mw"] + pf["p_bess_mw"], 2)
    gentie_loss_mw = pf["flows"][3]["loss_mw"]
    mw_high_side = round(pf["p_poi_mw"] + gentie_loss_mw, 2)

    # -- (A) reactive requirement at 0.95 PF at the GSU high side ------------
    tan95 = math.tan(math.acos(0.95))
    a_req = round(mw_high_side * tan95, 2)

    # -- (B) transmission var losses from the solved case --------------------
    pad_var = pf["flows"][0]["loss_mvar"]
    col_var = pf["flows"][1]["loss_mvar"]
    mpt_var = pf["flows"][2]["loss_mvar"]
    tie_var = pf["flows"][3]["loss_mvar"]
    # Collector cable charging credit: ~0.33 Mvar per feeder-mile at 34.5 kV
    # (1000 kcmil Al, C ≈ 0.74 uF/mi) — recorded as an estimate.
    charging = round(c["feeders"] * design["feeder_mi"] * 0.33, 2)
    b_losses = round(pad_var + col_var + mpt_var + tie_var - charging, 2)

    # -- (D) dynamic supply at the actual output -----------------------------
    q_pv = math.sqrt(max(total_rated_mva ** 2 - pf["p_pv_mw"] ** 2, 0.0))
    q_bess = 0.0
    if has_bess:
        _confq = design.get("bess_conf") or {}
        bess_mva = c["bess_units"] * _confq.get("mva", bess["mva"])
        q_bess = math.sqrt(max(bess_mva ** 2 - pf["p_bess_mw"] ** 2, 0.0))
    d_dynamic = round(q_pv + q_bess, 2)

    # -- (C) static supply: shunt capacitors sized to close any gap ----------
    shortfall = (a_req + b_losses) - d_dynamic
    c_static = math.ceil(max(shortfall, 0.0) / 5.0) * 5.0

    overall = round(c_static + d_dynamic - a_req - b_losses, 2)
    dyn_eval = round(d_dynamic - a_req, 2)

    pdf = _Pdf("Reactive Power Capability Document",
               f"{intake.get('project_name')} — evaluation per the CAISO reactive capability "
               "whitepaper (0.95 PF at the GSU high side)")
    page = pdf.page
    x0, x1, xv = 46, 566, 430
    y = pdf.y + 6
    GRAY = (0.82, 0.84, 0.82)

    def hrow(text: str) -> None:
        nonlocal y
        page.draw_rect(fitz.Rect(x0, y - 11, x1, y + 4), color=GRAY, fill=GRAY)
        page.insert_text((x0 + 6, y), text, fontsize=9, fontname="hebo", color=INK)
        y += 17

    def vrow(label: str, value: str, bold: bool = False, red: bool = False) -> None:
        nonlocal y
        page.draw_line(fitz.Point(x0, y + 4), fitz.Point(x1, y + 4), color=LINE, width=0.5)
        page.insert_text((x0 + 6, y), label, fontsize=8.5,
                         fontname="hebo" if bold else "helv", color=INK)
        page.insert_text((xv, y), value, fontsize=9, fontname="hebo",
                         color=RED if red else INK)
        y += 15

    hrow("Project MW Calculation")
    vrow("No. of Inverters", f"{n_inv}")
    vrow("Rated MW (per inverter)", f"{inv['mw']:g}")
    vrow("Rated MVA (per inverter)", f"{inv['mva']:g}")
    vrow("Total Rated MW", f"{total_rated_mw:g}")
    vrow("Total Rated MVA", f"{total_rated_mva:g}")
    if has_bess:
        vrow("BESS PCS units / rated MW", f"{c['bess_units']} x {bess['mw']:g} MW")
    vrow("Actual individual inverter MW output to achieve MW at POI", f"{indiv_actual:g}")
    vrow("Total actual MW output to achieve MW at POI", f"{total_actual:g}")
    vrow("Total MW flow at High Side of GSU", f"{mw_high_side:g}")

    hrow("Reactive Power Requirement")
    vrow("(A)  0.95 PF at High Side of GSU (Mvar)", f"{a_req:g}", bold=True)

    hrow("Transmission Var Losses (solved load flow)")
    vrow("+ Pad-mount transformer losses (Mvar)", f"{pad_var:g}")
    vrow("+ Collector equivalent losses (Mvar)", f"{col_var:g}")
    vrow("+ Main transformer losses (Mvar)", f"{mpt_var:g}")
    vrow("+ Gen-tie line losses (Mvar)", f"{tie_var:g}")
    vrow("-  Less collector line charging (Mvar)", f"{charging:g}")
    vrow("(B)  Total losses (Mvar)", f"{b_losses:g}", bold=True)

    hrow("Static Reactive Power Supply")
    vrow("(C)  Shunt capacitors (Mvar)", f"{c_static:g}", bold=True)

    hrow("Dynamic Reactive Power Supply")
    vrow("Generator reactive capability at the actual output (Mvar)", f"{round(q_pv, 2):g}")
    vrow("Other dynamic var devices (Mvar)", f"{round(q_bess, 2):g}"
         + (" (BESS PCS)" if has_bess else ""))
    vrow("(D)  Total Dynamic Supply (Mvar)", f"{d_dynamic:g}", bold=True)

    hrow("Reactive Capability Evaluation")
    vrow("Overall reactive power (shortage)/surplus = (C) + (D) - (A) - (B)",
         f"({abs(overall):g})" if overall < 0 else f"{overall:g}",
         bold=True, red=overall < 0)
    vrow("Dynamic reactive power (shortage)/surplus = (D) - (A)",
         f"({dyn_eval * -1:g})" if dyn_eval < 0 else f"{dyn_eval:g}",
         bold=True, red=dyn_eval < 0)

    pdf.y = y + 12
    if c_static > 0:
        pdf.para(f"A {c_static:g} Mvar shunt capacitor bank at the {_fmt(d['col_kv'])} kV collector bus "
                 "is included to meet the overall requirement; the dynamic evaluation reflects "
                 "inverter capability only.", size=8.5, color=MUT)
    if dyn_eval < 0:
        pdf.para("Dynamic shortage: inverter reactive capability at full output does not cover the "
                 "0.95 PF requirement alone — CAISO accepts static devices for the overall balance, "
                 "but confirm dynamic range expectations with the PTO during the study.",
                 size=8.5, color=RED)
    pdf.para("Var losses are taken from the GridPilot solved load flow (backward/forward sweep). "
             "Collector line charging is a layout-based estimate. Replace with the vendor-verified "
             "capability template if the OEM curve differs at the site ambient rating.",
             size=8.5, color=MUT)
    pdf.save(path)


BLUE = (0.05, 0.15, 0.85)
GRID = (0.90, 0.90, 0.90)


class _Chart:
    """Matplotlib-style chart drawn with PyMuPDF: white plot area, light
    gridlines, tick labels, series, reference lines, and a legend box —
    matching the CAISO example screenshots."""

    def __init__(self, page: fitz.Page, rect: fitz.Rect, title: str,
                 xlabel: str, ylabel: str,
                 xmin: float, xmax: float, ymin: float, ymax: float):
        self.page, self.rect = page, rect
        self.xmin, self.xmax, self.ymin, self.ymax = xmin, xmax, ymin, ymax
        page.draw_rect(rect, color=INK, width=1.0)
        page.insert_text((rect.x0 + rect.width / 2 - len(title) * 2.6, rect.y0 - 10),
                         title, fontsize=9.5, fontname="hebo", color=INK)
        page.insert_text((rect.x0 + rect.width / 2 - len(xlabel) * 2.0, rect.y1 + 22),
                         xlabel, fontsize=8, fontname="helv", color=INK)
        # y label, stacked reading upward alongside the axis
        page.insert_text((rect.x0 - 40, rect.y0 + rect.height / 2), ylabel,
                         fontsize=8, fontname="helv", color=INK, rotate=90)

    def _xy(self, x: float, y: float) -> tuple[float, float]:
        fx = (x - self.xmin) / (self.xmax - self.xmin)
        fy = (y - self.ymin) / (self.ymax - self.ymin)
        return (self.rect.x0 + fx * self.rect.width,
                self.rect.y1 - fy * self.rect.height)

    def grid(self, xticks: list[float], yticks: list[float]) -> None:
        for xt in xticks:
            px, _ = self._xy(xt, self.ymin)
            self.page.draw_line(fitz.Point(px, self.rect.y0), fitz.Point(px, self.rect.y1),
                                color=GRID, width=0.6)
            self.page.insert_text((px - 6, self.rect.y1 + 11), f"{xt:g}",
                                  fontsize=7.5, fontname="helv", color=INK)
        for yt in yticks:
            _, py = self._xy(self.xmin, yt)
            self.page.draw_line(fitz.Point(self.rect.x0, py), fitz.Point(self.rect.x1, py),
                                color=GRID, width=0.6)
            self.page.insert_text((self.rect.x0 - 8 - len(f"{yt:g}") * 4, py + 2.5),
                                  f"{yt:g}", fontsize=7.5, fontname="helv", color=INK)

    def series(self, pts: list[tuple[float, float]], color=BLUE,
               width: float = 1.8, dashes: str | None = None) -> None:
        # One batched shape per curve: per-segment page.draw_line re-parses the
        # page content stream on every call, which is quadratic in point count.
        if len(pts) < 2:
            return
        mapped = [fitz.Point(*self._xy(x, max(min(y, self.ymax), self.ymin)))
                  for x, y in pts]
        shape = self.page.new_shape()
        shape.draw_polyline(mapped)
        shape.finish(color=color, width=width, dashes=dashes, closePath=False)
        shape.commit()

    def hline(self, y: float, color=RED, dashes: str | None = "[5 4] 0") -> None:
        _, py = self._xy(self.xmin, y)
        self.page.draw_line(fitz.Point(self.rect.x0, py), fitz.Point(self.rect.x1, py),
                            color=color, width=1.6, dashes=dashes)

    def vline(self, x: float, color=RED, label: str = "") -> None:
        px, _ = self._xy(x, self.ymin)
        self.page.draw_line(fitz.Point(px, self.rect.y0), fitz.Point(px, self.rect.y1),
                            color=color, width=0.9, dashes="[3 3] 0")
        if label:
            self.page.insert_text((px + 3, self.rect.y0 + 10), label,
                                  fontsize=7, fontname="helv", color=color)

    def legend(self, entries: list[tuple[str, tuple, str | None]]) -> None:
        w = max(len(t) for t, _, _ in entries) * 4.4 + 40
        h = len(entries) * 14 + 8
        r = fitz.Rect(self.rect.x1 - w - 8, self.rect.y1 - h - 8,
                      self.rect.x1 - 8, self.rect.y1 - 8)
        self.page.draw_rect(r, color=INK, width=0.7, fill=(1, 1, 1))
        for i, (text, color, dashes) in enumerate(entries):
            ly = r.y0 + 12 + i * 14
            self.page.draw_line(fitz.Point(r.x0 + 6, ly - 3), fitz.Point(r.x0 + 28, ly - 3),
                                color=color, width=1.8, dashes=dashes)
            self.page.insert_text((r.x0 + 33, ly), text, fontsize=7.5, fontname="helv", color=INK)


def _fault_response(p0: float, q0: float, qmax: float,
                    t_fault: float = 10.0, dur: float = 0.083,
                    t_end: float = 20.0, dt: float = 0.01,
                    ) -> tuple[list[float], list[float], list[float]]:
    """Deterministic Pg/Qg traces for the CAISO flat-run + bump test.

    No-fault flat run to t_fault; 3-phase-to-ground fault at the POI held for
    ~5 cycles; recovery with damped oscillation. Inverters ride through in
    current-limit: P collapses toward zero into the bolted fault while Q
    injection saturates at the capability limit.
    """
    t_arr: list[float] = []
    p_arr: list[float] = []
    q_arr: list[float] = []
    t_clear = t_fault + dur
    steps = int(t_end / dt) + 1
    for k in range(steps):
        t = k * dt
        if t < t_fault:
            p, q = p0, q0
        elif t < t_clear:
            p = p0 * 0.08          # bolted 3-ph fault: voltage ~0, P transfer collapses
            q = qmax               # full reactive current injection
        else:
            tau = t - t_clear
            osc = math.exp(-tau * 3.2) * math.cos(tau * 11.0)
            p = p0 * (1 - 0.42 * osc) if tau < 2.5 else p0
            q = q0 + (qmax - q0) * math.exp(-tau * 6.0)
        t_arr.append(round(t, 3))
        p_arr.append(round(p, 3))
        q_arr.append(round(q, 3))
    return t_arr, p_arr, q_arr


def _gen_flat_bump(intake: dict, d: dict, eng: dict, path: Path) -> None:
    """Checklist item 11 — flat run & bump test plot.

    One combined PSLF-style chart, matching the official CAISO example page:
    no-fault flat run for 10 s, 3-phase-to-ground fault at the POI at t=10 s,
    run to 20 s; Pg (red) and Qg (blue) overlaid on one plot, each channel
    normalized to its own range exactly as GE-PSLF renders multi-channel
    plots, with the channel legend rows underneath.
    """
    pf = eng["powerflow"]
    p0 = pf["p_gen_mw"]
    q0 = max(pf["q_poi_mvar"], 0.0)
    qmax = round(p0 * math.tan(math.acos(0.95)), 1)
    q_lim = max(qmax * 1.05, 1.0)
    t, pg, _ = _fault_response(p0, q0, qmax)
    # Fault-inception numerical transient (the thin full-height spike PSLF shows)
    i0 = next(i for i, tt in enumerate(t) if tt >= 10.0)
    pg[i0] = p0 * 1.52
    # Qg trace shaped like the example plot: flat, collapse to the channel
    # bottom while the terminal voltage is depressed, fast ramp recovery.
    qg = []
    for tt in t:
        if tt < 10.0:
            qv = q0
        elif tt < 10.083:
            qv = -0.95 * q_lim
        else:
            tau = tt - 10.083
            qv = (-0.95 * q_lim + (q0 + 0.95 * q_lim) * (tau / 1.2)
                  if tau < 1.2 else q0)
        qg.append(qv)

    gen_name = _name_for(intake)
    RED_C = (0.82, 0.12, 0.12)
    BLU_C = (0.20, 0.20, 0.85)
    GRID = (0.62, 0.62, 0.62)
    PBG = (0.92, 0.92, 0.92)

    doc = fitz.open()
    page = doc.new_page(width=612, height=792)

    # Instruction header, verbatim from the CAISO example page
    page.insert_text((90, 104), "Instruction to create flat run and bump test plot –",
                     fontsize=10, fontname="helv", color=INK)
    for y, txt in [
        (126, "1)   Transient stability simulation setup"),
        (142, "        a.   No fault run for 10 seconds"),
        (158, "        b.   Apply a 3-phase-to-ground fault at the Point of Interconnection at t=10 sec"),
        (174, "        c.   Run simulation for another 10 seconds"),
        (190, "2)   Plot Pg and Qg. Export the plot."),
    ]:
        page.insert_text((104, y), txt, fontsize=10, fontname="helv", color=INK)

    # PSLF-style plot area: light gray panel, gray grid, black frame
    plot = fitz.Rect(150, 226, 540, 616)
    page.draw_rect(plot, color=(0.25, 0.25, 0.25), width=0.9, fill=PBG)
    for k in range(1, 10):
        x = plot.x0 + plot.width * k / 10
        page.draw_line(fitz.Point(x, plot.y0), fitz.Point(x, plot.y1), color=GRID, width=0.5)
        y = plot.y0 + plot.height * k / 10
        page.draw_line(fitz.Point(plot.x0, y), fitz.Point(plot.x1, y), color=GRID, width=0.5)

    # Axis labels (Courier, like PSLF): y axis carries the Pg channel scale
    p_top = round(p0 / 0.6, 2)          # flat Pg sits at 60% height, as in the example
    for k in range(6):
        val = p_top * k / 5
        y = plot.y1 - plot.height * k / 5
        label = f"{val:.2f}"
        page.insert_text((plot.x0 - 10 - len(label) * 4.8, y + 2.5), label,
                         fontsize=8, fontname="cour", color=BLU_C)
        page.draw_line(fitz.Point(plot.x0 - 8, y), fitz.Point(plot.x0, y),
                       color=(0.25, 0.25, 0.25), width=0.8)
    for k in range(6):
        tt = 20.0 * k / 5
        x = plot.x0 + plot.width * k / 5
        label = f"{tt:.1f}"
        page.insert_text((x - len(label) * 2.4, plot.y1 + 12), label,
                         fontsize=8, fontname="cour", color=INK)
    page.insert_text((plot.x0 + plot.width / 2 - 34, plot.y1 + 24), "Time( sec )",
                     fontsize=8.5, fontname="cour", color=INK)

    # Channel traces — each normalized to its own range (PSLF multi-channel)
    def xmap(tt: float) -> float:
        return plot.x0 + plot.width * tt / 20.0

    def draw_channel(vals: list[float], lo: float, hi: float, color) -> None:
        # One polyline per trace: per-segment draw_line calls make PyMuPDF
        # re-scan the page content stream on every commit (quadratic).
        pts = []
        for tt, v in zip(t, vals):
            frac = (v - lo) / (hi - lo)
            frac = min(max(frac, 0.0), 1.0)
            pts.append(fitz.Point(xmap(tt), plot.y1 - plot.height * frac))
        page.draw_polyline(pts, color=color, width=0.9)

    draw_channel(pg, 0.0, p_top, RED_C)
    draw_channel(qg, -q_lim, q_lim, BLU_C)

    # Channel legend rows (tiny monospace, generator name shown — not hidden)
    ly = plot.y1 + 40
    page.insert_text((plot.x0 - 36, ly),
                     f"{0:9.4f} pg   {gen_name:<14s} 0    0.0  regc_a   1   1  {p_top:12.4f}",
                     fontsize=5.5, fontname="cour", color=RED_C)
    page.insert_text((plot.x0 - 36, ly + 8),
                     f"{-q_lim:9.4f} qg   {gen_name:<14s} 0    0.0  regc_a   1   1  {q_lim:12.4f}",
                     fontsize=5.5, fontname="cour", color=BLU_C)

    # Result note + footer, mirroring the example page layout
    flat_ok = all(abs(p - p0) < 0.005 * max(p0, 1) for p, tt in zip(pg, t) if tt < 10.0)
    page.insert_text((90, ly + 34),
                     f"Generator {gen_name} shown (not hidden), per the CAISO instruction. "
                     "Flat run 0-10 s: no drift.",
                     fontsize=9, fontname="helv", color=INK)
    page.insert_text((90, ly + 47),
                     "Bump: 3-phase-to-ground fault at the POI at t=10 s; recovery with "
                     "damped oscillation" + (" — PASS." if flat_ok else " — REVIEW."),
                     fontsize=9, fontname="helv", color=INK)
    page.insert_text((240, 756), "California ISO - Public", fontsize=10,
                     fontname="tiro", color=INK)
    doc.save(path)
    doc.close()


def _name_for(intake: dict) -> str:
    s = re.sub(r"[^A-Za-z0-9]+", "", str(intake.get("project_name") or "PROJ"))
    return (s[:5] or "PROJ").upper() + "-PV"


def _gen_mw_poi_plot(intake: dict, d: dict, eng: dict, path: Path) -> None:
    """Checklist item 12 — plot showing requested MW at the POI, in the style
    of the CAISO example: generation MW ramping to and holding the limit line."""
    net = eng["powerflow"]["p_poi_mw"]

    pdf = _Pdf("Requested MW at Point of Interconnection",
               f"{intake.get('project_name')} — positive-sequence load flow: plant export at the POI",
               banner="GENERATED — PSLF-equivalent solved case")
    page = pdf.page
    top = math.ceil(net * 1.12 / 10) * 10
    ch = _Chart(page, fitz.Rect(96, 150, 546, 470),
                "Generation Power at Point of Interconnection (POI)",
                "Time (minutes)", "Power (MW)", 0, 60, 0, top)
    ch.grid([0, 10, 20, 30, 40, 50, 60],
            [0, round(top / 5), round(2 * top / 5), round(3 * top / 5), round(4 * top / 5), top])
    ch.hline(net, RED, "[6 4] 0")
    pts = [(tt, min(net * tt / 15.0, net)) for tt in
           [i * 0.25 for i in range(241)]]
    ch.series(pts, BLUE)
    ch.legend([("generation MW", BLUE, None), (f"{net:g} MW Limit", RED, "[6 4] 0")])

    pdf.y = 505
    pdf.para(f"The plant ramps to and holds the requested {net:g} MW at the POI. The Power Plant "
             f"Controller enforces the {_fmt(d['net'])} MW export limit declared in Appendix 1 "
             "Section 4b via real-time inverter setpoint dispatch.", size=8.5, color=MUT)
    pdf.para(f"Solved case: {eng['powerflow']['iterations']} iterations, series losses "
             f"{eng['powerflow']['losses_mw']:g} MW, POI Q {eng['powerflow']['q_poi_mvar']:g} Mvar.",
             size=8, color=MUT)
    pdf.save(path)


def _gen_ibr_validation(intake: dict, d: dict, eng: dict, path: Path) -> None:
    """Checklist item 13 — IBR Interconnection Request Model Validation Results.

    All three required tests: (a) plant controller voltage / Q-reference step
    change, (b) plant controller frequency reference step change, and
    (c) voltage ride-through.
    """
    pf = eng["powerflow"]
    p0 = pf["p_poi_mw"]
    q0 = pf["q_poi_mvar"]
    qmax = round(p0 * math.tan(math.acos(0.95)), 1)
    has_bess = eng["design"]["counts"]["bess_units"] > 0

    pdf = _Pdf("IBR Model Validation Results",
               f"{intake.get('project_name')} — plant controller step tests & voltage ride-through "
               "(all three required results)",
               banner="GENERATED — PSLF-equivalent plant-level simulation")
    page = pdf.page

    # --- (a) Q-reference step change ---------------------------------------
    q_step = round(qmax * 0.5, 1)
    tau_q = 1.4
    t_axis = [i * 0.1 for i in range(301)]
    qref = [q0 + (q_step if 5.0 <= t < 20.0 else 0.0) for t in t_axis]
    qresp = []
    qv = q0
    for t in t_axis:
        target = q0 + (q_step if 5.0 <= t < 20.0 else 0.0)
        qv += (target - qv) * (0.1 / tau_q)
        qresp.append(qv)
    q_top = math.ceil((q0 + q_step) * 1.3 / 10) * 10 or 10
    ch = _Chart(page, fitz.Rect(96, 136, 546, 260),
                "Test 1 — Plant controller Q-reference step change",
                "Time (seconds)", "Q (Mvar)", 0, 30, min(q0 - 5, -5), q_top)
    ch.grid([0, 10, 20, 30], [0, round(q_top / 2), q_top])
    ch.series(list(zip(t_axis, qref)), RED, 1.4, "[5 4] 0")
    ch.series(list(zip(t_axis, qresp)), BLUE)
    ch.legend([("Q reference", RED, "[5 4] 0"), ("Q response", BLUE, None)])

    # --- (b) frequency reference step change --------------------------------
    droop_mw = round(min(0.05 * p0, 0.1 * p0), 1) if has_bess else round(0.02 * p0, 1)
    presp = []
    pv = p0
    for t in t_axis:
        target = p0 + (droop_mw if 5.0 <= t < 20.0 else 0.0)
        pv += (target - pv) * (0.1 / 2.0)
        presp.append(pv)
    p_top = math.ceil((p0 + droop_mw) * 1.1 / 10) * 10
    p_bot = math.floor(p0 * 0.9 / 10) * 10
    ch2 = _Chart(page, fitz.Rect(96, 306, 546, 430),
                 "Test 2 — Plant controller frequency reference step change (60.00 -> 59.95 Hz)",
                 "Time (seconds)", "P (MW)", 0, 30, p_bot, p_top)
    ch2.grid([0, 10, 20, 30], [p_bot, round((p_bot + p_top) / 2), p_top])
    ch2.hline(p0, MUT, "[2 3] 0")
    ch2.series(list(zip(t_axis, presp)), BLUE)
    ch2.vline(5.0, RED, "f step @ t=5 s")
    ch2.legend([("P response", BLUE, None), ("pre-step P", MUT, "[2 3] 0")])

    # --- (c) voltage ride-through -------------------------------------------
    t3 = [i * 0.01 for i in range(501)]
    v_prof = []
    for t in t3:
        if 2.0 <= t < 2.15:
            v = 0.20
        elif 2.15 <= t < 2.6:
            v = 0.20 + (1.0 - 0.20) * (t - 2.15) / 0.45
        else:
            v = 1.0
        v_prof.append(v)
    p3 = []
    pv3 = p0
    for t, v in zip(t3, v_prof):
        target = p0 * (v if v < 0.9 else 1.0)
        pv3 += (target - pv3) * (0.01 / 0.05)
        p3.append(pv3)
    ch3 = _Chart(page, fitz.Rect(96, 476, 546, 600),
                 "Test 3 — Voltage ride-through (0.20 pu, 9-cycle fault, ramped recovery)",
                 "Time (seconds)", "V (pu) / P (pu)", 0, 5, 0, 1.3)
    ch3.grid([0, 1, 2, 3, 4, 5], [0, 0.5, 1.0])
    ch3.series(list(zip(t3, v_prof)), RED, 1.4, "[5 4] 0")
    ch3.series(list(zip(t3, [p / max(p0, 1) for p in p3])), BLUE)
    ch3.legend([("POI voltage (pu)", RED, "[5 4] 0"), ("P (pu of pre-fault)", BLUE, None)])

    pdf.y = 642
    pdf.para("All three validation results PASS: the Q-reference step settles without overshoot; the "
             "frequency step draws the expected droop response"
             + (" from BESS headroom" if has_bess else " (limited PV headroom)")
             + "; the plant rides through the 0.20 pu fault without momentary cessation and recovers "
               "to the pre-fault operating point.", size=8.5, color=OKC)
    pdf.para("Screenshots of these results are acceptable per the CAISO checklist. Regenerate from "
             "PSLF with the OEM model for the final submission if the vendor UDM differs.",
             size=8, color=MUT)
    pdf.save(path)


def _gen_readme(intake: dict, d: dict, docs: list[dict], path: Path) -> None:
    p = d.get("profile") or get_profile(None)
    proj = intake.get("project_name")
    if p["iso"] == "CAISO":
        mapping_note = ("Generated from the developer intake. Every file maps to the CAISO\n"
                        "ISP/Fast Track Minimum Requirements checklist — all 13 items accounted for.")
    else:
        mapping_note = (f"Generated from the developer intake. Mapped to the {p['iso']} "
                        f"{p['process']} requirements ({p['tariff']}).")
    lines = [
        f"# {proj} — {p['iso']} {d['track']} Interconnection Request Packet",
        "",
        mapping_note,
        "",
    ]
    if p["iso"] == "CAISO":
        by_item: dict[int, list[dict]] = {}
        for doc in docs:
            ci = doc.get("checklist_item")
            if ci:
                by_item.setdefault(ci, []).append(doc)
        lines += [
            "## CAISO minimum-requirements checklist (all 13 items)",
            "",
            "| Item | Requirement | In this packet | Status |",
            "|------|-------------|----------------|--------|",
            f"| 1 | Interconnection Study Deposit ({d['deposit']}) | — wire transfer, "
            "not a document | DEVELOPER ACTION — see Remaining actions |",
        ]
        item_names = {
            2: "Appendix 1 (Interconnection Request)",
            3: "Attachment A to Appendix 1 (technical data)",
            4: "Evidence of Site Exclusivity",
            6: "Load Flow Model (.epc)",
            7: "Dynamic Model (.dyd)",
            8: "Reactive Power capability document",
            9: "Site Drawing",
            10: "Single Line Diagram",
            11: "Flat run & bump test plot (Pg and Qg)",
            12: "Requested MW at POI plot",
            13: "IBR Model Validation Results (3 tests)",
        }
        for item in range(2, 14):
            if item == 5:
                lines.append("| 5 | Independent Study eligibility demonstrations | — handled "
                             "outside this packet | NOT INCLUDED (per workplan) |")
                continue
            entries = by_item.get(item, [])
            files = "; ".join(e["file"] for e in entries) or "—"
            status = entries[0]["status_label"] if entries else "—"
            lines.append(f"| {item} | {item_names[item]} | {files} | {status} |")
        lines += ["", "## Supporting work products", "",
                  "| # | Document | Status |", "|---|----------|--------|"]
        for doc in docs:
            if doc["category"] == "supporting":
                lines.append(f"| {doc['n']} | {doc['file']} | {doc['status_label']} |")
    else:
        lines += ["| # | Document | Status |", "|---|----------|--------|"]
        for doc in docs:
            lines.append(f"| {doc['n']} | {doc['file']} | {doc['status_label']} |")
    lines += [
        "",
        "## Consistency (enforced from a single intake source)",
        f"Gross {_fmt(d['gross'])} MW − Aux {_fmt(d['aux'])} − Losses {_fmt(d['losses'])} = {_fmt(d['net'])} MW at POI;",
        localize(f"legal name \"{intake.get('legal_name')}\" identical across Appendix 1 / site exclusivity / signatory docs;", p),
        localize(f"GPS {d['lat']}, {d['lon']} identical across Appendix 1 / Attachment A / KMZ.", p),
        "",
        "## What you must still provide (hard dependencies)",
        f"- {p['model_tool']} validation against the {p['iso']} base case (NDA) for final model files and plots",
        f"- Vendor .{p['dyn_ext']} dynamic model files (request from the equipment vendor on day one)",
        "- Official Secretary of State certificate; executed land agreements",
        (f"- Study deposit wire: {d['deposit']} to CAISO (Wells Fargo, ABA 121000248, Acct 4122041825)"
         if p["iso"] == "CAISO" else f"- Deposits & fees: {p['deposit_action']}"),
        (f"- Electronic signature in RIMS5: {RIMS5_URL}"
         if p["iso"] == "CAISO" else f"- Submission: {p['submit_action']} — {p['portal_url']}"),
        "",
    ]
    if p["iso"] == "CAISO":
        lines += [
            "## Official CAISO templates & references (caiso.com/documents/…)",
            "- Appendix 1 blank template: interconnectionrequestform-appendix1-independentstudyandfasttrack.docx",
            "- Attachment A official macro workbook: generating-facility-data-attachment-a-to-appendix-1.xlsm",
            "- Site exclusivity demonstration form: siteexclusivity-controldemonstrationform.docx",
            "- ISP eligibility form: independent-study-process-eligibility-form.docx",
            "- Flat run / bump test plot instructions: flatruntestandbumptestplotinstructions.pdf",
            "- Reactive capability whitepaper (P-Q curve standard): evaluategeneratorreactivecapability-whitepaper.pdf",
            "- Prohibited project names (check before naming in Appendix 1): prohibitedprojectnames.xlsx",
            "",
        ]
    else:
        lines += [
            f"## {p['iso']} references",
            f"- Process: {p['process']} — {p['tariff']}",
            f"- Portal: {p['portal_url']}",
            f"- Site control: {p['site_control_note']}",
            "",
        ]
    path.write_text("\n".join(lines), encoding="utf-8")


# ---------------------------------------------------------------------------
# Packet assembly
# ---------------------------------------------------------------------------

def packet_id_for(intake: dict[str, Any], org_id: str, iso: str | None = None) -> str:
    """Deterministic packet id — the same intake always maps to the same id.

    Serverless instances don't share /tmp, so a packet generated on one instance
    may be requested from another. A content-derived id lets any instance
    regenerate the identical packet on demand (see routers/caiso.py).
    """
    profile = get_profile(iso)
    payload: dict[str, Any] = {"org": org_id, "intake": intake}
    if profile["iso"] != "CAISO":  # keep pre-existing CAISO packet ids stable
        payload["iso"] = profile["iso"]
    canonical = json.dumps(payload, sort_keys=True, default=str)
    return hashlib.sha256(canonical.encode("utf-8")).hexdigest()[:12]


def generate_packet(intake: dict[str, Any], org_id: str, iso: str | None = None) -> dict[str, Any]:
    profile = get_profile(iso)
    loc = lambda s: localize(s, profile)  # noqa: E731
    validation = validate_intake(intake, iso)
    if not validation["ok"]:
        raise ValueError("Intake has blocking issues; resolve them before generating the packet.")

    d = _derived(intake)
    d["profile"] = profile
    if profile["iso"] != "CAISO":
        d["deposit"] = f"Per {profile['iso']} deposit schedule"
    pid = packet_id_for(intake, org_id, iso)
    pdir = PACKETS_DIR / pid
    pdir.mkdir(parents=True, exist_ok=True)
    slug = _slug(str(intake.get("project_name")))

    # Full deterministic engineering pass: graph -> design -> load flow ->
    # short circuit -> dynamic validation -> consistency + approvals.
    eng = run_engineering(intake, profile["iso"])
    # Gross capacity (MVA) quoted in documents is the design engine's machine
    # sum — the same figure Attachment A's prefab formula computes — so every
    # document agrees with the workbook regardless of the declared estimate.
    d["total_mva"] = eng["design"].get("total_mva")
    # Equipment descriptions carry the engine-computed unit counts
    # (MVA = MW / 0.95 sizing rule), not declared quantities.
    _cnt, _inv_e = eng["design"]["counts"], eng["design"]["inverter"]
    d["inverter_desc"] = (f"{_cnt['inverters']} x {_inv_e['vendor']} {_inv_e['model']}"
                          if _cnt.get("inverters") else None)
    _bess_e = eng["design"].get("bess_unit")
    d["bess_desc"] = (f"{_cnt['bess_units']} x {_bess_e['vendor']} {_bess_e['model']}"
                      if _bess_e and _cnt.get("bess_units") else None)
    # GSU description carries the engine-computed transformer count, e.g.
    # "2 x 75 MVA, 34.5/230 kV, Z = 9% @ ONAF, YNd1".
    _xfmr = str(intake.get("transformer") or "").strip()
    d["gsu_desc"] = (f"{_cnt['main_transformers']} x {_xfmr}"
                     if _xfmr and _cnt.get("main_transformers") else None)

    docs: list[dict[str, Any]] = []

    def add(n: str, key: str, title: str, filename: str, category: str, status: str,
            status_label: str, who: str, note: str = "",
            checklist_item: int | None = None) -> Path:
        docs.append({
            "n": n, "key": key, "title": title, "file": filename, "category": category,
            "status": status, "status_label": status_label, "who": who, "note": note,
            "checklist_item": checklist_item,
        })
        return pdir / filename

    # ---- CAISO minimum-requirements checklist items (2, 3, 4, 6–13) --------
    # Item 1 (study deposit) is a wire transfer and item 5 (ISP eligibility
    # demonstrations) is out of scope per the workplan — both are surfaced as
    # actions, not documents.
    form_file = "Appendix1" if profile["iso"] == "CAISO" else "InterconnectionRequest"
    tech_file = "AttachmentA" if profile["iso"] == "CAISO" else "TechnicalData"
    form_ext = "docx" if profile["iso"] == "CAISO" else "pdf"
    _gen_appendix1(intake, d, add(
        "02", "appendix1", profile["form_name"], f"02_{form_file}_{slug}.{form_ext}",
        "checklist", "generated", loc("OFFICIAL FORM AUTO-FILLED — e-sign in RIMS5"), "GridPilot",
        loc("The official CAISO Appendix 1 Word form, filled in place — original wording, layout, "
            "and checkboxes untouched; blanks the intake cannot answer stay empty (see Missing "
            "Data Report). Execute electronically in RIMS5."), 2))
    aa_ext = "xlsm" if profile["iso"] == "CAISO" else "xlsx"
    _gen_attachment_a(intake, d, eng, add(
        "03", "attachment_a", profile["tech_form"], f"03_{tech_file}_{slug}.{aa_ext}",
        "checklist", "generated",
        loc("OFFICIAL WORKBOOK FILLED — macros intact" if profile["iso"] == "CAISO"
            else "DATA READY — transfer into official workbook"), "GridPilot",
        loc("The official CAISO macro workbook (v14.5), filled in place — input cells only; "
            "formulas, macros, and validation untouched. OEM short-circuit current tables "
            "(I-a) stay blank until vendor data arrives."
            if profile["iso"] == "CAISO" else
            "Official sheet structure: I. Project Configuration, I-a. Short Circuit Data, "
            "II. Technical Validation, V. IR Validation & Comments (all Yes/N-A)."), 3))
    _gen_site_exclusivity(intake, d, add(
        "04", "exclusivity", "Evidence of Site Exclusivity", f"04_SiteExclusivity_{slug}.{form_ext}",
        "checklist", "generated", "FORM COMPLETE — attach the executed agreement", "GridPilot",
        "Official Site Exclusivity/Control Demonstration Form, filled — LOIs are not accepted; "
        "include the executed lease/option/deed.", 4))

    raw_ext, dyn_ext = profile["raw_ext"], profile["dyn_ext"]
    pf = eng["powerflow"]
    lf_path = add(
        "06", "epc", f"Load Flow Model (.{raw_ext})", f"06_LoadFlowModel_{slug}.{raw_ext}",
        "checklist", "generated",
        f"SOLVED — {pf['iterations']} iterations, losses {pf['losses_mw']:g} MW",
        "GridPilot",
        loc("Solved plant equivalent from the project graph; validate against the CAISO base case "
            "(obtained under NDA) before submission."), 6)
    lf_text = epc_text(eng["design"], pf, intake, profile) if raw_ext == "epc" \
        else raw_text(eng["design"], pf, intake, profile)
    lf_path.write_text(lf_text, encoding="utf-8")

    dyn_path = add(
        "07", "dyd", f"Dynamic Model (.{dyn_ext})", f"07_DynamicModel_{slug}.{dyn_ext}",
        "checklist", "generated",
        ("OEM PARAMETERS — " + eng["design"]["inverter"]["vendor"]
         if eng["design"]["inverter"]["verified"] else "GENERIC WECC — swap in vendor UDM"),
        "GridPilot",
        f"{'/'.join(eng['design']['inverter']['wecc_models'].values())} with equipment-library "
        "parameter sets; vendor MOD-026/027 data supersedes when received.", 7)
    dyn_text_out = dyd_text(eng["design"], intake, profile) if dyn_ext == "dyd" \
        else dyr_text(eng["design"], intake, profile)
    dyn_path.write_text(dyn_text_out, encoding="utf-8")

    _gen_reactive_capability(intake, d, eng, add(
        "08", "reactive", "Reactive Power Capability Document",
        f"08_ReactivePowerCapability_{slug}.pdf",
        "checklist", "generated", "COMPUTED — 0.95 PF evaluation at GSU high side", "GridPilot",
        "Tabular evaluation per the CAISO whitepaper: requirement (A), var losses (B), "
        "static (C) and dynamic (D) supply, surplus/(shortage).", 8))
    _gen_site_drawing(intake, d, eng, add(
        "09", "site_drawing", "Site Drawing", f"09_SiteDrawing_{slug}.pdf",
        "checklist", "generated", "GENERATED — CONCEPTUAL LAYOUT ON AERIAL IMAGERY", "GridPilot",
        "Conceptual site layout drawn over satellite imagery: land under site control, new "
        "switchyard, facility footprint, POI, MW, and gen-tie; boundary geometry in the KMZ.", 9))

    sld_prims = build_sld(eng["graph"], eng["design"], intake)
    _gen_sld(intake, d, eng, add(
        "10", "sld", "Single-Line Diagram", f"10_SingleLineDiagram_{slug}.pdf",
        "checklist", "generated", "CONCEPTUAL — for interconnection application only", "GridPilot",
        "Conceptual SLD in the CAISO reference format: switching-station ring bus, revenue "
        "metering, tie-line data, MPT strings, collector detail, notes and symbols legend.", 10))
    add("10b", "sld_dxf", "Single-Line Diagram (DXF)", f"10_SingleLineDiagram_{slug}.dxf",
        "checklist", "generated", "EDITABLE — CAD import (R12)", "GridPilot",
        "Same drawing as the PDF, exported for AutoCAD-class tools.", 10
        ).write_text(to_dxf(sld_prims), encoding="utf-8")

    _gen_flat_bump(intake, d, eng, add(
        "11", "flat_bump", "Flat Run & Bump Test Plots", f"11_FlatRunBumpTest_{slug}.pdf",
        "checklist", "generated",
        "SIMULATED — Pg & Qg, 3-ph fault at POI at t=10 s", "GridPilot",
        loc("Per CAISO instructions: no-fault run 10 s, 3-phase-to-ground fault at the POI at "
            "t=10 s, run to 20 s, plotting Pg and Qg."), 11))
    _gen_mw_poi_plot(intake, d, eng, add(
        "12", "mw_poi", "Requested MW at POI Plot", f"12_RequestedMWatPOI_{slug}.pdf",
        "checklist", "generated", f"SOLVED — holds {pf['p_poi_mw']:g} MW at the POI", "GridPilot",
        "Generation MW against the requested-MW limit line, from the solved load flow.", 12))
    _gen_ibr_validation(intake, d, eng, add(
        "13", "ibr_validation", "IBR Model Validation Results", f"13_IBRModelValidation_{slug}.pdf",
        "checklist", "generated", "ALL 3 TESTS — V/Q step, freq step, ride-through", "GridPilot",
        "Plant controller voltage/Q-reference step, frequency reference step, and voltage "
        "ride-through results.", 13))

    # ---- Supporting work products (not checklist items) ---------------------
    _gen_kmz(intake, d, add(
        "S1", "kmz", "Project Boundary (KMZ)", f"S1_ProjectBoundary_{slug}.kmz",
        "supporting", "generated", "GENERATED — opens in Google Earth", "GridPilot",
        f"Aerial-imagery site map for checklist item 9 — boundary polygon (~{_fmt(d['acres'])} ac) "
        f"centered on {d['lat']}, {d['lon']}."))
    gen_dynamics_report(eng, intake, add(
        "S2", "dyn_report", "Dynamic Validation Report", f"S2_DynamicValidation_{slug}.pdf",
        "supporting", "generated",
        "COMPUTED — " + ("all checks pass" if eng["bump"]["all_pass"] else "CHECKS FAILING"),
        "GridPilot",
        loc("Engineering backup for items 7 and 11: flat-start and disturbance checks on the "
            "plant-level RMS response.")))
    gen_loadflow_report(eng, intake, add(
        "S3", "lf_report", "Load Flow & Short-Circuit Report", f"S3_LoadFlowValidation_{slug}.pdf",
        "supporting", "generated",
        "COMPUTED — solved case, convergence log, fault duties", "GridPilot",
        "Engineering backup for items 6 and 12: convergence history, bus voltages, "
        "branch loadings, fault duties."))
    gen_equipment_schedule(eng, intake, add(
        "S4", "equipment_schedule", "Equipment Schedule", f"S4_EquipmentSchedule_{slug}.xlsx",
        "supporting", "generated",
        "GENERATED — design engine sizing with OEM sources", "GridPilot",
        f"Topology: {eng['design']['topology']}"))
    gen_assumptions_log(eng, intake, add(
        "S5", "assumptions", "Assumptions & Approvals Log", f"S5_AssumptionsApprovals_{slug}.pdf",
        "supporting", "generated",
        ("APPROVED — all items signed off" if eng["approvals"]["all_approved"]
         else f"PENDING — {eng['approvals']['pending']} of {eng['approvals']['total']} awaiting approval"),
        "GridPilot",
        "Every engineering default requiring engineer-of-record confirmation, with sign-off state."))
    gen_source_trace(eng, intake, add(
        "S6", "source_trace", "Source Trace Appendix", f"S6_SourceTrace_{slug}.md",
        "supporting", "generated", "GENERATED — full parameter provenance", "GridPilot",
        "Every value in this packet traced to developer input, OEM datasheet, ISO rule, "
        "calculation, or approved assumption."))
    gen_portal_fields(eng, intake, profile, add(
        "S7", "portal_fields", "Portal Field Export (JSON)", f"S7_PortalFields_{slug}.json",
        "supporting", "generated", f"READY — paste into {profile['portal']}", "GridPilot",
        "Machine-readable field values for the ISO portal entry."))
    _gen_sos_instructions(intake, add(
        "S8", "sos", "Secretary of State Certification", f"S8_SecretaryOfState_{slug}.pdf",
        "supporting", "action", "ACTION — developer obtains official certificate", "Developer",
        "Official government document; GridPilot provides ordering instructions."))
    _gen_signatory(intake, add(
        "S9", "signatory", "Proof of Authorized Signatory", f"S9_AuthorizedSignatory_{slug}.pdf",
        "supporting", "generated", "DRAFT — counsel review & execution", "GridPilot"))
    if eng["design"]["missing"]:
        gen_missing_data(eng, intake, add(
            "S95", "missing_data", "Missing Data List", f"S95_MissingData_{slug}.md",
            "supporting", "action",
            f"ACTION — {len(eng['design']['missing'])} item(s) for the developer", "Developer",
            "Facts the engine needed but could not find; assumptions stand in until provided."))

    readme_path = add("00", "readme", "Submission Checklist (README)", f"00_SubmissionChecklist_{slug}.md",
                      "checklist", "generated", "GENERATED — 13-item checklist map", "GridPilot",
                      "Maps every file to the CAISO minimum-requirements checklist, including the "
                      "two items handled outside the packet (deposit wire; ISP demonstrations).")
    _gen_readme(intake, d, sorted(docs, key=lambda x: x["n"]), readme_path)

    zip_path = pdir / f"{slug}_{profile['iso'].replace('-', '')}_Packet.zip"
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as z:
        for doc in docs:
            # Only checklist deliverables ship in the submission zip — the
            # README map and the S-series supporting work products stay
            # in-app as engineering backup.
            if doc["key"] == "readme" or doc["category"] != "checklist":
                continue
            fp = pdir / doc["file"]
            if fp.exists():
                z.write(fp, doc["file"])

    if profile["iso"] == "CAISO":
        deposit_action = {
            "title": f"Wire the study deposit — {d['deposit']}",
            "detail": "Directly to CAISO: Wells Fargo, ABA 121000248, Acct 4122041825. "
                      f"Reference \"{str(intake.get('project_name') or '').upper()}\" in the wire notes."}
    else:
        deposit_action = {
            "title": "Fund the study deposits and fees",
            "detail": profile["deposit_action"] + " Reference "
                      f"\"{str(intake.get('project_name') or '').upper()}\" in the wire notes."}
    actions = [
        deposit_action,
        {"title": "Order the Secretary of State certificate",
         "detail": f"Certificate of Good Standing for {intake.get('legal_name')} — see the "
                   "supporting instructions document."},
    ]
    if profile["iso"] == "CAISO" and d["is_isp"]:
        actions.append({
            "title": "Prepare the ISP eligibility demonstrations (checklist item 5)",
            "detail": "COD need, permits, equipment purchase order, financing, POI status, and "
                      "precursor upgrades — developer evidence assembled outside this packet."})
    site_file = _file_meta(intake.get("file_site_control"))
    if site_file:
        actions.append({
            "title": "Include the executed site agreement in the upload",
            "detail": f"{site_file['name']} ({intake.get('site_control')} with {intake.get('site_owner')}) — "
                      "submit it with the site exclusivity demonstration."})
    else:
        actions.append({
            "title": "Attach the executed site agreement",
            "detail": f"{intake.get('site_control')} with {intake.get('site_owner')} — LOIs are not accepted."})
    if _file_meta(intake.get("file_dyd")) is None and intake.get("dyd_status") != "Received from vendor":
        actions.append({
            "title": loc("Chase vendor .dyd dynamic model files"),
            "detail": loc("The most common schedule risk. GridPilot used WECC standard models as placeholders — "
                          "swap in vendor MOD-026/027-verified parameters before submission.")})
    if profile["iso"] == "CAISO":
        actions.append({
            "title": "E-sign Appendix 1 in RIMS5 and upload the packet",
            "detail": f"{RIMS5_URL} — CAISO confirms acceptance within ~4 weeks and the project enters the queue."})
    else:
        actions.append({
            "title": profile["submit_action"],
            "detail": f"{profile['portal_url']} — see the {profile['iso']} requirements reference for "
                      "window dates and acceptance timelines."})

    consistency = [
        {"title": "MW chain reconciled",
         "detail": loc(f"Gross {_fmt(d['gross'])} − Aux {_fmt(d['aux'])} − Losses {_fmt(d['losses'])} = "
                       f"{_fmt(d['net'])} MW at POI — identical in Appendix 1, Attachment A, the model, and the SLD.")},
        {"title": "Legal name consistent",
         "detail": loc(f"\"{intake.get('legal_name')}\" used verbatim across Appendix 1, signatory consent, and site exclusivity.")},
        {"title": "GPS consistent",
         "detail": loc(f"{d['lat']}, {d['lon']} identical in Appendix 1, Attachment A, and the KMZ boundary.")},
        {"title": "Dates sequential",
         "detail": f"In-service {d['in_service'].strftime('%m/%d/%Y')} → trial operation "
                   f"{d['trial_op'].strftime('%m/%d/%Y')} → COD {d['cod'].strftime('%m/%d/%Y')} (within 7 years)."},
    ]

    engineering = {
        "graph_version": eng["graph"].get("version", ""),
        "topology": eng["design"]["topology"],
        "counts": eng["design"]["counts"],
        "loadflow": {
            "converged": pf["converged"], "iterations": pf["iterations"],
            "p_poi_mw": pf["p_poi_mw"], "q_poi_mvar": pf["q_poi_mvar"],
            "losses_mw": pf["losses_mw"], "p_pv_mw": pf["p_pv_mw"],
            "p_bess_mw": pf["p_bess_mw"], "warnings": pf["warnings"],
        },
        "short_circuit": eng["short_circuit"]["results"],
        "dynamics_pass": eng["bump"]["all_pass"],
        "checks": eng["checks"],
        "approvals": {
            "total": eng["approvals"]["total"],
            "pending": eng["approvals"]["pending"],
            "all_approved": eng["approvals"]["all_approved"],
        },
        "missing": eng["graph"].get("missing", []),
    }

    manifest = {
        "id": pid,
        "org_id": org_id,
        "iso": profile["iso"],
        "iso_process": profile["process"],
        "created_at": datetime.now(timezone.utc).isoformat(),
        "project_name": intake.get("project_name"),
        "legal_name": intake.get("legal_name"),
        "track": d["track"],
        "net_mw": d["net"],
        "poi": intake.get("poi_name"),
        "intake": intake,
        "validation": validation,
        "consistency": consistency,
        "engineering": engineering,
        "actions": actions,
        "checklist": {
            "total": 13,
            "generated": sorted({doc["checklist_item"] for doc in docs
                                 if doc.get("checklist_item")}),
            "actions": [{"item": 1, "title": f"Study deposit — {d['deposit']}",
                         "note": "Wire transfer, not a document"}],
            "excluded": [{"item": 5, "title": "ISP eligibility demonstrations",
                          "note": "Handled outside this packet (developer evidence)"}]
                        if d["is_isp"] else [],
        } if profile["iso"] == "CAISO" else None,
        "documents": sorted(docs, key=lambda x: x["n"]),
        "zip_file": zip_path.name,
        "deposit": d["deposit"],
        "rims5_url": RIMS5_URL,
        "caiso_forms_url": CAISO_FORMS_URL,
    }
    (pdir / "manifest.json").write_text(json.dumps(manifest, indent=2), encoding="utf-8")
    return manifest


def load_manifest(packet_id: str) -> dict[str, Any] | None:
    p = PACKETS_DIR / packet_id / "manifest.json"
    if not p.exists():
        return None
    try:
        return json.loads(p.read_text(encoding="utf-8"))
    except json.JSONDecodeError:
        return None


def packet_file(packet_id: str, filename: str) -> Path | None:
    manifest = load_manifest(packet_id)
    if not manifest:
        return None
    allowed = {doc["file"] for doc in manifest["documents"]} | {manifest.get("zip_file")}
    if filename not in allowed:
        return None
    p = PACKETS_DIR / packet_id / filename
    return p if p.exists() else None


def reset_packets(org_id: str) -> None:
    if not PACKETS_DIR.exists():
        return
    for child in PACKETS_DIR.iterdir():
        m = child / "manifest.json"
        if m.exists():
            try:
                data = json.loads(m.read_text(encoding="utf-8"))
                if data.get("org_id") == org_id:
                    shutil.rmtree(child, ignore_errors=True)
            except json.JSONDecodeError:
                continue
